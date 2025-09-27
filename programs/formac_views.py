import json
import os
import random
import calendar, locale
import zipfile
from datetime import datetime, date

from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.files.storage import FileSystemStorage
from django.core.mail import send_mail, send_mass_mail
from django.db.models import Q
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import render, redirect

# Create your views here.
from django.urls import reverse
from django.utils import dateparse
from django.utils.text import slugify, phone2numeric
from django.utils.timezone import now

from programas.settings import MEDIA_URL, MEDIA_ROOT
from programs.models import Program, PhdStudent, Student, \
    FormationMember, InvestigationLine, PhdStudentTheme, CursStudentTheme, ColegStudentTheme, ProgramMember, \
    InvestigationProject, ProgramBackgrounds, MscStudent, ProgramEdition, MscStudentTheme, DipStudent, CursStudent, ColegStudent, \
    PostgMember, StudentFormationPlan, FormationMember, CursStudent, ColegStudent, FAQ, New
from programs.models import Document as FormacDoc


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def index(request):

    context={
        'members':FormationMember.objects.all(),
        'programs': Program.objects.filter(type__in=['curs','coleg']).order_by('-type'),
        'public_docs': FormacDoc.objects.filter(is_public=True),
        'news': New.objects.all().order_by('-date')[:5],
        'faqs': FAQ.objects.filter(title__icontains='[formac]').order_by('-date')[:5],
    }
    try:
        director = FormationMember.objects.get(charge='Director')
        context['director']= director

    except FormationMember.DoesNotExist:
        pass

    return render(request, 'programs/formac/formac_index.html', context)

@login_required
def home(request):
    try:
        formac_member=FormationMember.objects.get(user=request.user)
        context={
            'member':formac_member,
            'curs_programs': Program.objects.filter(type='curs').__len__(),
            'coleg_programs': Program.objects.filter(type='coleg').__len__(),
            
            'formac_members': FormationMember.objects.all().__len__(),
            
        }
        return render(request, 'programs/formac/formac_home.html', context)

    except FormationMember.DoesNotExist:
        return error_500(request,"Usted no es miembro de la Direccion de formación")




def error_500(request, error_message):
    context={

        'error_message':error_message,
    }
    try:
        context['member'] = FormationMember.objects.get(user=request.user)
    except Exception:
        pass

    return render(request,'programs/formac/formac_error_500.html', context)

def logout(request):
    logout(request)
    return HttpResponseRedirect(reverse('formac:formac_home'))

@login_required
def ajx_formac_students_by_program(request):
    response_data=[]
    labels = []
    data = []
    for program in Program.objects.filter(type='curs'):
        labels.append(program.slug.upper())
        data.append(program.student_set.all().__len__())

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    response_data.append(labels)
    response_data.append(data)




    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_formac_coleg_students_by_program(request):
    response_data=[]
    labels = []
    data = []
    for program in Program.objects.filter(type='coleg'):
        labels.append(program.slug.upper())
        data.append(program.colegstudent_set.filter(status='estudiante').__len__())

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    response_data.append(labels)
    response_data.append(data)




    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_formac_curs_students_by_program(request):
    response_data=[]
    labels = []
    data = []
    for program in Program.objects.filter(type='curs'):
        labels.append(program.slug.upper())
        data.append(program.cursstudent_set.filter(status='cursista').__len__())

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    response_data.append(labels)
    response_data.append(data)




    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )



@login_required
def programs(request, program_type):
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        context={
            'member': formac_member,
            'programs': Program.objects.filter(type=program_type),
            'program_type': program_type,

        }
        return render(request, "programs/formac/formac_programs_list.html", context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Dirección de formación')


@login_required
def members(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        context={
            'member': formac_member,
            'members': FormationMember.objects.all(),

        }
        return render(request, "programs/formac/formac_members_list.html", context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Dirección de Formación')



@login_required
def lines(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        context={
            'member': formac_member,
            'programs': Program.objects.all(),

        }
        return render(request, "programs/formac/formac_lines_list.html", context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Dirección de Formación')

@login_required
def lines(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        context={
            'member': formac_member,
            'programs': Program.objects.all(),

        }
        return render(request, "programs/formac/formac_lines_list.html", context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Dirección de formación')


@login_required
def program_students(request, program_slug, scope):
    program = Program.objects.get(slug = program_slug)
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        context = {
            'member': formac_member,
            'program': program,
            'scope':scope,
        }
        if program.type == 'curs':
            if scope == 'all':
                context['students'] = CursStudent.objects.filter(program=program)
            elif scope == 'requesters':
                context['students'] = CursStudent.objects.filter(program=program, status='solicitante')
            elif scope == 'approved':
                context['students'] = CursStudent.objects.filter(program=program, status='cursista')
            elif scope == 'graduated':
                context['students'] = CursStudent.objects.filter(program=program, status='graduado')
            else:
                return error_500(request,
                                 'El contexto "' + scope + '" no se admite. Debe ser "all", "approved", "requesters" o "graduated" ')
                
        elif program.type == 'coleg':
            if scope == 'all':
                context['students'] = ColegStudent.objects.filter(program=program)
            elif scope == 'requesters':
                context['students'] = ColegStudent.objects.filter(program=program, status='solicitante')
            elif scope == 'approved':
                context['students'] = ColegStudent.objects.filter(program=program, status='estudiante')
            elif scope == 'graduated':
                context['students'] = ColegStudent.objects.filter(program=program, status='graduado')
            else:
                return error_500(request,
                                 'El contexto "' + scope + '" no se admite. Debe ser "all", "approved", "requesters" o "graduated" ')
        

        return render(request, "programs/formac/formac_students_list.html", context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Dirección de Formación')


@login_required
def program_members(request, program_slug, scope):
    program = Program.objects.get(slug=program_slug)
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        context = {
            'member': formac_member,
            'program': program,
            'scope': scope,
        }

        if scope == 'all':
            context['program_members'] = ProgramMember.objects.filter(program=program)
        elif scope == 'committee':
            context['program_members'] = ProgramMember.objects.filter(Q(role='Director')| Q(role='Metodologo')| Q(role='Técnico'),program=program)
        elif scope == 'proffessor':
            context['program_members'] = ProgramMember.objects.filter(program=program, role='Profesor')
        elif scope == 'tuthor':
            context['program_members'] = ProgramMember.objects.filter(program=program, role='Tutor')
        else:
            return error_500(request,
                             'El contexto "' + scope + '" no se admite. Por favor acceda desde algún enlace válido del sitio ')

        return render(request, "programs/formac/formac_program_members_list.html", context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Dirección de Formación')

@login_required
def program_statistics(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        context={
            'member':formac_member,
            'program':program,
        }

        return render(request, 'programs/formac/formac_program_statistics.html', context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no tiene privilegios para acceder a esta página')


@login_required
def ajx_members_by_grade(request, program_slug):
    response_data=[]
    grades=[]
    data=[]

    program = Program.objects.get(slug=program_slug)
    for member in FormationMember.objects.filter(program=program):
        if not member.degree in grades:
            grades.append(member.degree)

    for degree in grades:
        data.append(FormationMember.objects.filter(program=program, degree=degree).__len__())

    response_data.append(grades)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_students_by_line(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    i=0
    for line in InvestigationLine.objects.filter(program=Program.objects.get(slug=program_slug)):
        i += 1
        labels.append(line.name.split()[0])
        if program.type == 'curs':
            data.append(CursStudentTheme.objects.filter(line=line).__len__())
        elif program.type == 'coleg':
            data.append(ColegStudentTheme.objects.filter(line=line).__len__())


    response_data.append(labels)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_students_by_edition(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for edition in ProgramEdition.objects.filter(program=program):
        labels.append('Edición '+str(edition.order))
        if program.type == 'curs':
            data.append(CursStudent.objects.filter(Q(status='cursista')|Q(status='graduado'), program=program, edition=edition).__len__())
        elif program.type == 'coleg':
            data.append(ColegStudent.objects.filter(Q(status='estudiante')|Q(status='graduado'), program=program, edition=edition).__len__())



    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_graduated_by_edition(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data = []
    labels = []
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for edition in ProgramEdition.objects.filter(program=program):
        labels.append('Edición ' + str(edition.order))
        if program.type == 'curs':
            data.append(
                CursStudent.objects.filter(program=program, edition=edition, status='graduado').__len__())
        elif program.type == 'coleg':
            data.append(
                ColegStudent.objects.filter(program=program, edition=edition, status='graduado').__len__())

    response_data.append(labels)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_students_by_gender(request, program_slug):

    response_data=[]
    program=Program.objects.get(slug=program_slug)
    if program.type == 'curs':
        response_data.append(CursStudent.objects.filter(program=program, gender='f').__len__())
        response_data.append(CursStudent.objects.filter(program=program, gender='m').__len__())
    elif program.type == 'coleg':
        response_data.append(ColegStudent.objects.filter(program=program, gender='f').__len__())
        response_data.append(ColegStudent.objects.filter(program=program,gender='m').__len__())
    
    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_students_by_age(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = ['<30 años','30-40','40-50','>50 años']
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    i=0
    if program.type == 'curs':
        data.append(CursStudent.objects.filter(program=program,birth_date__year__gte=now().year-30 ).__len__())
        data.append(CursStudent.objects.filter(program=program,birth_date__year__gte=now().year-40, birth_date__year__lt=now().year-30 ).__len__() )
        data.append(CursStudent.objects.filter(program=program,birth_date__year__gte=now().year-50, birth_date__year__lt=now().year-40 ).__len__() )
        data.append(CursStudent.objects.filter(program=program, birth_date__year__lte=now().year-50 ).__len__() )
    elif program.type == 'coleg':
        data.append(ColegStudent.objects.filter(program=program, birth_date__year__gte=now().year - 30).__len__())
        data.append(ColegStudent.objects.filter(program=program, birth_date__year__gte=now().year - 40,
                                              birth_date__year__lt=now().year - 30).__len__())
        data.append(ColegStudent.objects.filter(program=program, birth_date__year__gte=now().year - 50,
                                              birth_date__year__lt=now().year - 40).__len__())
        data.append(ColegStudent.objects.filter(program=program, birth_date__year__lte=now().year - 50).__len__())
    
    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def formac_by_line_projects_list(request, line_id):
    line = InvestigationLine.objects.get(pk=line_id)
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        context={
            'member':formac_member,
            'line': line,
            'projects': InvestigationProject.objects.filter(line=line),
        }
        return render(request, 'programs/formac/formac_by_line_projects_list.html', context)


    except FormationMember.DoesNotExist:
        return error_500(request,  'Usted no tiene privilegios para acceder a esta página')


@login_required
def create_formac_member(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)

        if formac_member.charge == 'Director':
            if request.method == 'POST':
                try:
                    user = User.objects.get(email=request.POST['email'])
                    try:
                        formac_member = FormationMember.objects.get(user=user)

                        return error_500(request, 'Ya existe un miembro de la direccion con ese correo' )
                    except FormationMember.DoesNotExist:
                        new_formac_member = FormationMember(
                            user=user,
                            charge=request.POST['charge'],
                            grade=request.POST['grade'],
                            phone=request.POST['phone'],
                            gender=request.POST['gender']

                        )

                        new_formac_member.save()

                        try:

                            picture = request.FILES['picture']
                            new_formac_member.picture = picture
                            new_formac_member.save()

                        except:
                            print( Exception)


                        return HttpResponseRedirect(reverse('formac:members'))

                except User.DoesNotExist:
                    passwd = 'FORMAC' + str(random.randint(1000000, 9999999))
                    user = User.objects.create_user(
                        request.POST['email'],
                        request.POST['email'],
                        passwd,  # Cambiar despues por contrase;a generada

                    )
                    user.first_name = request.POST['name']
                    user.last_name = request.POST['surename']
                    user.save()

                    new_formac_member = FormationMember(
                        user=user,
                        charge=request.POST['charge'],
                        grade=request.POST['grade'],
                        phone=request.POST['phone'],
                        gender=request.POST['gender']

                    )
                    new_formac_member.save()

                    try:

                        picture = request.FILES['picture']
                        new_formac_member.picture = picture
                        new_formac_member.save()

                    except:
                        print(Exception)


                    return HttpResponseRedirect(reverse('formac:members'))


            else:
                context = {
                    'member': formac_member,

                }
                return render(request, 'programs/formac/formac_create_member.html', context)
        else:
            return error_500(request, 'Usted no tiene privilegios para agregar nuevos miembros a la Dir. de Formación ')

    except FormationMember.DoesNotExist:
        return error_500(request,'Usted no tiene privilegios para acceder a esta página')


@login_required
def edit_formac_member(request, member_id):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        if formac_member.charge == 'Director':
            if request.method == 'POST':
                pmember = FormationMember.objects.get(pk=member_id)
                pmember.user.first_name = request.POST['name']
                pmember.user.last_name = request.POST['surename']
                pmember.user.email = request.POST['email']
                pmember.user.save()

                pmember.charge = request.POST['charge']
                pmember.gender = request.POST['gender']
                pmember.phone = request.POST['phone']
                pmember.grade = request.POST['grade']

                try:
                    pmember.picture = request.FILES['picture']
                except:
                    print(Exception)

                pmember.save()

                return HttpResponseRedirect(reverse('formac:members'))
            else:
                context={
                    'member': formac_member,
                    'pmember': FormationMember.objects.get(pk=member_id),
                }

                return render(request, 'programs/formac/formac_edit_member.html', context)
        else:
            return error_500(request, 'Usted no tiene privilegios para editar miembros de la Dir. de Formación ')
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no tiene privilegios para acceder a esta página')

@login_required
def ajx_delete_member(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        if formac_member.charge == 'Director':
            if request.method == 'POST':
                pmember = FormationMember.objects.get(pk=request.POST['member_id'])
                pmember.delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            else:
                return HttpResponse(
                    json.dumps([{'deleted': 2}]),
                    content_type="application/json"
                )


        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    except FormationMember.DoesNotExist:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )


@login_required
def ajx_delete_document(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        if formac_member.charge == 'Director':
            if request.method == 'POST':
                document = FormacDoc.objects.get(pk=request.POST['document_id'])
                try:
                    document.doc.delete()
                except:
                    pass

                document.delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            else:
                return HttpResponse(
                    json.dumps([{'deleted': 2}]),
                    content_type="application/json"
                )


        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    except FormationMember.DoesNotExist:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_formac_program_last_years_requests(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []
    data_1=[]
    data_2=[]

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for i in range(now().year-4,now().year+1):
        labels.append(i)
        if program.type == 'curs':
            data_1.append(CursStudent.objects.filter(program=program, request_date__year=i).__len__())
            data_2.append(CursStudent.objects.filter(program=program, init_date__year=i).__len__())
        elif program.type == 'coleg':
            data_1.append(ColegStudent.objects.filter(program=program, request_date__year=i).__len__())
            data_2.append(ColegStudent.objects.filter(program=program, init_date__year=i).__len__())
        
    data.append(data_1)
    data.append(data_2)

    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_formac_program_by_year_requests(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    year=request.POST['year']
    response_data=[]
    labels = []
    data = []
    data_1=[]
    data_2=[]
    meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 7: "Julio",
             8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for i in range(1,13):
        labels.append(meses[i])
        if program.type == 'curs':
            data_1.append(CursStudent.objects.filter(program=program, request_date__year=year,request_date__month=i).__len__())
            data_2.append(CursStudent.objects.filter(program=program, init_date__year=year,init_date__month=i).__len__())
        elif program.type == 'coleg':
            data_1.append(ColegStudent.objects.filter(program=program, request_date__year=year, request_date__month=i).__len__())
            data_2.append(ColegStudent.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())
        

    data.append(data_1)
    data.append(data_2)

    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_formac_usr_exists(request):

    if request.method=='POST':

        try:
            user= User.objects.get(email=request.POST['email'])
            try:
                formac_member = FormationMember.objects.get(user=user)

                return HttpResponse(
                    json.dumps([{'exists': 2}]),
                    content_type="application/json"
                )
            except FormationMember.DoesNotExist:
                return HttpResponse(
                    json.dumps([{'exists': 1}]),
                    content_type="application/json"
                )


        except User.DoesNotExist:
            return HttpResponse(
                json.dumps([{'exists': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'exists': 0}]),
            content_type="application/json"
        )

def formac_member_picture(request, member_id):
    fs = FileSystemStorage()
    # filename = Papers.objects.get(pk=paper_id).file_url +  str(Papers.objects.get(pk=paper_id).file)
    filename = FormationMember.objects.get(pk=member_id).picture.url
    if fs.exists(filename):
        with fs.open(filename) as img:
            response = HttpResponse(img, content_type='image/jpeg')
            return response
    else:
        return HttpResponse('Error')


@login_required
def formac_new_document(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        if request.method == 'POST':
            new_doc = FormacDoc(
                year=request.POST['year'],
                month=request.POST['month'],
                description=request.POST['description'],
                type=request.POST['type'],
                doc=request.FILES['doc'],
            )
            try:
                if request.POST['is_public'] == 'on':
                    new_doc.is_public = True

            except:
                pass
            new_doc.save()

            return HttpResponseRedirect(reverse('formac:documents', args=['all']))
        else:
            years=[]
            for year in range(now().year-4, now().year+1):
                years.append(year)
            context={
                'member':formac_member,
                'years':years,
                'current_year':now().year,
            }

            return render(request, 'programs/formac/formac_create_document.html', context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no tiene privilegios para acceder a esta página')

@login_required
def formac_edit_document(request, document_id):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        if request.method == 'POST':
            doc = FormacDoc.objects.get(pk=document_id)
            old_year = doc.year
            old_month = doc.month
            old_type = doc.type

            doc.year = request.POST['year']
            doc.month = request.POST['month']
            doc.type = request.POST['type']
            doc.description = request.POST['description']
            doc.save()

            try:
                if request.POST['is_public']== 'on':
                    doc.is_public = True
                    doc.save()
            except:
                doc.is_public = False
                doc.save()

            try:
                doc_file = request.FILES['doc']
                doc_ext = doc_file.name.split('.')[doc_file.name.split('.').__len__() - 1]
                doc_name = 'formac/docs/{0}/{1}_{2}_{3}.{4}'.format(doc.year, doc.type.capitalize(), doc.year,doc.month, doc_ext)



                fs = FileSystemStorage()

                filename = fs.save(doc_name, doc_file)
                old_doc = doc.doc.url

                # doc.doc.delete()
                doc.doc = filename
                doc.save()
                try:
                    fd = FileSystemStorage()
                    fd.delete(old_doc)
                    print('Archivo eliminado exitosamente')
                except:
                    print('No se pudo eliminar el archivo:', old_doc)


            except:
                print('No viene archivo, pero hay que mover el que esta')
                if str(old_year) != doc.year or old_month != doc.month or old_type != doc.type:
                    initial_path = doc.doc.url

                    print('Initial path:', initial_path)
                    doc_ext = initial_path.split('.')[initial_path.split('.').__len__() - 1]
                    old_name = doc.doc.name

                    doc.doc.name = '/formac/docs/{1}/{0}_{1}_{2}.{3}'.format(doc.type.capitalize(), doc.year, doc.month,
                                                                            doc_ext)
                    doc.save()
                    new_path = MEDIA_ROOT + doc.doc.name
                    print('New path:', new_path)
                    try:
                        if not os.path.exists(MEDIA_ROOT + '/formac/docs/{0}'.format(doc.year)):
                            os.makedirs(MEDIA_ROOT + '/formac/docs/{0}'.format(doc.year))
                        fn = FileSystemStorage(MEDIA_ROOT)
                        if fn.exists(doc.doc.name[1:]):
                            new_name = '/'+ fn.get_available_name(doc.doc.name[1:])
                            doc.doc.name = new_name
                            doc.save()
                            new_path = MEDIA_ROOT+doc.doc.name
                            os.rename(initial_path, new_path)

                        else:
                            os.rename(initial_path, new_path)
                    except:
                        print('Exception:', 'No se pudo mover el archivo')
                        # doc.doc.name = old_name
                        # doc.save
                else:
                    pass

            return HttpResponseRedirect(reverse('formac:documents', args=['all']))

        else:
            years=[]
            for year in range(now().year-4, now().year+1):
                years.append(year)
            context={
                'member':formac_member,
                'years':years,
                'document':FormacDoc.objects.get(pk=document_id),
            }

            return render(request, 'programs/formac/formac_edit_document.html', context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no tiene privilegios para acceder a esta página')


@login_required
def formac_documents(request, scope):
    print(scope)
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        context={
            'member':formac_member,

        }
        if scope == 'all':
            context['documents']= FormacDoc.objects.all()
        elif scope == 'brief':
            context['documents']=  FormacDoc.objects.filter(type='acta')
        elif scope == 'reports':
            context['documents']=  FormacDoc.objects.filter(type='informe')
        elif scope == 'resolutions':
            context['documents']=  FormacDoc.objects.filter(type='resolucion')
        else:
            return error_500(request,'El contexto es incorrecto. Intente acceder desde uno de los enlaces publicados en el sitio.')



        return render(request, 'programs/formac/formac_documents_list.html', context)
    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no tiene privilegios para acceder a esta página')


@login_required
def formac_document_view(request, document_id):

    doc = FormacDoc.objects.get(pk=document_id)

    fs = FileSystemStorage()

    filename =doc.doc.url

    if fs.exists(filename):
        file_name= filename.split('/')[filename.split('/').__len__()-1]
        doc_ext =filename.split('.')[filename.split('.').__len__()-1]

        if doc_ext =='doc' or doc_ext=='docx' or doc_ext == 'odt':

            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/doc')
                response['Content-Disposition'] =  "inline; filename=" +'"'+  file_name+'"'

                return response

        elif doc_ext == 'pdf' :
            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/pdf')
                response['Content-Disposition'] = "inline; filename=" +'"'+ file_name + '"'

                return response

    else:


        return error_500(request, 'No existe el archivo solicitado')

def formac_public_document_view(request, document_id):

    doc = FormacDoc.objects.get(pk=document_id)

    fs = FileSystemStorage()

    filename =doc.doc.url

    if doc.is_public:

        if fs.exists(filename) :
            file_name= filename.split('/')[filename.split('/').__len__()-1]
            doc_ext =filename.split('.')[filename.split('.').__len__()-1]

            if doc_ext =='doc' or doc_ext=='docx' or doc_ext == 'odt':

                with fs.open(filename) as brief:
                    response = HttpResponse(brief, content_type='application/doc')
                    response['Content-Disposition'] =  "inline; filename=" +'"'+  file_name+'"'

                    return response

            elif doc_ext == 'pdf' :
                with fs.open(filename) as brief:
                    response = HttpResponse(brief, content_type='application/pdf')
                    response['Content-Disposition'] = "inline; filename=" +'"'+ file_name + '"'

                    return response

        else:
            return error_500(request, 'No existe el archivo solicitado')
    else:

        return error_500(request, 'El documento solicitado no es publico')


@login_required
def docx_formac_report(request, scope):
    try:
        formac_member =FormationMember.objects.get(user=request.user)
        document = Document()
        docname = ''
        if scope == 'current_year':
            docname = 'Reporte_Formacion_' + str(now().year) + '.docx'

            document.add_heading('Resumen de los cursos', level=1)
            document.add_heading('Año '+ str(now().year), level=2)
            document.add_heading('Cursos', level=2)
            if Program.objects.filter(type='curs'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre'
                hdr_cells[1].text = 'Coordinador'
                hdr_cells[2].text = 'Email'

                for program in Program.objects.filter(type='curs'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = program.full_name
                    row_cells[1].text = str(ProgramMember.objects.get(program=program, role='Coordinador'))
                    row_cells[2].text = program.email
            else:
                document.add_heading('Aun no se registran cursos en este sitio', level=3)

            document.add_heading('Curso Colegio', level=2)
            if Program.objects.filter(type='coleg'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre'
                hdr_cells[1].text = 'Coordinador'
                hdr_cells[2].text = 'Email'

                for program in Program.objects.filter(type='coleg'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = program.full_name
                    row_cells[1].text = str(ProgramMember.objects.get(program=program, role='Coordinador'))
                    row_cells[2].text = program.email
            else:
                document.add_heading('Aun no se registran cursos de colegio en este sitio', level=3)

            
            document.add_heading('Detalles por programas', level=2)
            document.add_heading('Cursos', level=3)

            for program in Program.objects.filter(type='curs'):
                document.add_heading('Cursos de '+str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='cursista', init_date__year=now().year):
                    table = document.add_table(rows=1, cols=5)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[2].text = 'Fecha de ingreso'
                    hdr_cells[3].text = 'Edición'
                    hdr_cells[4].text = 'País'


                    for student in CursStudent.objects.filter(program=program, status='cursista', init_date__year=now().year):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        try:
                            row_cells[1].text = str(student.init_date)
                        except student.graduate_date.DoesNotExist:
                            row_cells[1].text = 'No declarada'
                        row_cells[2].text = str(student.init_date)
                        row_cells[3].text = str (student.edition)
                        row_cells[4].text = str(student.country).capitalize()
                else:
                    document.add_heading('No se registran nuevos ingresos al programa este año', level=3)


                document.add_heading('Solicitantes de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='solicitante', request_date__year=now().year):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de solicitud'
                    hdr_cells[2].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='solicitante', request_date__year=now().year):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.request_date)
                        row_cells[2].text = str(student.country).capitalize()


                else:
                    document.add_heading('No se registran solicitudes de ingreso al programa este año', level=3)

                document.add_heading('Graduados de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='graduado',
                                          graduate_date__year=now().year):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de egreso'
                    hdr_cells[2].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='graduado',
                                                          graduate_date__year=now().year):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.graduate_date)
                        row_cells[2].text = str(student.country).capitalize()


                else:
                    document.add_heading('No se registran graduados del programa este año', level=3)

            document.add_heading('Colegio', level=3)

            for program in Program.objects.filter(type='coleg'):
                document.add_heading('Estudiante de '+str(program.full_name), level=3)
                if ColegStudent.objects.filter(program=program, status='estudiante', init_date__year=now().year):
                    table = document.add_table(rows=1, cols=4)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de ingreso'
                    hdr_cells[2].text = 'Edición'
                    hdr_cells[3].text = 'Centro de Procedencia'


                    for student in ColegStudent.objects.filter(program=program, status='estudiante', init_date__year=now().year):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.init_date)
                        row_cells[2].text = str(student.edition)
                        row_cells[3].text= str(student.center)
                else:
                    document.add_heading('No se registran nuevos ingresos al programa este año', level=3)


                document.add_heading('Graduados de ' + str(program.full_name), level=3)
                if ColegStudent.objects.filter(program=program, status='graduado', request_date__year=now().year):
                    table = document.add_table(rows=1, cols=4)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de egreso'
                    hdr_cells[2].text = 'Edición'
                    hdr_cells[3].text = 'Centro de Procedencia'

                    for student in ColegStudent.objects.filter(program=program, status='graduado', request_date__year=now().year):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.graduate_date)
                        row_cells[2].text =str(student.edition)
                        row_cells[3].text = str(student.center).capitalize()


                else:
                    document.add_heading('No se registran graduados del programa este año', level=3)

            
        elif scope == 'last_year':
            docname = 'Reporte_Formacion_' + str(now().year - 1) + '.docx'
            document.add_heading('Resumen de los programas de formación', level=1)
            document.add_heading('Año '+ str(now().year-1), level=2)
            document.add_heading('Programas de formación', level=2)

            if Program.objects.filter(type='curs'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre'
                hdr_cells[1].text = 'Coordinador(es)'
                hdr_cells[2].text = 'Email'

                coordinators = ''

                for program in Program.objects.filter(type='curs'):
                    for coordinator in ProgramMember.objects.filter(program=program, role='Coordinador'):
                        coordinators = coordinators + ',' + coordinator.user.get_full_name()

                    row_cells = table.add_row().cells
                    row_cells[0].text = program.full_name
                    row_cells[1].text = coordinators
                    row_cells[2].text = program.email
            else:
                document.add_heading('Aun no se registran cursos en este sitio', level=3)

            document.add_heading('Colegio', level=2)

            if Program.objects.filter(type='coleg'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre'
                hdr_cells[1].text = 'Coordinador(es)'
                hdr_cells[2].text = 'Email'

                for program in Program.objects.filter(type='coleg'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = program.full_name
                    row_cells[1].text = coordinators
                    row_cells[2].text = program.email
            else:
                document.add_heading('Aun no se registran programas de maestría en este sitio', level=3)

            document.add_heading('Detalles por programas', level=2)
            document.add_heading('Cursos', level=3)

            for program in Program.objects.filter(type='curs'):
                document.add_heading('Cursos de '+str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='cursista', init_date__year=now().year-1):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de ingreso'
                    hdr_cells[2].text = 'País'


                    for student in CursStudent.objects.filter(program=program, status='cursista', init_date__year=now().year-1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        try:
                            row_cells[1].text = str(student.graduate_date)
                        except student.graduate_date.DoesNotExist:
                            row_cells[1].text = 'No declarada'

                        row_cells[2].text = str(student.init_date)
                        row_cells[3].text = str(student.country).capitalize()
                else:
                    document.add_heading('No se registran nuevos ingresos al programa este año', level=3)


                document.add_heading('Solicitantes de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='solicitante', request_date__year=now().year-1):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de solicitud'
                    hdr_cells[2].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='solicitante', request_date__year=now().year-1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.request_date)
                        row_cells[2].text = str(student.country).capitalize()


                else:
                    document.add_heading('No se registran solicitudes de ingreso al programa este año', level=3)

                document.add_heading('Graduados de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='graduado',
                                          graduate_date__year=now().year-1):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de egreso'
                    hdr_cells[2].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='graduado',
                                                          graduate_date__year=now().year-1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.graduate_date)
                        row_cells[2].text = str(student.country).capitalize()


                else:
                    document.add_heading('No se registran graduados del programa este año', level=3)

            document.add_heading('Colegio', level=3)

            for program in Program.objects.filter(type='coleg'):
                document.add_heading('Estudiante de '+str(program.full_name), level=3)
                if ColegStudent.objects.filter(program=program, status='estudiante', init_date__year=now().year-1):
                    table = document.add_table(rows=1, cols=5)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[2].text = 'Fecha de ingreso'
                    hdr_cells[3].text = 'Edición'
                    hdr_cells[4].text = 'Centro de Procedencia'


                    for student in ColegStudent.objects.filter(program=program, status='estudiante', init_date__year=now().year-1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        try:
                            row_cells[1].text = str(student.graduate_date)
                        except StudentFormationPlan.DoesNotExist:
                            row_cells[1].text = 'No declarada'

                        row_cells[2].text = str(student.init_date)
                        row_cells[3].text = str(student.edition)
                        row_cells[4].text = str(student.center).capitalize()
                else:
                    document.add_heading('No se registran nuevos ingresos al programa este año', level=3)


                document.add_heading('Graduados de ' + str(program.full_name), level=3)
                if ColegStudent.objects.filter(program=program, status='solicitante', graduate_date__year=now().year-1):
                    table = document.add_table(rows=1, cols=4)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de egreso'
                    hdr_cells[2].text = 'Edición'
                    hdr_cells[3].text = 'Centro de Procedencia'

                    for student in ColegStudent.objects.filter(program=program, status='solicitante', graduate_date__year=now().year-1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.graduate_date)
                        row_cells[2].text = str (student.edition)
                        row_cells[3].text = str(student.center).capitalize()


                else:
                    document.add_heading('No se registran graduados del programa este año', level=3)

            
        elif scope == 'five_years':
            document.add_heading('Resumen de los programas de formación', level=1)
            document.add_heading('Período '+ str(now().year-6) + '-' + str(now().year-1), level=2)
            document.add_heading('Programas de Formación', level=2)
            docname = 'Reporte_Formacion_' + str(now().year-6) + '-' + str(now().year-1) +  '.docx'

            document.add_heading('Programas de Formación', level=2)

            if Program.objects.filter(type='curs'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre'
                hdr_cells[1].text = 'Coordinador'
                hdr_cells[2].text = 'Email'

                for program in Program.objects.filter(type='curs'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = program.full_name
                    row_cells[1].text = str(ProgramMember.objects.get(program=program, role='Coordinador'))
                    row_cells[2].text = program.email
            else:
                document.add_heading('Aun no se registran programas doctorales en este sitio', level=3)

            document.add_heading('Colegio', level=2)

            if Program.objects.filter(type='coleg'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre'
                hdr_cells[1].text = 'Coordinador'
                hdr_cells[2].text = 'Email'

                for program in Program.objects.filter(type='coleg'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = program.full_name
                    row_cells[1].text = str(ProgramMember.objects.get(program=program, role='Coordinador'))
                    row_cells[2].text = program.email
            else:
                document.add_heading('Aun no se registran programas de maestría en este sitio', level=3)

            
            document.add_heading('Detalles por programas', level=2)
            document.add_heading('Cursos', level=3)

            for program in Program.objects.filter(type='curs'):
                document.add_heading('Cursistas de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='cursista',
                                          init_date__gt=date(now().year - 6,1,1), init_date__lt=date(now().year - 1,12,31)):
                    table = document.add_table(rows=1, cols=4)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[2].text = 'Fecha de ingreso'
                    hdr_cells[3].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='cursista',
                                                          init_date__gt=date(now().year - 6,1,1), init_date__lt=date(now().year - 1,12,31)):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        try:
                            row_cells[1].text = str(student.graduate_date)
                        except student.graduate_date.DoesNotExist:
                            row_cells[1].text = 'No declarada'

                        row_cells[2].text = str(student.init_date)
                        row_cells[3].text = str(student.country).capitalize()
                else:
                    document.add_heading('No se registran ingresos al programa en el periodo', level=3)

                document.add_heading('Solicitantes de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='solicitante',
                                          request_date__gt=date(now().year - 6,1,1), request_date__lt=date(now().year - 1,12,31)):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Fecha de solicitud'
                    hdr_cells[2].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='solicitante',
                                                          request_date__gt=date(now().year - 6,1,1), request_date__lt=date(now().year - 1,12,31)):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.request_date)
                        row_cells[2].text = str(student.country).capitalize()


                else:
                    document.add_heading('No se registran solicitudes de ingreso al programa en el periodo', level=3)

                document.add_heading('Graduados de ' + str(program.full_name), level=3)
                if CursStudent.objects.filter(program=program, status='graduado',
                                          graduate_date__gt=date(now().year - 6,1,1), graduate_date__lt=date(now().year - 1,12,31)):
                    table = document.add_table(rows=1, cols=3)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Año de egreso'
                    hdr_cells[2].text = 'País'

                    for student in CursStudent.objects.filter(program=program, status='graduado',
                                                          graduate_date__gt=date(now().year - 6, 1, 1),
                                                          graduate_date__lt=date(now().year - 1, 12, 31)):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.graduate_date.year)
                        row_cells[2].text = str(student.country).capitalize()


                else:
                    document.add_heading('No se registran graduados del programa en el periodo', level=3)

            document.add_heading('Colegio', level=3)

            for program in Program.objects.filter(type='coleg'):
                document.add_heading('Estudiante de ' + str(program.full_name), level=3)
                if ColegStudent.objects.filter(program=program, status='estudiante', init_date__gt=date(now().year - 6,1,1), init_date__lt=date(now().year - 1,12,31)):
                    table = document.add_table(rows=1, cols=5)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[2].text = 'Fecha de ingreso'
                    hdr_cells[3].text = 'Edición'
                    hdr_cells[4].text = 'Centro de Procedencia'

                    for student in ColegStudent.objects.filter(program=program, status='estudiante',
                                                             init_date__gt=date(now().year - 6,1,1), init_date__lt=date(now().year - 1,12,31)):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        try:
                            row_cells[1].text = str(student.graduate_date)
                        except student.graduate_date.DoesNotExist:
                            row_cells[1].text = 'No declarada'

                        row_cells[2].text = str(student.init_date)
                        row_cells[3].text = str(student.center).capitalize()
                else:
                    document.add_heading('No se registran nuevos ingresos al programa este año', level=3)

                document.add_heading('Graduados de ' + str(program.full_name), level=3)
                if ColegStudent.objects.filter(program=program, status='graduado', graduate_date__gt=date(now().year - 6,1,1), graduate_date__lt=date(now().year - 1,12,31)):
                    table = document.add_table(rows=1, cols=5)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Nombre y apellidos'
                    hdr_cells[1].text = 'Año de egreso'
                    hdr_cells[3].text = 'Edición'
                    hdr_cells[4].text = 'Centro de Procedencia'

                    for student in ColegStudent.objects.filter(program=program, status='graduado',
                                                             graduate_date__gt=date(now().year - 6,1,1), graduate_date__lt=date(now().year - 1,12,31)):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.get_full_name())
                        row_cells[1].text = str(student.graduate_date.year)
                        row_cells[2].text =str (student.edition)
                        row_cells[3].text = str(student.center).capitalize()


                else:
                    document.add_heading('No se registran graduados del programa este año', level=3)

            

        # docpath = MEDIA_ROOT + '/formac/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        docpath = MEDIA_ROOT + '/formac/' + docname

        document.save(docpath)

        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="' + docname + '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')



    except FormationMember.DoesNotExist:
        return error_500(request, 'Usted no es miembro de la Direccion de Formación')
@login_required
def ajx_students_by_center(request, program_slug):
    from django.db.models import Count
    from django.http import JsonResponse
    from django.shortcuts import get_object_or_404
    from .models import Program, CursStudent, ColegStudent
    
    program = get_object_or_404(Program, slug=program_slug)
    
    centers = CursStudent.objects.filter(
        program=program
    ).exclude(center__isnull=True).exclude(center__exact='').values(
        'center'
    ).annotate(
        total=Count('id')
    ).order_by('-total')

    centers = ColegStudent.objects.filter(
        program=program
    ).exclude(center__isnull=True).exclude(center__exact='').values(
        'center'
    ).annotate(
        total=Count('id')
    ).order_by('-total')
    
    labels = [c['center'] for c in centers]
    data = [c['total'] for c in centers]
    
    return JsonResponse({
        'labels': labels,
        'data': data,
        'total': sum(data)
    })


# Preguntas
@login_required
def create_faq(request):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        
    except FormationMember.DoesNotExist:
        return error_500(request, None, 'Usted no está vinculado a ningún programa de formación.')

    if request.method == 'POST':
        faq = FAQ(
            # program=program, 
            title='[formac]'+ request.POST.get('faq.title', '').strip(),
            content=request.POST.get('faq_content', '').strip()
        )
        if 'img_faq' in request.FILES:
            faq.img = request.FILES['img_faq']
        faq.save()
        return redirect(reverse('formac:faq_list'))

    context = {
        
        'member': formac_member
    }
    return render(request, 'programs/formac/create_faq.html', context)
@login_required
def faq_list(request):
    faqs = FAQ.objects.filter(title__icontains='[formac]')
    context={'faqs': faqs}
            
    return render(request, 'programs/formac/faq_list.html', context)

@login_required
def read_faq(request, faq_id):

        context={
            'faq': FAQ.objects.get(pk=faq_id)
        }       
        return render(request, 'programs/formac/read_faq.html', context)    


@login_required
def edit_faq(request, faq_id):
    try:
        formac_member = FormationMember.objects.get(user=request.user)
        
    except FormationMember.DoesNotExist:
        return error_500(request, None, 'Usted no está vinculado a ningún programa de formación.')
    faq = FAQ.objects.get(pk=faq_id)
    if request.method == 'POST':
        faq.title = request.POST['faq.title']
        faq.content = request.POST['faq.content']
        try:
            faq.img = request.FILES['img_faq']
        except:
            pass
        faq.save()
        return HttpResponseRedirect(reverse('formac:read_faq', args=[faq_id]))
    else:
        context={
            'member': formac_member,
            'faq':faq,
            
        }
        return render(request, 'programs/formac/edit_faq.html', context)
    
    
@login_required
def ajx_delete_faq(request, faq_id):
    if request.method == 'POST':
        faq_id = request.POST.get('faq_id', '').strip()
        if not faq_id.isdigit():
            return JsonResponse({'deleted': 0})

        try:
            FAQ.objects.get(pk=faq_id).delete()
            return JsonResponse({'deleted': 1})
        except FAQ.DoesNotExist:
            return JsonResponse({'deleted': 0})
    else:
        return JsonResponse({'deleted': 2})


