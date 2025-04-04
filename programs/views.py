import csv as csv
import io
import json
import os
import random
import secrets
import shutil
import string
import textwrap
import zipfile
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet
styleSheet = getSampleStyleSheet()
style = styleSheet['BodyText']

from django.contrib import messages
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.contrib.messages import success
from django.core.files.storage import FileSystemStorage
from django.core.mail import send_mail
from django.db.models import Q
from django.forms import forms
from django.http import HttpResponse, HttpResponseNotFound, HttpResponseRedirect, Http404, FileResponse
from django.shortcuts import render, get_object_or_404
from django.template.defaultfilters import upper

# Create your views here.
from django.urls import reverse
from django.utils.text import slugify
from django.utils.timezone import now

from programas.settings import MEDIA_ROOT, STATIC_ROOT, DIR_COM_EMAIL, MEDIA_URL, NO_REPLY_EMAIL
from programs.forms import FileUploadForm, AnnouncementForm, PhdStudentThesisForm, PhdThesisCommentForm
from programs.models import Program, ProgramInitRequirements, PhdStudent, Student, \
    ProgramMember, ProgramFinishRequirements, InvestigationLine, PhdStudentTheme, \
    InvestigationProject, ProgramBackgrounds, MscStudent, ProgramEdition, MscStudentTheme, DipStudent, Tuthor, \
    ProgramBrief, CGCBrief, CNGCBrief, Course, CourseEvaluation, CourseProfessor, StudentFormationPlan, \
    FormationPlanActivities, InnerAreas, ProgramDocument, ProgramFileDoc, StudentFileDocument, Message, CGCDocument, \
    ProgramSpeciality, New, MessageSended, Requester, PhdStudentThesis, PhdAnnouncement, PhdThesisComment, \
    PhdDefenseCourtMember, FAQ
from programs.templatetags.extra_tags import finish_requirements_accomplished, \
    init_requirements_accomplished
from programs.utils import user_is_program_cs, user_is_program_member, utils_send_email, user_is_program_student, \
    create_new_tuthor, request_send_email
from docx import Document
from PyPDF2 import PdfFileWriter, PdfFileReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas

def index(request, program_slug):
    program = get_object_or_404(Program, slug=program_slug)
    if not request.user.is_authenticated :
        # program = Program.objects.get(slug=program_slug)
        # faqs = FAQ.objects.filter(program=program)
        context = {
            'program': program,
            'lines': InvestigationLine.objects.filter(program=program),
            'public_docs':ProgramDocument.objects.filter(program=program, is_public=True),
            'news': New.objects.filter(program=program).order_by('-date')[:5],
            'faqs': FAQ.objects.filter(program=program).order_by('-date')[:5],
        }
        if program.self_request:
            context['init_requirenments']=program.programfiledoc_set.filter(type='student')

        if program.type == 'phd':
            context['announcements'] = PhdAnnouncement.objects.filter( phd_student__student__program = program, date__gte = now() )
            return render(request, 'programs/phd_index.html', context)
        elif program.type == 'msc':
            return render(request, 'programs/msc_index.html', context)
        elif program.type == 'dip':
            return render(request, 'programs/dip_index.html', context)

    else:
        return HttpResponseRedirect(reverse('programs:home', args=[program_slug]))

@login_required
def home(request, program_slug):

    try:
        program = Program.objects.get(slug=program_slug)

        if request.user.is_authenticated:
            if program.type == 'phd':

                context = {
                    'program': program,
                    'requesters': PhdStudent.objects.filter(student__program=program, status='solicitante').__len__(),
                    'doctorands': PhdStudent.objects.filter(student__program=program, status='doctorando').__len__(),
                    'graduated': PhdStudent.objects.filter(student__program=program, status='graduado').__len__(),
                    'last_requesters': PhdStudent.objects.filter(student__program=program,
                                                                 status='solicitante').order_by(
                        '-student__request_date')[:4],
                    'last_aproved': PhdStudent.objects.filter(student__program=program, status='doctorando').order_by(
                        '-student__init_date')[:4],
                    'last_graduated': PhdStudent.objects.filter(student__program=program, status='graduado').order_by(
                        '-student__graduate_date')[:4],

                }
                try:
                    context['member'] = ProgramMember.objects.get(user=request.user, program=program)
                except ProgramMember.DoesNotExist:
                    try:
                        student = Student.objects.get(user=request.user, program=program)
                        context['student'] = student
                        try:
                            formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                        except StudentFormationPlan.DoesNotExist:
                            return HttpResponseRedirect(
                                reverse('programs:create_formation_plan', args=[program_slug, student.id]))

                    except Student.DoesNotExist:
                        logout(request)
                        raise Http404('No hay profesor o estudiante de este programa con ese usuario')

                return render(request, 'programs/phd_home.html', context)
            elif program.type == 'msc':
                context = {
                    'program': program,
                    'requesters': MscStudent.objects.filter(program=program, status='solicitante').__len__(),
                    'masters': MscStudent.objects.filter(program=program, status='maestrante').__len__(),
                    'graduated': MscStudent.objects.filter(program=program, status='graduado').__len__(),
                    'last_requesters': MscStudent.objects.filter(program=program, status='solicitante').order_by(
                        '-request_date')[:4],
                    'last_aproved': MscStudent.objects.filter(program=program, status='maestrante').order_by(
                        '-init_date')[
                                    :4],
                    'last_graduated': MscStudent.objects.filter(program=program, status='graduado').order_by(
                        '-graduate_date')[
                                      :4],

                }
                try:
                    context['member'] = ProgramMember.objects.get(user=request.user, program=program)
                except ProgramMember.DoesNotExist:
                    try:
                        context['student'] = MscStudent.objects.get(user=request.user, program=program)
                    except MscStudent.DoesNotExist:
                        logout(request)
                        raise Http404('No hay profesor o estudiante de este programa con ese usuario')

                return render(request, 'programs/msc_home.html', context)
            elif program.type == 'dip':
                context = {
                    'program': program,
                    'requesters': DipStudent.objects.filter(program=program, status='solicitante').__len__(),
                    'diplomantes': DipStudent.objects.filter(program=program, status='diplomante').__len__(),
                    'graduated': DipStudent.objects.filter(program=program, status='graduado').__len__(),
                    'last_requesters': DipStudent.objects.filter(program=program, status='solicitante').order_by(
                        '-request_date')[:4],
                    'last_aproved': DipStudent.objects.filter(program=program, status='diplomante').order_by(
                        '-init_date')[
                                    :4],
                    'last_graduated': DipStudent.objects.filter(program=program, status='graduado').order_by(
                        '-graduate_date')[
                                      :4],

                }
                try:
                    context['member'] = ProgramMember.objects.get(user=request.user, program=program)
                except ProgramMember.DoesNotExist:
                    try:
                        context['student'] = DipStudent.objects.get(user=request.user, program=program)
                    except DipStudent.DoesNotExist:
                        logout(request)
                        raise Http404('No hay profesor o estudiante de este programa con ese usuario')

                return render(request, 'programs/dip_home.html', context)

            else:
                return HttpResponseNotFound("Tipo de programa no soportado") 
        else:
            return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
    except Program.DoesNotExist:
        if program_slug == 'favicon.ico':
            return HttpResponseNotFound()
        else:
            return HttpResponse('El programa no existe: ', program_slug)


    # TODO Verificar lo del user is authenticate

@login_required
def new_phd_announcement(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    phd_student = PhdStudent.objects.get(student=Student.objects.get(pk=student_id))
    if program.type == 'phd':
        if request.method == 'POST':
            announcement_form = AnnouncementForm(request.POST, request.FILES, initial={'phd_student':phd_student, 'thesis': phd_student.phdstudentthesis})
            if announcement_form.is_valid():

                try:
                    for i in range(1,4):
                        court_member = PhdDefenseCourtMember(
                            thesis = phd_student.phdstudentthesis,
                            name = request.POST['cmember_{0}_name'.format(i)],
                            lastname = request.POST['cmember_{0}_lastname'.format(i)],
                            center = request.POST['cmember_{0}_center'.format(i)],
                            role = 'Miembro'
                        )
                        court_member.save()
                except Exception as e:
                    messages.error(request, 'No fue posible guardar uin miembro del tribunal')
                    return HttpResponseRedirect(
                        reverse('programs:new_phd_announcement', args=[program_slug, student_id]))
                try:
                    for i in range(1, 3):
                        court_member = PhdDefenseCourtMember(
                            thesis=phd_student.phdstudentthesis,
                            name=request.POST['smember_{0}_name'.format(i)],
                            lastname=request.POST['smember_{0}_lastname'.format(i)],
                            center=request.POST['smember_{0}_center'.format(i)],
                            role='Suplente'
                        )
                        court_member.save()
                except Exception as e:
                    messages.error(request, 'No fue posible guardar un suplente del tribunal')
                    return HttpResponseRedirect(
                        reverse('programs:new_phd_announcement', args=[program_slug, student_id]))

                announcement = announcement_form.save()
                create_announcement_pdf(request, program_slug, announcement.id)
                messages.success(request, ('Nueva convocatoria creada exitosamente'))
                return HttpResponseRedirect(reverse('programs:program_announcements', args=[program_slug]))

            else:
                print(announcement_form.errors)
                messages.error(request, announcement_form.errors)
                return HttpResponseRedirect(reverse('programs:new_phd_announcement', args=[program_slug, student_id]))

        else:
            context = {
                'student':phd_student.student,
                'announcement_form': AnnouncementForm(),
                'program':program,

            }
            return render(request, 'programs/new_phd_announcement.html', context)
    else:
        pass

@login_required
def edit_phd_announcement(request, program_slug, announcement_id):
    program = Program.objects.get(slug=program_slug)
    announcement = PhdAnnouncement.objects.get(pk=announcement_id)
    if program.type == 'phd':
        if request.method == 'POST':
            announcement_form = AnnouncementForm(request.POST, instance=announcement)
            if announcement_form.is_valid():
                try:
                    for court_member in PhdDefenseCourtMember.objects.filter(thesis = announcement.phd_student.phdstudentthesis, role = "Miembro" ):
                        court_member.name = request.POST['cmember_{0}_name'.format(court_member.id)]
                        court_member.lastname = request.POST['cmember_{0}_lastname'.format(court_member.id)]
                        court_member.center = request.POST['cmember_{0}_center'.format(court_member.id)]

                        court_member.save()
                except Exception as e:
                    messages.error(request, 'No fue posible guardar uin miembro del tribunal')
                    return HttpResponseRedirect(
                        reverse('programs:edit_phd_announcement', args=[program_slug, announcement_id]))
                try:
                    for court_member in PhdDefenseCourtMember.objects.filter(
                            thesis=announcement.phd_student.phdstudentthesis, role="Suplente"):
                        court_member.name = request.POST['smember_{0}_name'.format(court_member.id)]
                        court_member.lastname = request.POST['smember_{0}_lastname'.format(court_member.id)]
                        court_member.center = request.POST['smember_{0}_center'.format(court_member.id)]

                        court_member.save()
                except Exception as e:
                    messages.error(request, 'No fue posible guardar un suplente del tribunal')
                    return HttpResponseRedirect(
                        reverse('programs:edit_phd_announcement', args=[program_slug, announcement_id]))

                announcement_form.save()
                messages.success(request, ('Nueva convocatoria editada exitosamente'))
                return HttpResponseRedirect(reverse('programs:program_announcements', args=[program_slug]))

            else:
                print(announcement_form.errors)
                messages.error(request, announcement_form.errors)
                return HttpResponseRedirect(reverse('programs:edit_phd_announcement', args=[program_slug, announcement_id]))

        else:
            context = {
                'student':announcement.phd_student.student,
                'announcement_form': AnnouncementForm(instance=announcement),
                'announcement':announcement,
                'program':program,

            }
            return render(request, 'programs/edit_phd_announcement.html', context)
    else:
        pass


@login_required
def create_student(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            try:
                user = User.objects.get(email=request.POST['student_email'])
                try:
                    Student.objects.get(user=user, program=program)
                    return error_500(request, program, 'El estudiante ya existe en este programa')
                except Student.DoesNotExist:
                    student = Student(
                        user=user,
                        program=program,
                        gender=request.POST['gender'],
                        dni=request.POST['student_dni'],
                        birth_date=request.POST['student_birth_date'],


                    )
                    try:
                        if request.POST['is_master']=='True':
                            student.is_master=True
                            student.msc_title = request.POST['msc_title']
                    except:
                        pass
                    student.save()

                    formation_plan = StudentFormationPlan(
                        phdstudent=student,
                        elaboration_date=now(),
                        last_update_date=now(),
                        planned_end_year=int(request.POST['student_planned_end_year']),

                    )
                    formation_plan.save()

                    utils_send_email(request, 'wm', program.email, student, '', '', program, '*********')

                    try:
                        student.picture = request.FILES['picture']
                        student.save()

                    except:
                        pass

                    if program.type == 'phd':
                        new_student = PhdStudent(
                            student=student,
                            status='solicitante',
                            category=request.POST['student_category'],
                            center=request.POST['student_center'],
                        )
                        new_student.save()
                        for i in range(1,int(request.POST['total_tuthors'])+1):
                            create_new_tuthor(request,program,
                                              request.POST['tuthor_name_' + str(i)],
                                              request.POST['tuthor_lastname_' + str(i)],
                                              request.POST['tuthor_institution_' + str(i)], request.POST['tuthor_email_' + str(i)], new_student)
                        new_theme = PhdStudentTheme(
                            phd_student=new_student,
                            description=request.POST['theme'],
                            line=InvestigationLine.objects.get(pk=request.POST['investigation_line'])
                        )

                        try:
                            project=InvestigationProject.objects.get(pk=request.POST['investigation_project'])
                            new_theme.project = project,
                        except:
                            pass

                        new_theme.save()
                    else:
                        return HttpResponse('Tipo de programa aun por crear')

                    for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):

                        if 'student_requirement_' + str(requirement.id) in request.POST:
                            new_student_requirement = StudentFileDocument(
                                student=student,
                                program_file_document=requirement,
                                accomplished=True,
                            )
                            if requirement.get_old:
                                try:
                                    new_student_requirement.caducity_date = request.POST['doc_caducity_date_'+str(requirement.id)]
                                except:
                                    pass

                            new_student_requirement.save()

                        else:
                            new_student_requirement = StudentFileDocument(
                                student=student,
                                program_file_document=requirement,
                            )
                            new_student_requirement.save()

                    return HttpResponseRedirect(reverse('programs:students_list', args=[program_slug, 'all']))
            except User.DoesNotExist:
                passwd = program_slug + str(random.randint(1000000, 9999999))
                user = User.objects.create_user(
                    request.POST['student_email'],
                    request.POST['student_email'],
                    passwd,  # Cambiar despues por contrase;a generada

                )
                user.first_name = request.POST['student_name']
                user.last_name = request.POST['student_surename']
                user.save()

            student = Student(
                user=user,
                program=program,
                gender=request.POST['gender'],
                dni=request.POST['student_dni'],
                birth_date=request.POST['student_birth_date'],


            )
            student.save()

            formation_plan = StudentFormationPlan(
                phdstudent=student,
                elaboration_date=now(),
                last_update_date=now(),
                planned_end_year=int(request.POST['student_planned_end_year']),

            )
            formation_plan.save()

            utils_send_email(request, 'wm', program.email, student, '', '', program, passwd)

            try:
                student.picture=request.FILES['picture']
                student.save()

            except:
                pass

            if program.type=='phd':
                new_student = PhdStudent(
                    student=student,
                    status='solicitante',
                    category=request.POST['student_category'],
                    center=request.POST['student_center'],
                )
                new_student.save()
                for i in range(1,int(request.POST['total_tuthors'])+1):
                    print(create_new_tuthor(request,
                                            program,
                                            request.POST['tuthor_name_' + str(i)],
                                            request.POST['tuthor_lastname_' + str(i)],
                                            request.POST['tuthor_institution_' + str(i)],
                                            request.POST['tuthor_email_' + str(i)],
                                            new_student))

                new_theme = PhdStudentTheme(
                    phd_student=new_student,
                    description=request.POST['theme'],
                    line=InvestigationLine.objects.get(pk=request.POST['investigation_line']),

                )
                try:
                    project = InvestigationProject.objects.get(pk=request.POST['investigation_project'])
                    new_theme.project = project

                except:
                    pass

                new_theme.save()
            else:
                return HttpResponse('Tipo de programa aun por crear')

            for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):

                if 'student_requirement_' + str(requirement.id) in request.POST:
                    new_student_requirement = StudentFileDocument(
                        student=student,
                        program_file_document=requirement,
                        accomplished=True,
                    )
                    if requirement.get_old:
                        try:
                            new_student_requirement.caducity_date = request.POST[
                                'doc_caducity_date_' + str(requirement.id)]
                        except:
                            pass
                    new_student_requirement.save()

                else:
                    new_student_requirement = StudentFileDocument(
                        student=student,
                        program_file_document=requirement,
                    )
                    new_student_requirement.save()

            return HttpResponseRedirect(reverse('programs:students_list', args=[program_slug, 'all']))
        else:
            context = {
                'program': program,
                'init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'lines': InvestigationLine.objects.filter(program=program),
                'inner_areas':InnerAreas.objects.all(),
                'specialities': ProgramSpeciality.objects.filter(program=program)
            }
            if Program.objects.get(slug=program_slug).type == 'phd':
                return render(request, 'programs/create_phd_student.html', context)
            else:
                return HttpResponse('El programa no es un doctorado')
    else:
        return HttpResponse('Error, acceso solo a coordinadores y secretarios')

def download_announcement_pdf(request, program_slug, announcement_id):
    announcement = PhdAnnouncement.objects.get(pk=announcement_id)
    conv_path = '{0}/{1}/Convs/Convocatoria_{1}_{2}.pdf'.format(MEDIA_ROOT, program_slug, announcement.phd_student.id)
    if os.path.exists(conv_path):
        return FileResponse(open(create_announcement_pdf(request, program_slug, announcement_id), 'rb'))
    else:
        create_announcement_pdf(request, program_slug, announcement_id)
        return HttpResponseRedirect(reverse('programs:download_announcement_pdf', args=[program_slug, announcement_id]))

def create_announcement_pdf(request, program_slug, announcement_id):
    program = Program.objects.get(slug=program_slug)
    announcement = PhdAnnouncement.objects.get(pk=announcement_id)
    student_name = '{0} {1}'.format(announcement.phd_student.student.user.first_name, announcement.phd_student.student.user.last_name)
    thesis_title = announcement.phd_student.phdstudenttheme.description

    try:

        packet = io.BytesIO()
        # create a new PDF with Reportlab
        can = canvas.Canvas(packet, pagesize=letter)


        # can.setFontSize(int(request.POST['font_size']))

        can.setFillColor('000000')
        can.setFont('Times-Roman', 12)
        can.drawString(202,659, student_name)
        title_len = 90
        if len(thesis_title) > title_len:
            wrap_text = textwrap.wrap(thesis_title, width=title_len)
            for item in wrap_text:
                if wrap_text.index(item) <=1:
                    can.drawString(166,631- wrap_text.index(item) * 12, item)
        else:
            can.drawString(166, 631, thesis_title)

        can.drawString(202,604, program.full_name)
        can.drawString(165,576, announcement.sponsor)

        tutors = ''
        for tutor in announcement.phd_student.tuthor_set.all():
            tutors = tutors+'Dr.C. {0} {1},'.format(tutor.professor.user.first_name, tutor.professor.user.last_name)

        can.drawString(132, 548, tutors)
        can.drawString(123, 521, str(announcement.date.date()))
        can.drawString(239, 521, str(announcement.date.time()))
        can.drawString(390, 521, announcement.type)
        can.drawString(123, 494, announcement.place)
        if announcement.type == 'On-line':
            can.drawString(83, 465, 'URL de la sala: ')

            p = Paragraph(' <a href="{0}" color=blue>dé click aquí para acceder a la sala virtual de la defensa</a>'.format(announcement.url_vc), style)


            p.wrap(800, 800)
            p.drawOn(can, 160, 463)


        can.drawString(150, 410, announcement.approval_resolution)

        i=0
        for member in PhdDefenseCourtMember.objects.filter(thesis = announcement.phd_student.phdstudentthesis, role = "Miembro"):
            can.drawString(83, 370-i*18, "Dr.C. {0} {1}".format(member.name, member.lastname))
            i+=1

        i=0
        for member in PhdDefenseCourtMember.objects.filter(thesis = announcement.phd_student.phdstudentthesis, role = "Suplente"):
            can.drawString(83, 296-i*18, "Dr.C. {0} {1}".format(member.name, member.lastname))
            i+=1

        url1 = request.scheme +'://'+request.META['HTTP_HOST']+announcement.phd_student.phdstudentthesis.file.url
        url2 = request.scheme +'://'+request.META['HTTP_HOST']+reverse('programs:new_phd_thesis_comment', args=[program_slug,announcement.phd_student.phdstudentthesis.id])
        p1 = Paragraph(' <a href="{0}" color=blue>aquí (dé click sobre el texto azul)</a>'.format(url1), style)
        p2 = Paragraph(' <a href="{0}" color=blue>Para comentar la tesis puede dar click aquí</a>'.format(url2), style)
        p1.wrap(800,800)
        p1.drawOn(can,250,243)
        p2.wrap(800, 800)
        p2.drawOn(can, 350, 216)


        # if len(authors_text) > authors_string_large:
        #     wrap_text = textwrap.wrap(authors_text, width=authors_string_large)
        #     for item in wrap_text:
        #         can.drawString(x_authors,
        #                        y_authors - wrap_text.index(item) * int(request.POST['authors_font_size']),
        #                        item)
        # else:
        #     can.drawString(x_authors, y_authors, authors_text)

        # title_text = upper('Titulo de la ponencia, Titulo de la ponencia, Titulo de la ponencia, Titulo de la ponencia, Titulo de la ponencia')
        # can.setFillColor(request.POST['title_font_color'])
        # can.setFont(request.POST['title_font_name'], int(request.POST['title_font_size']))
        # if len(title_text) > title_string_large:
        #     wrap_text = textwrap.wrap(title_text, width=title_string_large)
        #     for item in wrap_text:
        #         can.drawString(x_name,
        #                        y_name - wrap_text.index(item) * int(request.POST['title_font_size']), item)
        # else:
        #     can.drawString(x_name, y_name, title_text)

        can.showPage()
        can.save()

        # move to the beginning of the StringIO buffer
        packet.seek(0)
        new_pdf = PdfFileReader(packet)
        # read your existing PDF

        existing_pdf = PdfFileReader(open(STATIC_ROOT+'/programs/docs/CONVOCATORIA.pdf', "rb"))
        output = PdfFileWriter()
        # add the "watermark" (which is the new pdf) on the existing page
        page = existing_pdf.getPage(0)
        page.mergePage(new_pdf.getPage(0))
        output.addPage(page)
        # finally, write "output" to a real file
        conv_path = '{0}/{1}/Convs/Convocatoria_{1}_{2}.pdf'.format(MEDIA_ROOT, program_slug, announcement.phd_student.id)
        if not os.path.exists('{0}/{1}/Convs'.format(MEDIA_ROOT, program_slug)):
            os.makedirs('{0}/{1}/Convs'.format(MEDIA_ROOT, program_slug))

        outputStream = open(conv_path, "wb")
        output.write(outputStream)
        outputStream.close()
        url3 = '{0}//{1}/announcements/{2}/pdf'.format(request.scheme,request.META['HTTP_HOST'], announcement_id )
        print(url3)
        send_mail('Nueva convocatoria de defensa pública en '+program.full_name, 'Buenos días!! Nuestro programa anuncia una nueva defensa. '
                                                                                 'El pdf con los datos de la defensa está '
                                                                                 'disponible en '+url3+
                                                                                 ' .Saludos',
                  program.email,[DIR_COM_EMAIL],fail_silently=True)

        return conv_path
    except Exception as e:
        return str(e)


    return HttpResponse("Test")

@login_required
def create_msc_student(request, program_slug, edition_id):
    program=Program.objects.get(slug=program_slug)
    edition = ProgramEdition.objects.get(pk = edition_id)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            try:
                user = User.objects.get(email=request.POST['student_email'])
                try:
                    student = MscStudent.objects.get(user=user, program=program)
                    return error_500(request, program, 'Estudiante matriculado previamente en el programa')
                except MscStudent.DoesNotExist:
                    #Se crea el MscStudent

                    student = MscStudent(
                        user=user,
                        edition=edition,
                        program=program,
                        gender=request.POST['gender'],
                        dni=request.POST['student_dni'],
                        birth_date=request.POST['student_birth_date'],
                        country=request.POST['student_country']
                    )
                    student.save()

                    for i in range(1, int(request.POST['total_tuthors']) + 1):
                        create_new_tuthor(request, program,
                                          request.POST['tuthor_name_' + str(i)],
                                          request.POST['tuthor_lastname_' + str(i)],
                                          request.POST['tuthor_institution_' + str(i)],
                                          request.POST['tuthor_email_' + str(i)], student)

                    utils_send_email(request, 'wm', program.email, student, '', '', program, '*********')

                    try:
                        student.picture = request.FILES['picture']
                        student.save()

                    except:
                        pass

                    if program.type == 'msc':

                        new_theme = MscStudentTheme(
                            student=student,
                            description=request.POST['theme'],
                        )
                        try:
                            new_theme.project = InvestigationProject.objects.get(
                                pk=request.POST['investigation_project'])
                            new_theme.line = InvestigationProject.objects.get(
                                pk=request.POST['investigation_project']).line,

                        except:
                            pass

                        new_theme.save()
                    else:
                        return HttpResponse('Tipo de programa aun por crear')

                    for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):

                        if 'student_requirement_' + str(requirement.id) in request.POST:
                            new_student_requirement = StudentFileDocument(
                                msc_student=student,
                                program_file_document=requirement,
                                accomplished=True,
                            )
                            if requirement.get_old:
                                new_student_requirement.caducity_date = request.POST['caducity_date']

                            new_student_requirement.save()

                        else:
                            new_student_requirement = StudentFileDocument(
                                msc_student=student,
                                program_file_document=requirement,
                                accomplished=False,
                            )
                            if requirement.get_old:
                                new_student_requirement.caducity_date = None
                            new_student_requirement.save()

                    return HttpResponseRedirect(reverse('programs:create_msc_student', args=[program_slug, edition_id]))


            except User.DoesNotExist:
                passwd = program_slug + str(random.randint(1000000, 9999999))
                user = User.objects.create_user(
                    request.POST['student_email'],
                    request.POST['student_email'],
                    passwd,  # Cambiar despues por contrase;a generada

                )
                user.first_name = request.POST['student_name']
                user.last_name = request.POST['student_surename']
                user.save()

            student = MscStudent(
                user=user,
                edition=edition,
                program=program,
                gender=request.POST['gender'],
                dni=request.POST['student_dni'],
                birth_date=request.POST['student_birth_date'],
                country=request.POST['student_country']
            )
            student.save()

            for i in range(1, int(request.POST['total_tuthors']) + 1):
                create_new_tuthor(request, program,
                                  request.POST['tuthor_name_' + str(i)],
                                  request.POST['tuthor_lastname_' + str(i)],
                                  request.POST['tuthor_institution_' + str(i)],
                                  request.POST['tuthor_email_' + str(i)], student)

            utils_send_email(request, 'wm', program.email, student, '', '', program, passwd)

            try:
                student.picture=request.FILES['picture']
                student.save()

            except:
                pass

            if program.type=='msc':

                new_theme=MscStudentTheme(
                    student=student,
                    description=request.POST['theme'],
                )
                try:
                    new_theme.project=InvestigationProject.objects.get(pk=request.POST['investigation_project'])
                    new_theme.line=InvestigationProject.objects.get(pk=request.POST['investigation_project']).line,

                except:
                    pass

                new_theme.save()
            else:
                return HttpResponse('Tipo de programa aun por crear')

            for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):

                if 'student_requirement_' + str(requirement.id) in request.POST:
                    new_student_requirement = StudentFileDocument(
                        msc_student=student,
                        program_file_document=requirement,
                        accomplished=True,
                    )
                    if requirement.get_old:
                        new_student_requirement.caducity_date = request.POST['caducity_date']

                    new_student_requirement.save()

                else:
                    new_student_requirement = StudentFileDocument(
                        msc_student=student,
                        program_file_document=requirement,
                        accomplished=False,
                    )
                    if requirement.get_old:
                        new_student_requirement.caducity_date = None
                    new_student_requirement.save()

            return HttpResponseRedirect(reverse('programs:create_msc_student', args=[program_slug, edition_id]))
        else:
            context = {
                'program': program,
                'edition': edition,
                'init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'projects': InvestigationProject.objects.filter(program=program),
                'member': ProgramMember.objects.get(user=request.user, program = program)
            }
            if Program.objects.get(slug=program_slug).type == 'msc':
                return render(request, 'programs/create_msc_student.html', context)
            else:
                return HttpResponse('El programa no es una maestria')
    else:
        return error_500(request,program, 'Usted no tiene privilegios para agregar estudiantes en este programa')

@login_required
def create_dip_student(request, program_slug, edition_id):
    program=Program.objects.get(slug=program_slug)
    edition = ProgramEdition.objects.get(pk = edition_id)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            try:
                user = User.objects.get(email=request.POST['student_email'])
                try:
                    student = DipStudent.objects.get(user=user, program=program)
                    return error_500(request, program, 'Estudiante matriculado previamente en el programa')
                except DipStudent.DoesNotExist:
                    #Se crea el MscStudent

                    student = DipStudent(
                        user=user,
                        edition=edition,
                        program=program,
                        gender=request.POST['gender'],
                        dni=request.POST['student_dni'],
                        birth_date=request.POST['student_birth_date'],
                        country=request.POST['student_country']
                    )
                    student.save()

                    utils_send_email(request, 'wm', program.email, student, '', '', program, '*********')

                    try:
                        student.picture = request.FILES['picture']
                        student.save()

                    except:
                        pass

                    return HttpResponseRedirect(reverse('programs:create_msc_student', args=[program_slug, edition_id]))


            except User.DoesNotExist:
                passwd = program_slug + str(random.randint(1000000, 9999999))
                user = User.objects.create_user(
                    request.POST['student_email'],
                    request.POST['student_email'],
                    passwd,  # Cambiar despues por contrase;a generada

                )
                user.first_name = request.POST['student_name']
                user.last_name = request.POST['student_surename']
                user.save()

            student = DipStudent(
                user=user,
                edition=edition,
                program=program,
                gender=request.POST['gender'],
                dni=request.POST['student_dni'],
                birth_date=request.POST['student_birth_date'],
                country=request.POST['student_country']
            )
            student.save()

            utils_send_email(request, 'wm', program.email, student, '', '', program, passwd)

            try:
                student.picture=request.FILES['picture']
                student.save()

            except:
                pass

            return HttpResponseRedirect(reverse('programs:create_dip_student', args=[program_slug, edition_id]))
        else:
            context = {
                'program': program,
                'edition': edition,
            }
            if program.type == 'dip':
                return render(request, 'programs/create_dip_student.html', context)
            else:
                return HttpResponse('El programa no es un diplomado')
    else:
        return HttpResponse('Error, acceso solo a coordinadores y secretarios')


@login_required
def ajx_new_phd_thesis(request, program_slug, student_id):
    program=Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)

    if request.method == 'POST':
        thesis_form = PhdStudentThesisForm(request.POST, request.FILES)
        if thesis_form.is_valid():
            thesis=thesis_form.save()

            return HttpResponse(
                json.dumps([{'uploaded': '1', 'url':thesis.file.url }]),
                content_type="application/json")
        else:

            return HttpResponse(
                json.dumps([{'uploaded': '0', 'errors':thesis_form.errors }]),
                content_type="application/json")


def ajx_new_phd_thesis_comment(request, program_slug, thesis_id):
    program=Program.objects.get(slug=program_slug)
    thesis = PhdStudentThesis.objects.get(pk=thesis_id)

    if request.method == 'POST':
        thesis_comment_form = PhdThesisCommentForm(request.POST)
        if thesis_comment_form.is_valid():
            thesis=thesis_comment_form.save()

            return HttpResponse(
                json.dumps([{'uploaded': '1'}]),
                content_type="application/json")
        else:
            print(thesis_comment_form.errors)
            return HttpResponse(
                json.dumps([{'uploaded': '0', 'errors':thesis_comment_form.errors }]),
                content_type="application/json")

def new_phd_thesis_comment(request, program_slug, thesis_id):
    program=Program.objects.get(slug=program_slug)
    thesis = PhdStudentThesis.objects.get(pk=thesis_id)



    if request.method == 'POST':
        thesis_comment_form = PhdThesisCommentForm(request.POST)
        if thesis_comment_form.is_valid():
            thesis_comment_form.save()

            messages.success(request, 'Comentario guardado exitosamente. Gracias por dejarnos saber su opinión!!!')

            return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
        else:
            print(thesis_comment_form.errors)
    else:
        thesis_comment_form = PhdThesisCommentForm()

    context = {
        'program': program,
        'comment_form':thesis_comment_form,
        'thesis':thesis,
    }
    return render(request, 'programs/phd_comment.html', context)




@login_required
def ajx_rm_phd_thesis(request, program_slug, student_id):
    program=Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)

    if request.method == 'POST':
        thesis = student.phdstudent.phdstudentthesis
        try:
            thesis.delete()

            return HttpResponse(
                json.dumps([{'deleted': '1'}]),
                content_type="application/json")
        except:
            return HttpResponse(
                json.dumps([{'deleted': '0'}]),
                content_type="application/json")

@login_required
def ajx_rm_phd_thesis_comment(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    comment = PhdThesisComment.objects.get(pk=request.POST['comment_id'])

    if request.method == 'POST':

        try:
            comment.delete()

            return HttpResponse(
                json.dumps([{'deleted': '1'}]),
                content_type="application/json")
        except:
            return HttpResponse(
                json.dumps([{'deleted': '0'}]),
                content_type="application/json")


@login_required
def students_list(request, program_slug, scope):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_member(request.user, program) or user_is_program_student(request.user,program):
        if program.type == 'phd':
            if scope == 'all':
                context = {
                    'program': program,
                    'students': Student.objects.filter(program=program),
                    'scope': 'all',

                }
            elif scope == 'requesters':
                context = {
                    'program': program,
                    'students': Student.objects.filter(program=program, phdstudent__status='solicitante'),
                    'scope': 'Solicitantes',
                }
            elif scope == 'aproved':
                context = {
                    'program': program,
                    'students': Student.objects.filter(program=program, phdstudent__status='doctorando'),
                    'scope': 'Doctorandos',
                }
            elif scope == 'graduated':
                context = {
                    'program': program,
                    'students': Student.objects.filter(program=program, phdstudent__status='graduado'),
                    'scope': 'Graduados',
                }

            context['en_scope'] = scope

            if user_is_program_member(request.user, program):
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            elif user_is_program_student(request.user, program):
                context['student']=Student.objects.get(user=request.user, program=program)

            return render(request, 'programs/students_list.html', context)
        elif program.type == 'msc':
            if scope == 'all':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program),
                    'scope': 'all',
                }
            elif scope == 'requesters':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, status='solicitante'),
                    'scope': 'Solicitantes',
                }
            elif scope == 'aproved':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, status='maestrante'),
                    'scope': 'Maestrantes',
                }
            elif scope == 'graduated':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, status='graduado'),
                    'scope': 'Graduados',
                }

            return render(request, 'programs/students_list.html', context)
        elif program.type == 'dip':
            return HttpResponse('Aun no se implementa este tipo de programas')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta página')

@login_required
def phd_thesis_comments(request, program_slug, thesis_id):
    program = Program.objects.get(slug=program_slug)
    context = {
        'program':program,
        'comments': PhdThesisComment.objects.filter(thesis=PhdStudentThesis.objects.get(pk=thesis_id)),
        'member': ProgramMember.objects.get(user=request.user, program=program),
        'thesis': PhdStudentThesis.objects.get(pk=thesis_id),
    }
    return render(request, 'programs/comments_list.html', context)



@login_required
def msc_edition_students_list(request, program_slug, edition_id, scope):
    program=Program.objects.get(slug=program_slug)
    edition=ProgramEdition.objects.get(pk=edition_id)
    if user_is_program_member(request.user, program) or user_is_program_student(request.user,program):
        if program.type == 'msc':
            if scope == 'all':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, edition=edition),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'all',
                }
            elif scope == 'requesters':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, edition=edition, status='solicitante'),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'Solicitantes',
                }
            elif scope == 'aproved':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, edition=edition, status='maestrante'),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'Maestrantes',
                }
            elif scope == 'graduated':
                context = {
                    'program': program,
                    'students': MscStudent.objects.filter(program=program, edition=edition, status='graduado'),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'Graduados',
                }

            try:
                context['member']= ProgramMember.objects.get(user= request.user, program = program)
            except ProgramMember.DoesNotExist:
                pass

            return render(request, 'programs/msc_students_list.html', context)
        else:
            return error_500(request, program, 'El programa debe ser una maestria')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta página')

@login_required
def dip_edition_students_list(request, program_slug, edition_id, scope):
    program=Program.objects.get(slug=program_slug)
    edition=ProgramEdition.objects.get(pk=edition_id)
    if user_is_program_member(request.user, program) or user_is_program_student(request.user,program):
        if program.type == 'dip':
            if scope == 'all':
                context = {
                    'program': program,
                    'students': DipStudent.objects.filter(program=program, edition=edition),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'all',
                }
            elif scope == 'requesters':
                context = {
                    'program': program,
                    'students': DipStudent.objects.filter(program=program, edition=edition, status='solicitante'),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'Solicitantes',
                }
            elif scope == 'aproved':
                context = {
                    'program': program,
                    'students': DipStudent.objects.filter(program=program, edition=edition, status='diplomante'),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'Maestrantes',
                }
            elif scope == 'graduated':
                context = {
                    'program': program,
                    'students': DipStudent.objects.filter(program=program, edition=edition, status='graduado'),
                    'edition': ProgramEdition.objects.get(pk=edition_id),
                    'scope': 'Graduados',
                }

            return render(request, 'programs/dip_students_list.html', context)
        else:
            return error_500(request, program, 'El programa debe ser un diplomadoa')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta página')

@login_required
def msc_all_students_list(request, program_slug, scope):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_member(request.user, program) or user_is_program_student(request.user,program):
        if program.type == 'msc':
            if scope == 'all':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': MscStudent.objects.filter(program=program),
                    'scope': 'all',
                }
            elif scope == 'requesters':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': MscStudent.objects.filter(program=program, status='solicitante'),
                    'scope': 'Solicitantes',
                }
            elif scope == 'aproved':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': MscStudent.objects.filter(program=program, status='maestrante'),
                    'scope': 'Maestrantes',
                }
            elif scope == 'graduated':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': MscStudent.objects.filter(program=program, status='graduado'),
                    'scope': 'Graduados',
                }
            try:
                context['student']=MscStudent.objects.get(user=request.user)
            except MscStudent.DoesNotExist:
                try:
                    context['member']=ProgramMember.objects.get(user=request.user, program=program)
                except ProgramMember.DoesNotExist:
                    pass
            context['en_scope']=scope
            return render(request, 'programs/msc_students_list.html', context)
        else:
            return error_500(request, program, 'El programa debe ser una maestria')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta página')

@login_required
def dip_all_students_list(request, program_slug, scope):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_member(request.user, program) or user_is_program_student(request.user,program):
        if program.type == 'dip':
            if scope == 'all':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': DipStudent.objects.filter(program=program),
                    'scope': 'all',
                }
            elif scope == 'requesters':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': DipStudent.objects.filter(program=program, status='solicitante'),
                    'scope': 'Solicitantes',
                }
            elif scope == 'aproved':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': DipStudent.objects.filter(program=program, status='diplomante'),
                    'scope': 'Diplomantes',
                }
            elif scope == 'graduated':
                context = {
                    'program': program,
                    'editions': ProgramEdition.objects.filter(program=program),
                    'students': DipStudent.objects.filter(program=program, status='graduado'),
                    'scope': 'Graduados',
                }

            return render(request, 'programs/dip_students_list.html', context)
        else:
            return error_500(request, program, 'El programa debe ser un diplomado.')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta página')

@login_required
def members_list(request, program_slug, scope):
    program=Program.objects.get(slug=program_slug)
    if scope == 'all':
        context = {
            'program': program,
            'members': ProgramMember.objects.filter(program=program),
            'scope': 'all',
        }
    elif scope == 'comite':
        context = {
            'program': program,
            'members': ProgramMember.objects.filter(Q(role='Miembro')|Q(role='Coordinador')|Q(role='Secretario'), program=program),
            'scope': 'Comité académico',
        }
    elif scope == 'professors':
        context = {
            'program': program,
            'members': ProgramMember.objects.filter(program=program, role='Profesor'),
            'scope': 'Profesores',
        }
    elif scope == 'tutors':
        context = {
            'program': program,
            'members': ProgramMember.objects.filter(program=program, role='Tutor'),
            'scope': 'Tutores',
        }
    if user_is_program_member(request.user, program):
        context['member'] = ProgramMember.objects.get(user=request.user, program=program)
    elif user_is_program_student(request.user, program):
        if program.type == 'phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type == 'msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type == 'dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)


    return render(request, 'programs/members_list.html', context)

@login_required
def edit_student(request, program_slug, student_id):
    program= Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            user=student.user
            user.first_name=request.POST['student_name']
            user.last_name=request.POST['student_surename']
            user.email=request.POST['student_email']
            user.save()

            Student.objects.filter(pk=student_id).update(
                phone=request.POST['student_phone'],
                country=request.POST['student_country'],
                gender=request.POST['gender'],
                dni=request.POST['student_dni'],
                birth_date=request.POST['student_birth_date'],

            )
            if 'request_date' in request.POST and not request.POST['request_date'] == '':
                Student.objects.filter(pk=student_id).update(
                    request_date=request.POST['request_date']
                )
            if 'init_date' in request.POST and not request.POST['init_date'] == '':
                Student.objects.filter(pk=student_id).update(
                    init_date=request.POST['init_date']
                )
            if 'graduate_date' in request.POST and not request.POST['graduate_date'] == '':
                Student.objects.filter(pk=student_id).update(
                    graduate_date=request.POST['graduate_date']
                )
            try:
                if request.POST['is_master']== 'True':
                    Student.objects.filter(pk=student_id).update(
                        is_master=True,
                        msc_title=request.POST['msc_title'],
                    )
                else:
                    Student.objects.filter(pk=student_id).update(
                        is_master=False,
                        msc_title='',
                    )
            except:
                Student.objects.filter(pk=student_id).update(
                    is_master=False,
                    msc_title='',
                )
            try:
                if request.POST['have_prorrogue']== 'True':
                    Student.objects.filter(pk=student_id).update(
                        have_prorrogue=True,
                        prorrogue_end_date=request.POST['prorrogue_end_date'],
                    )
                else:
                    Student.objects.filter(pk=student_id).update(
                        have_prorrogue=False,
                        prorrogue_end_date=None,
                    )
            except:
                Student.objects.filter(pk=student_id).update(
                    have_prorrogue=False,
                    prorrogue_end_date=None,
                )




            PhdStudent.objects.filter(student=student).update(
                status=request.POST['student_status'],
                category=request.POST['student_category'],
                center=request.POST['student_center'],
            )

            student_theme, st_created = PhdStudentTheme.objects.get_or_create(
                phd_student=PhdStudent.objects.get(student=Student.objects.get(pk=student_id)),
            )
            student_theme.line = InvestigationLine.objects.get(pk=request.POST['investigation_line'])
            student_theme.description = request.POST['theme']

            student_theme.save()


            try:
                student_theme.project = InvestigationProject.objects.get(pk=request.POST['investigation_project'])
                student_theme.save()
            except:
                pass

            try:
                formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                formation_plan.planned_end_year = int(request.POST['student_planned_end_year'])
                formation_plan.last_update_date = now()
                formation_plan.save()
            except StudentFormationPlan.DoesNotExist:
                formation_plan = StudentFormationPlan(
                    phdstudent=student,
                    elaboration_date=now(),
                    last_update_date=now(),
                    planned_end_year=int(request.POST['student_planned_end_year']),

                )
                formation_plan.save()




            try:
                if request.FILES['student_picture']:
                    student.picture=request.FILES['student_picture']
                    student.save()
            except:
                pass

            for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):
                if 'student_new_requirement_' + str(requirement.id) in request.POST:
                    s_i_r=StudentFileDocument.objects.get(student=student, program_file_document=requirement)
                    s_i_r.accomplished=True
                    if requirement.get_old:
                        try:
                            s_i_r.caducity_date = request.POST['doc_caducity_date_' + str(requirement.id)]
                        except:
                            s_i_r.caducity_date = None

                    s_i_r.save()


                else:
                    s_i_r = StudentFileDocument.objects.get(student=student,
                                                            program_file_document=requirement)
                    s_i_r.caducity_date = None
                    s_i_r.accomplished = False
                    s_i_r.save()

            print(student.phdstudent.status)

            if student.phdstudent.status == 'doctorando' or student.phdstudent.status == 'graduado':
                for requirement in ProgramFileDoc.objects.filter(program=program, is_finish_requirenment=True):
                    if 'student_new_f_requirement_' + str(requirement.id) in request.POST:
                        s_f_r=StudentFileDocument.objects.get(student=student, program_file_document=requirement)
                        s_f_r.accomplished=True
                        if requirement.get_old:
                            try:
                                s_f_r.caducity_date = request.POST['doc_caducity_date_' + str(requirement.id)]
                            except:
                                s_f_r.caducity_date = None
                        s_f_r.save()
                    else:
                        s_f_r = StudentFileDocument.objects.get(student=student,
                                                                program_file_document=requirement)
                        s_f_r.accomplished = False
                        s_f_r.caducity_date = None
                        s_f_r.save()

            return HttpResponseRedirect(reverse('programs:students_list', args=[program_slug,'all']))
        else:
            context = {
                'member': ProgramMember.objects.get(user=request.user, program=program),
                'program': program,
                'doc_student': Student.objects.get(pk=student_id),
                'phd_student':PhdStudent.objects.get(student=Student.objects.get(pk=student_id), student__program=program),
                'init_requirements': ProgramInitRequirements.objects.filter(program=program),
                'new_init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'new_finish_requirements': ProgramFileDoc.objects.filter(program=program, is_finish_requirenment=True),
                'finish_requirements': ProgramFinishRequirements.objects.filter(program=program),
                'lines':InvestigationLine.objects.filter(program=program),
                'projects': InvestigationProject.objects.filter(program=program, line=PhdStudent.objects.get(student=Student.objects.get(pk=student_id)).phdstudenttheme.line),
                'inner_areas': InnerAreas.objects.all(),
                'specialities': ProgramSpeciality.objects.filter(program=program),
            }
            return render(request, 'programs/edit_phd_student.html', context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar estudiantes de este programa.')

@login_required
def view_student_profile(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    user_is_student = False
    if program.type == 'phd':
        try:
            if Student.objects.get(user=request.user, program=program) == Student.objects.get(pk=student_id, program=program):
                user_is_student = True
        except Student.DoesNotExist:
            pass
    elif program.type == 'msc':
        try:
            if MscStudent.objects.get(user=request.user, program=program) == MscStudent.objects.get(pk=student_id, program=program):
                user_is_student = True
        except MscStudent.DoesNotExist:
            pass
    elif program.type == 'dip':
        try:
            if DipStudent.objects.get(user=request.user, program=program) == DipStudent.objects.get(pk=student_id, program=program):
                user_is_student = True
        except DipStudent.DoesNotExist:
            pass
    if user_is_program_member(request.user, program) or user_is_student:
        if program.type == 'phd':
            context = {
                'program': program,
                'student': Student.objects.get(pk=student_id),
                'phd_student': PhdStudent.objects.get(student=Student.objects.get(pk=student_id),
                                                      student__program=program),
                'init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'finish_requirements': ProgramFileDoc.objects.filter(program=program, is_finish_requirenment=True),
                'projects': InvestigationProject.objects.filter(program=program),
                'inner_areas': InnerAreas.objects.all(),
                'messages': Message.objects.filter(Q(phd_student_receiver=Student.objects.get(pk=student_id))|Q(sender=request.user))
            }
            try:
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            except ProgramMember.DoesNotExist:
                pass
            return render(request, 'programs/phd_student_profile.html', context)
        elif program.type == 'msc':
            context = {
                'program': program,
                'student': MscStudent.objects.get(pk=student_id),
                'init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'finish_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'projects': InvestigationProject.objects.filter(program=program),
                'edition':MscStudent.objects.get(pk=student_id).edition,
                'inner_areas': InnerAreas.objects.all(),
                'messages': Message.objects.filter(
                    Q(msc_student_receiver=MscStudent.objects.get(pk=student_id)) | Q(sender=request.user))

            }
            try:
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            except ProgramMember.DoesNotExist:
                pass
            return render(request, 'programs/msc_student_profile.html', context)
        elif program.type == 'dip':
            context = {
                'program': program,
                'student': DipStudent.objects.get(pk=student_id),
                'init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'finish_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'projects': InvestigationProject.objects.filter(program=program),
                'edition': DipStudent.objects.get(pk=student_id).edition,
                'inner_areas': InnerAreas.objects.all(),
                'messages': Message.objects.filter(
                    Q(dip_student_receiver=DipStudent.objects.get(pk=student_id)) | Q(sender=request.user))

            }
            try:
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            except ProgramMember.DoesNotExist:
                pass
            return render(request, 'programs/dip_student_profile.html', context)
        else:
            return error_500(request, program, 'Tipo de programa no implementado.')

    else:
        return error_500(request, program, 'Usted no tiene acceso al perfil de este estudiante')


@login_required
def view_program_member_profile(request, program_slug, member_id):
    program = Program.objects.get(slug=program_slug)
    member = ProgramMember.objects.get(pk=member_id)
    if user_is_program_cs(request.user, program) or request.user == member.user:
        context={
            'program':program,
            'professor': member,
            'member': ProgramMember.objects.get(user=request.user, program =program)
        }
        if request.user == member.user:
            context['messages']=Message.objects.filter(program_receiver=member)
            context['messages_sended']=MessageSended.objects.filter(sender=member.user)

        return render(request, 'programs/program_member_profile.html', context)
    else:
        return error_500(request, program, "Usted no puede ver el perfil de este miembro del claustro")

@login_required
def edit_msc_student(request, program_slug, edition_id, student_id):
    program= Program.objects.get(slug=program_slug)
    student = MscStudent.objects.get(pk=student_id)
    
    if user_is_program_cs(request.user, program) or request.user == student.user:
        if request.method == 'POST':
            user=MscStudent.objects.get(pk=student_id).user
            user.first_name=request.POST['student_name']
            user.last_name=request.POST['student_surename']
            user.email=request.POST['student_email']
            user.save()

            MscStudent.objects.filter(pk=student_id).update(
                phone=request.POST['student_phone'],
                country=request.POST['student_country'],
                gender=request.POST['gender'],
                dni=request.POST['student_dni'],
                birth_date=request.POST['student_birth_date']

            )
            if 'request_date' in request.POST and not request.POST['request_date'] == '':
                MscStudent.objects.filter(pk=student_id).update(
                    request_date=request.POST['request_date']
                )
            if 'init_date' in request.POST and not request.POST['init_date'] == '':
                MscStudent.objects.filter(pk=student_id).update(
                    init_date=request.POST['init_date']
                )
            if 'graduate_date' in request.POST and not request.POST['graduate_date'] == '':
                MscStudent.objects.filter(pk=student_id).update(
                    graduate_date=request.POST['graduate_date']
                )

            MscStudent.objects.filter(pk=student_id).update(
                status=request.POST['student_status']
            )

            student_theme, created = MscStudentTheme.objects.get_or_create(
                student=MscStudent.objects.get(pk=student_id),

            )
            student_theme.description = request.POST['theme']

            student_theme.save()
            try:
                student_theme.project = InvestigationProject.objects.get(pk=request.POST['investigation_project'])
                student_theme.line = InvestigationProject.objects.get(pk=request.POST['investigation_project']).line
                student_theme.save()
            except:
                pass

            try:
                if request.FILES['student_picture']:
                    student=MscStudent.objects.get(pk=student_id)
                    student.picture=request.FILES['student_picture']
                    student.save()
            except:
                pass

            for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):
                if 'student_new_requirement_' + str(requirement.id) in request.POST:
                    s_i_r=StudentFileDocument.objects.get(msc_student=MscStudent.objects.get(pk=student_id), program_file_document=requirement)
                    s_i_r.accomplished=True
                    if requirement.get_old:
                        try:
                            s_i_r.caducity_date = request.POST['doc_caducity_date_' + str(requirement.id)]
                        except:
                            s_i_r.caducity_date = None
                    s_i_r.save()
                else:
                    s_i_r = StudentFileDocument.objects.get(msc_student=MscStudent.objects.get(pk=student_id),
                                                            program_file_document=requirement)

                    s_i_r.accomplished = False
                    s_i_r.caducity_date = None
                    s_i_r.save()
            if student.status == 'maestrante' or student.status == 'graduado':
                for requirement in ProgramFileDoc.objects.filter(program=program, is_finish_requirenment=True):
                    if 'student_new_f_requirement_' + str(requirement.id) in request.POST:
                        s_f_r = StudentFileDocument.objects.get(msc_student=MscStudent.objects.get(pk=student_id),
                                                                program_file_document=requirement)
                        if requirement.get_old:
                            try:
                                s_f_r.caducity_date = request.POST['doc_caducity_date_' + str(requirement.id)]
                            except:
                                s_f_r.caducity_date = None

                        s_f_r.accomplished=True
                        s_f_r.save()
                    else:
                        s_f_r = StudentFileDocument.objects.get(msc_student=MscStudent.objects.get(pk=student_id),
                                                                program_file_document=requirement)
                        s_f_r.caducity_date = None
                        s_f_r.accomplished = False
                        s_f_r.save()

            return HttpResponseRedirect(reverse('programs:msc_all_students_list', args=[program_slug,'all']))
        else:
            context = {
                'program': program,
                'member': ProgramMember.objects.get(user=request.user, program=program),
                'edition': ProgramEdition.objects.get(pk=edition_id),
                'msc_student': MscStudent.objects.get(pk=student_id),
                'new_init_requirements': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'new_finish_requirements': ProgramFileDoc.objects.filter(program=program, is_finish_requirenment=True),
                'projects': InvestigationProject.objects.filter(program=program),
                'is_student': request.user == student.user,
            }
            return render(request, 'programs/edit_msc_student.html', context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar estudiantes de este programa.')



@login_required
def edit_dip_student(request, program_slug, edition_id, student_id):
    program= Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            user=DipStudent.objects.get(pk=student_id).user
            user.first_name=request.POST['student_name']
            user.last_name=request.POST['student_surename']
            user.email=request.POST['student_email']
            user.save()

            DipStudent.objects.filter(pk=student_id).update(
                phone=request.POST['student_phone'],
                country=request.POST['student_country'],
                gender=request.POST['gender'],
                dni=request.POST['student_dni'],
                birth_date=request.POST['student_birth_date']

            )
            if 'request_date' in request.POST and not request.POST['request_date'] == '':
                DipStudent.objects.filter(pk=student_id).update(
                    request_date=request.POST['request_date']
                )
            if 'init_date' in request.POST and not request.POST['init_date'] == '':
                DipStudent.objects.filter(pk=student_id).update(
                    init_date=request.POST['init_date']
                )
            if 'graduate_date' in request.POST and not request.POST['graduate_date'] == '':
                DipStudent.objects.filter(pk=student_id).update(
                    graduate_date=request.POST['graduate_date']
                )

            DipStudent.objects.filter(pk=student_id).update(
                status=request.POST['student_status']
            )



            try:
                if request.FILES['student_picture']:
                    student=DipStudent.objects.get(pk=student_id)
                    student.picture=request.FILES['student_picture']
                    student.save()
            except:
                pass

            # for requirement in ProgramInitRequirements.objects.filter(program=program):
            #     if 'student_requirement_' + str(requirement.id) in request.POST:
            #         s_i_r=StudentInitRequirement.objects.get(msc_student=MscStudent.objects.get(pk=student_id), requirement=requirement)
            #         s_i_r.accomplished=True
            #         s_i_r.save()
            #     else:
            #         s_i_r = StudentInitRequirement.objects.get(msc_student=MscStudent.objects.get(pk=student_id),
            #                                                    requirement=requirement)
            #         s_i_r.accomplished = False
            #         s_i_r.save()
            # for requirement in ProgramFinishRequirements.objects.filter(program=program):
            #     if 'student_f_requirement_' + str(requirement.id) in request.POST:
            #         s_f_r=StudentFinishRequirement.objects.get(msc_student=MscStudent.objects.get(pk=student_id), requirement=requirement)
            #         s_f_r.accomplished=True
            #         s_f_r.save()
            #     else:
            #         s_f_r = StudentFinishRequirement.objects.get(msc_student=MscStudent.objects.get(pk=student_id),
            #                                                      requirement=requirement)
            #         s_f_r.accomplished = False
            #         s_f_r.save()

            return HttpResponseRedirect(reverse('programs:dip_all_students_list', args=[program_slug,'all']))
        else:
            context = {
                'program': program,
                'edition': ProgramEdition.objects.get(pk=edition_id),
                'student': DipStudent.objects.get(pk=student_id),
                # 'init_requirements': ProgramInitRequirements.objects.filter(program=program),
                # 'finish_requirements': ProgramFinishRequirements.objects.filter(program=program),
                # 'projects': InvestigationProject.objects.filter(program=program),
            }
            return render(request, 'programs/edit_dip_student.html', context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar estudiantes de este programa.')


def error_500(request, program, error_message):
    context={
        'program':program,
        'error_message':error_message,
    }

    if user_is_program_member(request.user, program):
        context['member']=ProgramMember.objects.get(user=request.user, program=program)
    elif user_is_program_student(request.user, program):
        if program.type == 'phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type == 'msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type == 'dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)

    return render(request,'programs/error_500.html', context)

@login_required
def create_professor(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            try:

                user = User.objects.get(email=request.POST['email'])
                try:
                    professor = ProgramMember.objects.get(user=user, program=program)
                    return error_500(request, program,
                                     'Ya existe un profesor con este email en el sistema, verifique que no lo intenta duplicar.')
                except ProgramMember.DoesNotExist:
                    professor = ProgramMember(
                        program=program,
                        user=user,
                        role=request.POST['role'],
                        institution=request.POST['institution'],
                        degree=request.POST['grade'],
                        phone=request.POST['phone'],
                        country=request.POST['country'],
                        fb_contact=request.POST['fb_contact'],
                        tw_contact=request.POST['tw_contact'],
                        ln_contact=request.POST['ln_contact'],
                        sex=request.POST['gender'],
                        birth_date=request.POST['member_birth_date'],

                    )
                    if professor.role == 'Coordinador':
                        professor.weight = 1
                    elif professor.role == 'Secretario':
                        professor.weight = 2
                    elif professor.role == 'Miembro':
                        professor.weight = 3
                    elif professor.role == 'Profesor':
                        professor.weight = 4
                    elif professor.role == 'Tutor':
                        professor.weight = 5

                    try:
                        professor.save()
                        # send_mail('Hola','Usuario creado',program.email,[professor.user.email,'boris_perez@unah.edu.cu'], fail_silently=False)
                        utils_send_email(request, 'wm', program.email, professor, '', '', program, '********')
                        try:
                            professor.picture = request.FILES['picture']
                            professor.save()
                        except:
                            pass
                        return HttpResponseRedirect(reverse('programs:members_list', args=[program_slug, 'all']))
                    except:
                        return error_500(request, program, 'Ha ocurrido un error al guardar los datos del nuevo profesor')

                    return HttpResponseRedirect(reverse('programs:members_list', args=[program_slug, 'all']))

            except User.DoesNotExist:
                passwd = program_slug + str(random.randint(1000000, 9999999))
                user = User.objects.create_user(
                    request.POST['email'],
                    request.POST['email'],
                    passwd,  # Cambiar a password generada luego #
                )
                user.first_name = request.POST['name']
                user.last_name = request.POST['surename']
                user.save()

                professor = ProgramMember(
                    program=program,
                    user=user,
                    role=request.POST['role'],
                    institution=request.POST['institution'],
                    degree=request.POST['grade'],
                    phone=request.POST['phone'],
                    country=request.POST['country'],
                    fb_contact=request.POST['fb_contact'],
                    tw_contact=request.POST['tw_contact'],
                    ln_contact=request.POST['ln_contact'],
                    sex=request.POST['gender'],
                    birth_date=request.POST['member_birth_date'],

                )
                if professor.role == 'Coordinador':
                    professor.weight = 1
                elif professor.role == 'Secretario':
                    professor.weight = 2
                elif professor.role == 'Miembro':
                    professor.weight = 3
                elif professor.role == 'Profesor':
                    professor.weight = 4
                elif professor.role == 'Tutor':
                    professor.weight = 5

                try:
                    professor.save()
                    # send_mail('Hola','Usuario creado',program.email,[professor.user.email,'boris_perez@unah.edu.cu'], fail_silently=False)
                    utils_send_email(request, 'wm', program.email, professor, '', '', program, passwd)
                    try:
                        professor.picture = request.FILES['picture']
                        professor.save()
                    except:
                        pass
                    return HttpResponseRedirect(reverse('programs:members_list', args=[program_slug, 'all']))
                except:
                    pass

        else:
            context={
                'program':program,
                'member': ProgramMember.objects.get(user= request.user, program = program)
            }
            return render(request, 'programs/create_professor.html', context)
    else:
        return error_500(request,program, 'Usted no tiene privilegios para agregar profesores en este programa')
@login_required
def edit_member(request, program_slug, member_id):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            user=ProgramMember.objects.get(pk=member_id).user
            user.first_name=request.POST['name']
            user.last_name=request.POST['surename']

            try:
                usuario = User.objects.get(email=request.POST['email'])
                if usuario != ProgramMember.objects.get(pk=member_id).user:
                    return error_500(request,program,'El correo que ha agregado ya está en uso por otro usuario')
            except User.DoesNotExist:
                user.email=request.POST['email']

            user.save()
            ProgramMember.objects.filter(pk=member_id).update(
                role=request.POST['role'],
                institution=request.POST['institution'],
                degree=request.POST['grade'],
                phone=request.POST['phone'],
                country=request.POST['country'],
                fb_contact=request.POST['fb_contact'],
                tw_contact=request.POST['tw_contact'],
                ln_contact=request.POST['ln_contact'],
                sex=request.POST['gender'],
                birth_date=request.POST['member_birth_date'],
            )
            professor=ProgramMember.objects.get(pk=member_id)
            if request.POST['role'] == 'Coordinador':
                professor.weight = 1
            elif request.POST['role'] == 'Secretario':
                professor.weight = 2
            elif request.POST['role'] == 'Miembro':
                professor.weight = 3
            elif request.POST['role'] == 'Profesor':
                professor.weight = 4
            elif request.POST['role'] == 'Tutor':
                professor.weight = 5
            professor.save()
            try:
                professor.picture=request.FILES['picture']
                professor.save()
            except:
                print('Excepcion sin foto')

            return HttpResponseRedirect(reverse('programs:members_list',args=[program_slug, 'all']))
        else:
            context={
                'program':program,
                'member':ProgramMember.objects.get(pk=member_id)
            }
            return render(request,'programs/edit_professor.html',context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar miembros de este programa.')


#Devuelve 0 si el user no existe, 1 si existe pero no es miembro del programa, 2 si existe y es miembro del programa
@login_required
def ajx_usr_exists(request,program_slug):
    program = Program.objects.get(slug=program_slug)

    if request.method=='POST':

        try:
            user= User.objects.get(email=request.POST['email'])
            if user_is_program_member(user, program ) or user_is_program_student(user, program):
                return HttpResponse(
                    json.dumps([{'exists': 2}]),
                    content_type="application/json"
                )
            else:
                return HttpResponse(
                    json.dumps([{'exists': 1}]),
                    content_type="application/json"
                )


        except User.DoesNotExist:
            return HttpResponse(
                json.dumps([{'exists': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'exists': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_student_exists(request,program_slug):
    program = Program.objects.get(slug=program_slug)
    
    if request.method=='POST':
        if program.type == 'phd':

            try:
                student= Student.objects.get(user__username=request.POST['email'], program=program)

                return HttpResponse(
                    json.dumps([{'exists': 1}]),
                    content_type="application/json"
                )

            except Student.DoesNotExist:
                pass
        elif program.type == 'msc':
            try:
                student = MscStudent.objects.get(user__username=request.POST['email'],program=program)
                return HttpResponse(
                    json.dumps([{'exists': 1}]),
                    content_type="application/json"
                )

            except MscStudent.DoesNotExist:
                pass
        elif program.type == 'dip':
            try:
                student = DipStudent.objects.get(user__username=request.POST['email'],program=program)
                return HttpResponse(
                    json.dumps([{'exists': 1}]),
                    content_type="application/json"
                )

            except DipStudent.DoesNotExist:
                pass
    else:
        return HttpResponse(
            json.dumps([{'exists': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_program_member_tuthor(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    data_response=[]

    if request.method == 'POST':

        try:
            member=ProgramMember.objects.get(program=program, user__email=request.POST['tuthor_email'])
            data_response.append(1)
            data_response.append([member.user.first_name, member.user.last_name, member.institution])
        except ProgramMember.DoesNotExist:
            data_response.append(0)


    return HttpResponse(
        json.dumps(data_response),
        content_type="application/json"
    )



@login_required
def create_line(request, program_slug):
    program= Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user,program):
        if request.method=='POST':
            new_line = InvestigationLine(
                program=program,
                name=request.POST['line_name']
            )
            new_line.save()
            return HttpResponseRedirect(reverse('programs:program_lines', args=[program_slug]))

        else:
            context={
                'program':program,
                'member': ProgramMember.objects.get(user=request.user, program = program)
            }
            return render(request, 'programs/create_line.html', context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para crear líneas.')

@login_required
def edit_line(request, program_slug, line_id):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            InvestigationLine.objects.filter(pk=line_id).update(
                name=request.POST['line_name']
            )
            return HttpResponseRedirect(reverse('programs:program_lines', args=[program_slug]))
        else:
            context={
                'program':program,
                'line':InvestigationLine.objects.get(pk=line_id),
            }
            return render(request, 'programs/edit_line.html', context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar esta linea')




@login_required
def program_lines(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    context={
        'program': program,
        'lines': InvestigationLine.objects.filter(program=Program.objects.get(slug=program_slug)),
    }
    if user_is_program_member(request.user, program):
        context['member'] = ProgramMember.objects.get(user=request.user, program=program)
    elif user_is_program_student(request.user, program):
        if program.type =='phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type =='msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type =='dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)

    return render(request, 'programs/lines_list.html', context)


@login_required
def program_annoucements(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    context={
        'program': program,

    }
    if user_is_program_member(request.user, program):
        context['member'] = ProgramMember.objects.get(user=request.user, program=program)
        context['announcements'] = PhdAnnouncement.objects.filter(phd_student__student__program=program)

    elif user_is_program_student(request.user, program):
        if program.type =='phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
            context['announcements'] = PhdAnnouncement.objects.filter(phd_student__student__program= program)
        elif program.type =='msc':
            # TODO
            pass
        elif program.type =='dip':
            # TODO
            pass

    return render(request, 'programs/announcements_list.html', context)


@login_required
def students_by_line(request, program_slug, line_id):
    program=Program.objects.get(slug=program_slug)
    line=InvestigationLine.objects.get(pk=line_id)
    context={
        'program':program,
        'line':line,
        'students':Student.objects.filter(phdstudent__phdstudenttheme__line=line)
    }
    try:
        context['member']= ProgramMember.objects.get(user = request.user, program = program)
    except ProgramMember.DoesNotExist:
        pass

    if user_is_program_student(request.user, program):
        if program.type =='phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type =='msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type =='dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)

    return render(request, 'programs/students_by_line.html', context)

@login_required
def create_project(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method=='POST':
            new_project=InvestigationProject(
                program=program,
                line=InvestigationLine.objects.get(pk=request.POST['line']),
                name=request.POST['project_name'],
                institution=request.POST['project_institution'],
                init_date=request.POST['project_init_date'],
                end_date=request.POST['project_end_date'],
            )
            new_project.save()
            return HttpResponseRedirect(reverse('programs:projects_list', args=[program_slug]))
        else:
            context={
                'program':program,
                'lines':InvestigationLine.objects.filter(program=program),
                'member': ProgramMember.objects.get(user=request.user, program = program )
            }
            return render(request, 'programs/create_project.html', context)
    else:
        return error_500(request,program,'Usted no tiene proivilegios para crear proyectos')


@login_required
def create_phd_speciality(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if program.type == 'phd':
            if request.method == 'POST':
                new_speciality = ProgramSpeciality(
                    program=program,
                    name=request.POST['speciality_name'],
                    code=request.POST['speciality_code'],
                )
                new_speciality.save()
                return HttpResponseRedirect(reverse('programs:phd_specialities', args=[program_slug]))
            else:
                context = {
                    'program': program,
                    'member': ProgramMember.objects.get(user=request.user, program=program)
                }
                return render(request, 'programs/create_speciality.html', context)
        else:
            return error_500(request, program, 'Solo los doctorados tienen especialidades.')
    else:
        return error_500(request,program,'Usted no tiene privilegios para agregar especialidades aquí')

@login_required
def edit_phd_speciality(request, program_slug, speciality_id):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if program.type == 'phd':
            speciality = ProgramSpeciality.objects.get(pk=speciality_id)
            if request.method == 'POST':
                speciality.name = request.POST['speciality_name']
                speciality.code = request.POST['speciality_code']
                speciality.save()
                return HttpResponseRedirect(reverse('programs:phd_specialities', args=[program_slug]))
            else:
                context = {
                    'program': program,
                    'speciality':speciality,
                    'member': ProgramMember.objects.get(user=request.user, program=program)
                }
                return render(request, 'programs/edit_speciality.html', context)
        else:
            return error_500(request, program, 'Solo los doctorados tienen especialidades.')
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar especialidades aquí')


@login_required
def phd_specialities(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        context={
            'program':program,
            'specialities':ProgramSpeciality.objects.filter(program=program)
        }
        if user_is_program_member(request.user, program):
            context['member'] = ProgramMember.objects.get(user=request.user, program=program)
        elif user_is_program_student(request.user, program):
            context['student'] = Student.objects.get(user=request.user, program=program)
        else:
            return error_500(request, program, 'Sólo los estudiantes o profesores del programa pueden ver las especialidades')

        return render(request, 'programs/specialities_list.html',context)
    else:
        return error_500(request, program, 'Sólo los doctorados tienen especialidades')

@login_required
def projects_list(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    context={
        'program':program,
        'projects':InvestigationProject.objects.filter(program=program)
    }
    if user_is_program_member(request.user, program):
        context['member'] = ProgramMember.objects.get(user=request.user, program=program)
    elif user_is_program_student(request.user, program):
        if program.type =='phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type =='msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type =='dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)
    return render(request, 'programs/projects_list.html',context)

@login_required
def edit_project(request, program_slug, project_id):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method=='POST':
            InvestigationProject.objects.filter(pk=project_id).update(
                line=InvestigationLine.objects.get(pk=request.POST['line']),
                name=request.POST['project_name'],
                institution=request.POST['project_institution'],
                init_date=request.POST['project_init_date'],
                end_date=request.POST['project_end_date'],
            )

            return HttpResponseRedirect(reverse('programs:projects_list', args=[program_slug]))
        else:
            context={
                'program':program,
                'lines':InvestigationLine.objects.filter(program=program),
                'project':  InvestigationProject.objects.get(pk=project_id),
            }
            return render(request, 'programs/edit_project.html', context)
    else:
        return error_500(request,program,'Usted no tiene privilegios para editar proyectos')

@login_required
def ajx_delete_student(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            student_id=request.POST['student_id']
            if program.type == 'phd':
                try:
                    Student.objects.get(pk=student_id).delete()
                    return HttpResponse(
                        json.dumps([{'deleted': 1}]),
                        content_type="application/json"
                    )
                except:
                    return HttpResponse(
                        json.dumps([{'deleted': 0}]),
                        content_type="application/json"
                    )
            elif program.type == 'msc':
                try:
                    MscStudent.objects.get(pk=student_id).delete()
                    return HttpResponse(
                        json.dumps([{'deleted': 1}]),
                        content_type="application/json"
                    )
                except:
                    return HttpResponse(
                        json.dumps([{'deleted': 0}]),
                        content_type="application/json"
                    )
            elif program.type == 'dip':
                try:
                    DipStudent.objects.get(pk=student_id).delete()
                    return HttpResponse(
                        json.dumps([{'deleted': 1}]),
                        content_type="application/json"
                    )
                except:
                    return HttpResponse(
                        json.dumps([{'deleted': 0}]),
                        content_type="application/json"
                    )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )


@login_required
def ajx_delete_member(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            member_id=request.POST['member_id']
            try:
                ProgramMember.objects.get(pk=member_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_course_professor(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            professor_id=int(request.POST['professor_id'])
            print(professor_id)
            try:
                CourseProfessor.objects.get(pk=professor_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_program_edition(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            edition_id=request.POST['edition_id']
            try:
                ProgramEdition.objects.get(pk=edition_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )
@login_required
def ajx_create_tuthor(request, program_slug, student_id):
    program=Program.objects.get(slug=program_slug)

    if program.type == 'phd':
        student = PhdStudent.objects.get(student=Student.objects.get(pk=student_id, program=program))
    elif program.type == 'msc':
        student = MscStudent.objects.get(pk=student_id, program=program)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':

            try:
                new_tuthor = create_new_tuthor(request, program,request.POST['tuthor_name'],request.POST['tuthor_lastname'],
                                               request.POST['tuthor_institution'],request.POST['tuthor_email'],student )
                if new_tuthor[0]:
                    tuthor = Tuthor.objects.get(pk=new_tuthor[1])
                    return HttpResponse(
                        json.dumps([{'created': 1,'id':tuthor.id, 'name':tuthor.professor.user.first_name,'lastname':tuthor.professor.user.last_name,
                                     'email':tuthor.professor.user.email,'institution':tuthor.professor.institution}]),
                        content_type="application/json"
                    )
                else:
                    return HttpResponse(
                        json.dumps([{'created': 0}]),
                        content_type="application/json"
                    )

            except:
                return HttpResponse(
                    json.dumps([{'created': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'created': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'created': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_create_activity(request, program_slug, student_id):
    program=Program.objects.get(slug=program_slug)
    student=Student.objects.get(pk=student_id)

    if program.type == 'phd':
        if request.user == student.user:
            if request.method == 'POST':
                formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                new_activity = FormationPlanActivities(
                    formation_plan=formation_plan,
                    init_date=request.POST['init_date'],
                    end_date=request.POST['end_date'],
                    description=request.POST['description'],

                )
                new_activity.save()
                formation_plan.last_update_date = now().date()
                formation_plan.save()
                return HttpResponse(
                    json.dumps([{'created': 1,'init_date':new_activity.init_date,'end_date':new_activity.end_date,
                                 'description':new_activity.description, 'id':new_activity.id}]),
                    content_type="application/json"
                )

            else:
                return HttpResponse(
                    json.dumps([{'created': 0}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'created': 3}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'created': 2}]),
            content_type="application/json"
        )

def ajx_change_activity_status(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)
    if request.user == student.user:
        if request.method == 'POST':
            activity = FormationPlanActivities.objects.get(pk=request.POST['activity_id'])
            if activity.status == 'pending':
                activity.status = 'ready'
            else:
                activity.status = 'pending'
            activity.save()
            return HttpResponse(
                json.dumps([{'edited': 1, 'status':activity.status }]),
                content_type="application/json"
            )
        else:
            return HttpResponse(
                json.dumps([{'edited': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'edited': 2}]),
            content_type="application/json"
        )


@login_required
def ajx_update_filedoc(request, program_slug):
    form = FileUploadForm(data=request.POST, files=request.FILES)

    program=Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        student = Student.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(student=student, program_file_document= ProgramFileDoc.objects.get(pk=request.POST['doc_id']))

    elif  program.type == 'msc':
        student = MscStudent.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(msc_student=student, program_file_document= ProgramFileDoc.objects.get(pk=request.POST['doc_id']))

    elif  program.type == 'dip':
        student = DipStudent.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(dip_student=student, program_file_document= ProgramFileDoc.objects.get(pk=request.POST['doc_id']))

    if request.method == 'POST':
        if student.user == request.user or user_is_program_cs(request.user, program):
            try:
                filedoc.file.delete()
            except:
                print('Archivo no eliminado por algopull ')
                pass

            try:
                file = request.FILES['file']
                doc_ext = file.name.split('.')[:: - 1][0]
                doc_name = slugify(file.name[:file.name.rindex('.')])

                file.name = doc_name+'.'+doc_ext
                filedoc.file = file
                filedoc.save()
                return HttpResponse(
                    json.dumps([{'updated': 1}]),
                    content_type="application/json"
                )
            except:
                print('Problemas con el archivo')
                return HttpResponse(
                    json.dumps([{'updated': 0, 'errors':form.errors}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'updated': 3, 'errors':form.errors}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': 2, 'errors':form.errors}]),
            content_type="application/json"
        )

@login_required
def ajx_upgrade_filedoc(request, program_slug):
    form = FileUploadForm(data=request.POST, files=request.FILES)

    program=Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        student = Student.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(student=student, program_file_document= ProgramFileDoc.objects.get(pk=request.POST['doc_id']))

    elif  program.type == 'msc':
        student = MscStudent.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(msc_student=student, program_file_document= ProgramFileDoc.objects.get(pk=request.POST['doc_id']))

    elif  program.type == 'dip':
        student = DipStudent.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(dip_student=student, program_file_document= ProgramFileDoc.objects.get(pk=request.POST['doc_id']))

    if request.method == 'POST':
        if student.user == request.user or user_is_program_cs(request.user, program):


            try:
                doc_file = request.FILES['file']
                fs = FileSystemStorage()
                doc_ext = doc_file.name.split('.')[:: - 1][0]
                doc_name = slugify(doc_file.name[:doc_file.name.rindex('.')])
                doc_file_name = doc_name +'.'+ doc_ext
                new_doc_name = 'program_{0}/students/{1}/docs/{2}'.format(program.slug,student.id, doc_file_name)


                filename = fs.save(new_doc_name, doc_file)
                filedoc.file.delete()
                filedoc.file = filename
                filedoc.save()

                return HttpResponse(
                    json.dumps([{'updated': 1}]),
                    content_type="application/json"
                )
            except:
                print('Problemas con el archivo')
                return HttpResponse(
                    json.dumps([{'updated': 0, 'errors':form.errors}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'updated': 3, 'errors':form.errors}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': 2, 'errors':form.errors}]),
            content_type="application/json"
        )

@login_required
def ajx_upload_background(request, program_slug):
    form = FileUploadForm(data=request.POST, files=request.FILES)

    program=Program.objects.get(slug=program_slug)

    if request.method == 'POST':
        if user_is_program_cs(request.user, program):
            try:
                new_bg = ProgramBackgrounds(
                    program = program,
                    background= request.FILES['program_background']
                )

                new_bg.save()
                return HttpResponse(
                    json.dumps([{'updated': '1','bg_id':new_bg.id, 'bg_name':new_bg.background.name}]),
                    content_type="application/json"
                )
            except:
                print('Problemas con el archivo del nuevo bg')
                return HttpResponse(
                    json.dumps([{'updated': '4', 'errors':form.errors}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'updated': '3', 'errors':form.errors}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': '2', 'errors':form.errors}]),
            content_type="application/json"
        )

@login_required
def ajx_new_requirenment(request, program_slug, scope):
    form = FileUploadForm(data=request.POST, files=request.FILES)

    program=Program.objects.get(slug=program_slug)

    if request.method == 'POST':
        if user_is_program_cs(request.user, program):
            if scope == 'init':
                try:
                    new_init_r = ProgramFileDoc(
                        program = program,
                        doc_name=request.POST['init_r_name'],
                        is_init_requirenment=True,
                        type=request.POST['init_r_type']
                    )
                    try:
                        jfn = request.POST['init_r_nat']
                        new_init_r.just_for_nationals = True
                    except:
                        pass

                    try:
                        get_old = request.POST['init_r_get_old']
                        new_init_r.get_old = True
                    except:
                        pass

                    new_init_r.save()
                    return HttpResponse(
                        json.dumps([{'updated': '1','r_id':new_init_r.id, 'r_name':new_init_r.doc_name,
                                     'r_get_old':new_init_r.get_old, 'r_jfn': new_init_r.just_for_nationals, 'r_type':new_init_r.type}]),
                        content_type="application/json"
                    )
                except:
                    print('Problemas con el requisito')
                    return HttpResponse(
                        json.dumps([{'updated': '0', 'errors':form.errors}]),
                        content_type="application/json"
                    )
            elif scope == 'finish':
                try:
                    new_finish_r = ProgramFileDoc(
                        program=program,
                        doc_name=request.POST['finish_r_name'],
                        is_finish_requirenment=True,
                        type=request.POST['finish_r_type']
                    )
                    try:
                        jfn = request.POST['finish_r_nat']
                        new_finish_r.just_for_nationals = True
                    except:
                        pass

                    try:
                        get_old = request.POST['finish_r_get_old']
                        new_finish_r.get_old = True
                    except:
                        pass

                    new_finish_r.save()
                    return HttpResponse(
                        json.dumps([{'updated': '1', 'r_id': new_finish_r.id, 'r_name': new_finish_r.doc_name,
                                     'r_get_old': new_finish_r.get_old, 'r_jfn': new_finish_r.just_for_nationals,
                                     'r_type': new_finish_r.type}]),
                        content_type="application/json"
                    )
                except:
                    print('Problemas con el requisito')
                    return HttpResponse(
                        json.dumps([{'updated': '0', 'errors': form.errors}]),
                        content_type="application/json"
                    )

            else:
                return HttpResponse(
                    json.dumps([{'updated': '4', 'errors': form.errors}]),
                    content_type="application/json"
                )


        else:
            return HttpResponse(
                json.dumps([{'updated':'3', 'errors':form.errors}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': '2', 'errors':form.errors}]),
            content_type="application/json"
        )


@login_required
def ajx_edit_requirenment(request, program_slug):
    form = FileUploadForm(data=request.POST, files=request.FILES)

    program = Program.objects.get(slug=program_slug)

    if request.method == 'POST':
        if user_is_program_cs(request.user, program):

            try:
                r = ProgramFileDoc.objects.get(pk=request.POST['r_id'])
                r.doc_name = request.POST['edit_r_name']
                r.type = request.POST['edit_r_type']

                try:
                    jfn = request.POST['edit_r_nat']
                    r.just_for_nationals = True
                except:
                    r.just_for_nationals = False

                try:

                    get_old = request.POST['edit_r_get_old']
                    print(get_old)
                    r.get_old = True
                except:
                    r.get_old = False

                r.save()
                r_data = {'updated': '1', 'r_id': r.id, 'r_name': r.doc_name,
                          'r_get_old': r.get_old, 'r_jfn': r.just_for_nationals,
                          'r_type': r.type}
                if r.is_init_requirenment:
                    r_data['r_context'] = 'init'
                elif r.is_finish_requirenment:
                    r_data['r_context'] = 'finish'
                return HttpResponse(
                    json.dumps([r_data]),
                    content_type="application/json"
                )
            except ProgramFileDoc.DoesNotExist:
                print('Problemas con el requisito')
                return HttpResponse(
                    json.dumps([{'updated': '0', 'errors': form.errors}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'updated': '3', 'errors': form.errors}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': '2', 'errors': form.errors}]),
            content_type="application/json"
        )

@login_required
def ajx_r_data(request, program_slug):

    program = Program.objects.get(slug=program_slug)

    if request.method == 'POST':
        if user_is_program_cs(request.user, program):
            r = ProgramFileDoc.objects.get(pk=request.POST['r_id'])

            r_data = {'recovered': '1', 'r_id': r.id, 'r_name': r.doc_name,
                      'r_get_old': r.get_old, 'r_jfn': r.just_for_nationals,
                      'r_type': r.type}
            if r.is_init_requirenment:
                r_data['r_context'] = 'init'
            elif r.is_finish_requirenment:
                r_data['r_context'] = 'finish'

            json_data = []
            json_data.append(r_data)
            return HttpResponse(

                json.dumps(json_data),
                content_type="application/json"
            )


        else:
            return HttpResponse(
                json.dumps([{'recovered': '3'}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'recovered': '2'}]),
            content_type="application/json"
        )

@login_required
def ajx_edit_background(request, program_slug):
    form = FileUploadForm(data=request.POST, files=request.FILES)

    program=Program.objects.get(slug=program_slug)

    if request.method == 'POST':
        try:
            bg = ProgramBackgrounds.objects.get(pk=request.POST['bg_id'])
        except:
            return HttpResponse(
                json.dumps([{'updated': '4', 'errors': form.errors}]),
                content_type="application/json"
            )


        if user_is_program_cs(request.user, program):
            try:
                bg.background.delete()
            except:
                pass

            try:
                bg.background = request.FILES['program_background']
                bg.save()
                return HttpResponse(
                    json.dumps([{'updated': '1','bg_id':bg.id, 'bg_name': bg.background.name}]),
                    content_type="application/json"
                )
            except:
                print('Problemas con el archivo')
                return HttpResponse(
                    json.dumps([{'updated': '0', 'errors':form.errors}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'updated': '3', 'errors':form.errors}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': '2', 'errors':form.errors}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_bg(request, program_slug):
    program = Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user, program ):
        if request.method=='POST':
            bg = ProgramBackgrounds.objects.get(pk=request.POST['bg_id'])
            try:
                bg.background.delete()
            except:
                pass

            try:

                bg.delete()
                return HttpResponse(
                    json.dumps([{'deleted': '1'}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': '2'}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': '0'}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': '3'}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_r(request, program_slug):
    program = Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user, program ):
        if request.method=='POST':
            r = ProgramFileDoc.objects.get(pk=request.POST['r_id'])

            try:
                r_context = ''
                if r.is_init_requirenment:
                    r_context = 'init'
                elif r.is_finish_requirenment:
                    r_context = 'finish'

                r.delete()
                return HttpResponse(
                    json.dumps([{'deleted': '1', 'r_context':r_context}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': '2'}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': '0'}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': '3'}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_filedoc(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        student = Student.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(student=student, program_file_document=ProgramFileDoc.objects.get(
            pk=request.POST['doc_id']))

    elif program.type == 'msc':
        student = MscStudent.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(msc_student=student, program_file_document=ProgramFileDoc.objects.get(
            pk=request.POST['doc_id']))

    elif program.type == 'dip':
        student = DipStudent.objects.get(pk=request.POST['student_id'])
        filedoc = StudentFileDocument.objects.get(dip_student=student, program_file_document=ProgramFileDoc.objects.get(
            pk=request.POST['doc_id']))

    if request.method == 'POST':
        if student.user == request.user or user_is_program_cs(request.user, program):
            try:
                filedoc.file.delete()
                filedoc.save()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                print('Problemas con el archivo')
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'updated': 3}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'updated': 2}]),
            content_type="application/json"
        )


def ajx_edit_activity(request, program_slug, student_id):
    program=Program.objects.get(slug=program_slug)
    student=Student.objects.get(pk=student_id)
    activity = FormationPlanActivities.objects.get(pk=request.POST['activity_id'])

    if program.type == 'phd':
        if request.user == student.user:
            if request.method == 'POST':
                activity.init_date=request.POST['init_date']
                activity.end_date=request.POST['end_date']
                activity.description=request.POST['description']
                activity.status = request.POST['status']
                activity.save()
                formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                formation_plan.last_update_date = now().date()
                formation_plan.save()
                return HttpResponse(
                    json.dumps([{'edited': 1,'init_date':activity.init_date,'end_date':activity.end_date,
                                 'description':activity.description,'status':activity.status, 'id':activity.id}]),
                    content_type="application/json"
                )

            else:
                return HttpResponse(
                    json.dumps([{'edited': 0}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'edited': 3}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'edited': 2}]),
            content_type="application/json"
        )

def ajx_delete_activity(request, program_slug, student_id):
    program=Program.objects.get(slug=program_slug)
    student=Student.objects.get(pk=student_id)
    activity = FormationPlanActivities.objects.get(pk=request.POST['activity_id'])

    if program.type == 'phd':
        if request.user == student.user:
            if request.method == 'POST':
                try:
                    activity.delete()
                    formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                    formation_plan.last_update_date = now().date()
                    formation_plan.save()

                    return HttpResponse(
                        json.dumps([{'deleted': 1}]),
                        content_type="application/json"
                    )
                except:
                    return HttpResponse(
                        json.dumps([{'deleted': 0}]),
                        content_type="application/json"
                    )

            else:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )

        else:
            return HttpResponse(
                json.dumps([{'deleted': 3}]),
                content_type="application/json"
            )

    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_import_courses(request, program_slug, edition_id):
    program=Program.objects.get(slug=program_slug)
    edition = ProgramEdition.objects.get(pk=edition_id)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            origin_edition = ProgramEdition.objects.get(pk=request.POST['edition_id'])
            courses=[]

            for course in Course.objects.filter(program=program, edition=origin_edition):
                new_course = Course(
                    program=program,
                    edition=edition,
                    name=course.name,
                    description=course.description,
                )
                new_course.save()
                for professor in course.courseprofessor_set.all():
                    new_professor = CourseProfessor(
                        course=new_course,
                        professor=professor.professor,
                    )
                    new_professor.save()

                courses.append(new_course.name)

            return HttpResponse(
                json.dumps([{'imported': 1}, courses]),
                content_type="application/json"
            )


        else:
            return HttpResponse(
                json.dumps([{'imported': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'imported': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_tuthor(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            tuthor_id=request.POST['tuthor_id']
            try:
                Tuthor.objects.get(pk=tuthor_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )


@login_required
def ajx_delete_line(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            line_id=request.POST['line_id']
            try:
                InvestigationLine.objects.get(pk=line_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )


@login_required
def ajx_delete_announcement(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            announcement_id=request.POST['announcement_id']
            try:
                PhdAnnouncement.objects.get(pk=announcement_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )


@login_required
def ajx_delete_speciality(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            speciality_id=request.POST['speciality_id']
            try:
                ProgramSpeciality.objects.get(pk=speciality_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 2}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 3}]),
            content_type="application/json"
        )


@login_required
def ajx_mark_message_readed(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    message = Message.objects.get(pk=request.POST['message_id'])

    if program.type == 'phd':
        try:
            receiver_user = message.phd_student_receiver.user
        except:
            receiver_user = message.program_receiver.user


    elif program.type == 'msc':
        try:
            receiver_user = message.msc_student_receiver.user
        except:
            receiver_user = message.program_receiver.user
    elif program.type == 'dip':
        try:
            receiver_user = message.dip_student_receiver.user
        except:
            receiver_user = message.program_receiver.user

    if request.user == receiver_user or request.user == message.sender:
        if request.method=='POST':
            try:
                message.readed = True
                message.save()
                return HttpResponse(
                    json.dumps([{'readed': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'readed': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'readed': 2}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'readed': 3}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_message(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    message = Message.objects.get(pk=request.POST['message_id'])

    if program.type == 'phd':
        try:
            receiver_user = message.phd_student_receiver.user
        except:
            receiver_user = message.program_receiver.user

    elif program.type == 'msc':
        try:
            receiver_user = message.msc_student_receiver.user
        except:
            receiver_user = message.program_receiver.user
    elif program.type == 'dip':
        try:
            receiver_user = message.dip_student_receiver.user
        except:
            receiver_user = message.program_receiver.user

    if request.user == receiver_user:
        if request.method=='POST':
            try:
                message.delete()
                if program.type == 'phd':
                    messages_count = Message.objects.filter(phd_student_receiver__user=receiver_user).__len__()
                elif program.type == 'msc':
                    messages_count = Message.objects.filter(msc_student_receiver__user=receiver_user).__len__()
                elif program.type == 'dip':
                    messages_count = Message.objects.filter(dip_student_receiver__user=receiver_user).__len__()
                return HttpResponse(
                    json.dumps([{'deleted': 1,'messages_count': messages_count}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 2}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 3}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_course(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            course_id=request.POST['component_id']
            try:
                Course.objects.get(pk=course_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_new(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            new_id=request.POST['new_id']
            try:
                New.objects.get(pk=new_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 2}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 3}]),
            content_type="application/json"
        )

@login_required
def ajx_edit_eval(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            eval_id=request.POST['eval_id']
            try:
                evaluation = CourseEvaluation.objects.get(pk=eval_id)
                evaluation.value = request.POST['eval']
                evaluation.save()
                return HttpResponse(
                    json.dumps([{'saved': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'saved': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'saved': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'saved': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_eval(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            eval_id=request.POST['eval_id']
            try:
                CourseEvaluation.objects.get(pk=eval_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )

@login_required
def ajx_delete_project(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            project_id=request.POST['project_id']
            try:
                InvestigationProject.objects.get(pk=project_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_this_year_requests(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []
    data_1=[]
    data_2=[]
    meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 7: "Julio",
             8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for i in range(1,now().month+1):
        labels.append(meses[i])
        if program.type == 'phd':
            data_1.append(PhdStudent.objects.filter(Q(status='solicitante')|Q(status='doctorando')|Q(status='graduado'),student__program=program, student__request_date__year=now().year,student__request_date__month=i).__len__())
            data_2.append(PhdStudent.objects.filter(Q(status='doctorando')|Q(status='graduado'),student__program=program, student__init_date__year=now().year,student__init_date__month=i).__len__())
        elif program.type == 'msc':
            data_1.append(MscStudent.objects.filter(program=program,request_date__year=now().year, request_date__month=i).__len__())
            data_2.append(MscStudent.objects.filter(program=program, init_date__year=now().year, init_date__month=i).__len__())
        elif program.type == 'dip':
            data_1.append(DipStudent.objects.filter(program=program, request_date__year=now().year,
                                                    request_date__month=i).__len__())
            data_2.append(
                DipStudent.objects.filter(program=program, init_date__year=now().year, init_date__month=i).__len__())

    data.append(data_1)
    data.append(data_2)

    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_by_year_requests(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    year=request.POST['year']
    response_data=[]
    labels = []
    data = []
    data_1=[]
    data_2=[]
    meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 7: "Julio",
             8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    if int(year) != now().year:
        for i in range(1, 13):
            labels.append(meses[i])
            if program.type == 'phd':
                data_1.append(
                    Student.objects.filter(program=program, request_date__year=year, request_date__month=i).__len__())
                data_2.append(
                    Student.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())
            elif program.type == 'msc':
                data_1.append(MscStudent.objects.filter(program=program, request_date__year=year,
                                                        request_date__month=i).__len__())
                data_2.append(
                    MscStudent.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())
            elif program.type == 'dip':
                data_1.append(DipStudent.objects.filter(program=program, request_date__year=year,
                                                        request_date__month=i).__len__())
                data_2.append(
                    DipStudent.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())
    else:
        for i in range(1, now().month + 1):
            labels.append(meses[i])
            if program.type == 'phd':
                data_1.append(Student.objects.filter(program=program, request_date__year=year,
                                                     request_date__month=i).__len__())
                data_2.append(
                    Student.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())
            elif program.type == 'msc':
                data_1.append(MscStudent.objects.filter(program=program, request_date__year=year,
                                                        request_date__month=i).__len__())
                data_2.append(
                    MscStudent.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())
            elif program.type == 'dip':
                data_1.append(DipStudent.objects.filter(program=program, request_date__year=year,
                                                        request_date__month=i).__len__())
                data_2.append(
                    DipStudent.objects.filter(program=program, init_date__year=year, init_date__month=i).__len__())

    data.append(data_1)
    data.append(data_2)

    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_students_by_country(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    countries = []
    data = []
    if program.type == 'phd':
        for student in Student.objects.filter(program=program):
            if not student.country in countries:
                countries.append(student.country)

        for country in countries:
            data.append(Student.objects.filter(program=program,country=country).__len__())

    if program.type == 'msc':
        for student in MscStudent.objects.filter(program=program):
            if not student.country in countries:
                countries.append(student.country)
        for country in countries:
            data.append(MscStudent.objects.filter(program=program,country=country).__len__())
    if program.type == 'dip':
        for student in DipStudent.objects.filter(program=program):
            if not student.country in countries:
                countries.append(student.country)

        for country in countries:
            data.append(DipStudent.objects.filter(program=program,country=country).__len__())

    response_data.append(countries)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_students_by_edition(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for edition in ProgramEdition.objects.filter(program=program):
        labels.append('Edición '+str(edition.order))
        if program.type == 'msc':
            data.append(MscStudent.objects.filter(Q(status='maestrante')|Q(status='graduado'), program=program, edition=edition).__len__())
        elif program.type == 'dip':
            data.append(DipStudent.objects.filter(Q(status='diplomante')|Q(status='graduado'), program=program, edition=edition).__len__())



    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_graduated_by_edition(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data = []
    labels = []
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for edition in ProgramEdition.objects.filter(program=program):
        labels.append('Edición ' + str(edition.order))
        if program.type == 'msc':
            data.append(
                MscStudent.objects.filter(program=program, edition=edition, status='graduado').__len__())
        elif program.type == 'dip':
            data.append(
                DipStudent.objects.filter(program=program, edition=edition, status='graduado').__len__())

    response_data.append(labels)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_last_years_requests(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []
    data_1=[]
    data_2=[]

    # locale.setlocale(locale.LC_ALL, 'es-ES')

    for i in range(now().year-4,now().year+1):
        labels.append(i)
        if program.type == 'phd':
            data_1.append(PhdStudent.objects.filter(Q(status='solicitante') | Q(status='doctorando') | Q(status='graduado'), student__program=program, student__request_date__year=i).__len__())
            data_2.append(PhdStudent.objects.filter(Q(status='doctorando') | Q(status='graduado'), student__program=program, student__init_date__year=i).__len__())

        elif program.type == 'msc':
            data_1.append(MscStudent.objects.filter(Q(status='solicitante')| Q(status='maestrante') | Q(status='graduado'),program=program,request_date__year=i).__len__())
            data_2.append(MscStudent.objects.filter(Q(status='maestrante') | Q(status='graduado'),program=program,init_date__year=i).__len__())
        elif program.type == 'dip':
            data_1.append(DipStudent.objects.filter(Q(status='solicitante')| Q(status='diplomante') | Q(status='graduado'),program=program,request_date__year=i).__len__())
            data_2.append(DipStudent.objects.filter(Q(status='diplomante') | Q(status='graduado'),program=program,init_date__year=i).__len__())

    data.append(data_1)
    data.append(data_2)

    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )
@login_required
def ajx_line_projects(request, program_slug):
    if request.method == 'POST':
        line = InvestigationLine.objects.get(pk=request.POST['line_id'])

        if InvestigationProject.objects.filter(line=line).__len__()>0:
            json_data = [{'status': 1}]
            for project in InvestigationProject.objects.filter(line=line):
                project_data = {}
                project_data['project_id']=project.id
                project_data['project_name']=project.name
                json_data.append(project_data)

            return HttpResponse(
                json.dumps(json_data),
                content_type="application/json"
            )
        else:
            return HttpResponse(
                json.dumps([{'status': 0}]),
                content_type="application/json"
            )


    else:
        return HttpResponse(
            json.dumps([{'status': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_students_by_line(request, program_slug, scope):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    i=0
    for line in InvestigationLine.objects.filter(program=Program.objects.get(slug=program_slug)):
        i += 1
        labels.append(line.name.split()[0])
        if program.type == 'phd':
            if scope == 'all':
                data.append(PhdStudentTheme.objects.filter(line=line).__len__())
            elif scope == 'requesters':
                data.append(PhdStudentTheme.objects.filter(line=line, phd_student__status='solicitante').__len__())
            elif scope == 'aproved':
                data.append(PhdStudentTheme.objects.filter(line=line, phd_student__status='doctorando').__len__())
            elif scope == 'graduated':
                data.append(PhdStudentTheme.objects.filter(line=line, phd_student__status='graduado').__len__())
        elif program.type == 'msc':
            if scope == 'all':
                data.append(MscStudentTheme.objects.filter(line=line).__len__())
            elif scope == 'requesters':
                data.append(MscStudentTheme.objects.filter(line=line, student__status ='solicitante').__len__())
            elif scope == 'aproved':
                data.append(MscStudentTheme.objects.filter(line=line, student__status ='maestrante').__len__())
            elif scope == 'graduated':
                data.append(MscStudentTheme.objects.filter(line=line, student__status ='graduado').__len__())



    response_data.append(labels)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_students_by_age(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = ['<30 años','30-40','40-50','>50 años']
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    i=0
    if program.type == 'phd':
        data.append(Student.objects.filter(program=program,birth_date__year__gte=now().year-30 ).__len__())
        data.append(Student.objects.filter(program=program,birth_date__year__gte=now().year-40, birth_date__year__lt=now().year-30 ).__len__() )
        data.append(Student.objects.filter(program=program,birth_date__year__gte=now().year-50, birth_date__year__lt=now().year-40 ).__len__() )
        data.append(Student.objects.filter(program=program, birth_date__year__lte=now().year-50 ).__len__() )
    elif program.type == 'msc':
        data.append(MscStudent.objects.filter(program=program, birth_date__year__gte=now().year - 30).__len__())
        data.append(MscStudent.objects.filter(program=program, birth_date__year__gte=now().year - 40,
                                              birth_date__year__lt=now().year - 30).__len__())
        data.append(MscStudent.objects.filter(program=program, birth_date__year__gte=now().year - 50,
                                              birth_date__year__lt=now().year - 40).__len__())
        data.append(MscStudent.objects.filter(program=program, birth_date__year__lte=now().year - 50).__len__())
    elif program.type == 'dip':
        data.append(DipStudent.objects.filter(program=program, birth_date__year__gte=now().year - 30).__len__())
        data.append(DipStudent.objects.filter(program=program, birth_date__year__gte=now().year - 40,
                                              birth_date__year__lt=now().year - 30).__len__())
        data.append(DipStudent.objects.filter(program=program, birth_date__year__gte=now().year - 50,
                                              birth_date__year__lt=now().year - 40).__len__())
        data.append(DipStudent.objects.filter(program=program, birth_date__year__lte=now().year - 50).__len__())


    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_next_years_defenses(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    response_data=[]
    labels = []
    data = []

    for year in range(now().year,now().year+6):
        labels.append(year)
        if program.type=='phd':
            data.append(StudentFormationPlan.objects.filter(phdstudent__program=program,phdstudent__phdstudent__status='Doctorando', planned_end_year=year).__len__())
        else:
            data.append('Error')

    response_data.append(labels)
    response_data.append(data)


    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_member_personal_msg(request, program_slug ):
    program=Program.objects.get(slug=program_slug)
    if request.method == 'POST' and request.POST['msg_body'].__len__() <= 4000:
        try:
            new_message = Message(
                sender=request.user,
                program_receiver=ProgramMember.objects.get(pk=request.POST['member_id']),
                subject=request.POST['msg_subject'],
                body=request.POST['msg_body'],

            )
            new_message.save()
            sended_message = MessageSended(
                sender=request.user,
                context='personal',
                program_receiver=ProgramMember.objects.get(pk=request.POST['member_id']),
                subject=request.POST['msg_subject'],
                body=request.POST['msg_body'],

            )
            sended_message.save()
            send_mail(request.POST['msg_subject'], request.POST['msg_body'],request.user.email,
                      [ProgramMember.objects.get(pk=request.POST['member_id']).user.email],
                      fail_silently=False,html_message=request.POST['msg_body'])
            return HttpResponse(
                json.dumps([{'sended': 1}]),
                content_type="application/json"
            )
        except:
            return HttpResponse(
                json.dumps([{'sended': 0}]),
                content_type="application/json"
            )
    elif request.method == 'POST' and request.POST['msg_body'].__len__() > 4000:
        return HttpResponse(
            json.dumps([{'sended': 2}]),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps([{'sended': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_member_massive_msg(request, program_slug ):
    program=Program.objects.get(slug=program_slug)
    if request.method == 'POST' and request.POST['msg_body'].__len__() <= 4000:
        try:
            email_list = []
            if request.POST['msg_scope'] == 'comite':
                for member in ProgramMember.objects.filter(Q(role='Coordinador')|Q(role='Secretario')|Q(role='Miembro'), program=program  ):
                    email_list.append(member.user.email)
                    new_message = Message(
                        sender=request.user,
                        program_receiver=member,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()
            elif request.POST['msg_scope']=='professors':
                for member in ProgramMember.objects.filter(program=program , role='Profesor' ):
                    email_list.append(member.user.email)
                    new_message = Message(
                        sender=request.user,
                        program_receiver=member,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()
            elif request.POST['msg_scope']=='tuthors':
                for member in ProgramMember.objects.filter(program=program , role='Tutor' ):
                    email_list.append(member.user.email)
                    new_message = Message(
                        sender=request.user,
                        program_receiver=member,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()

            elif request.POST['msg_scope']=='all':
                for member in ProgramMember.objects.filter(program=program ):
                    email_list.append(member.user.email)
                    new_message = Message(
                        sender=request.user,
                        program_receiver=member,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()

            if email_list.__len__()<=20:
                send_mail(request.POST['msg_subject'], request.POST['msg_body'],request.user.email,
                          email_list, fail_silently=False, html_message=request.POST['msg_body'])
            else:
                count = email_list.__len__() // 4
                rest = email_list.__len__() % 4

                for i in range(email_list.__len__() // count):
                    print(i)
                    send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                              request.user.email,
                              email_list[count * i:count * (i + 1)], fail_silently=False, html_message=request.POST['msg_body'])

                    if rest != 0:
                        send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                                  request.user.email,
                                  email_list[4 * count:email_list.__len__()], fail_silently=False,
                                  html_message=request.POST['msg_body'])

            sended_message = MessageSended(
                sender=request.user,
                context='profesores',
                subject=request.POST['msg_subject'],
                body=request.POST['msg_body'],

            )
            sended_message.save()

            return HttpResponse(
                json.dumps([{'sended': 1}]),
                content_type="application/json"
            )
        except:
            return HttpResponse(
                json.dumps([{'sended': 0}]),
                content_type="application/json"
            )
    elif request.method == 'POST' and request.POST['msg_body'].__len__() > 4000:
        return HttpResponse(
            json.dumps([{'sended': 2}]),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps([{'sended': 0}]),
            content_type="application/json"
        )


@login_required
def ajx_all_massive_msg(request, program_slug ):
    program=Program.objects.get(slug=program_slug)
    if request.method == 'POST' and request.POST['msg_body'].__len__() <= 4000:
        try:
            email_list = []
            for member in ProgramMember.objects.filter(program=program):
                email_list.append(member.user.email)

                new_message = Message(
                    sender=request.user,
                    program_receiver=member,
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                new_message.save()

            if program.type == 'phd':
                for student in Student.objects.filter(program=program):
                    email_list.append(student.user.email)
                    new_message = Message(
                        sender=request.user,
                        phd_student_receiver=student,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()
            elif program.type == 'msc':
                for student in MscStudent.objects.filter(program=program):
                    email_list.append(student.user.email)
                    new_message = Message(
                        sender=request.user,
                        msc_student_receiver=student,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()
            elif program.type == 'dip':
                for student in DipStudent.objects.filter(program=program):
                    email_list.append(student.user.email)
                    new_message = Message(
                        sender=request.user,
                        dip_student_receiver=student,
                        subject=request.POST['msg_subject'],
                        body=request.POST['msg_body'],

                    )
                    new_message.save()

            if email_list.__len__()<=20:
                send_mail(request.POST['msg_subject'], request.POST['msg_body'],request.user.email,
                          email_list, fail_silently=False, html_message=request.POST['msg_body'])
            else:
                send_number = email_list.__len__() // 4
                rest = email_list.__len__() % 4

                for i in range(email_list.__len__() // send_number):
                    send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                              request.user.email,
                              email_list[send_number * i:send_number * (i + 1)], fail_silently=False, html_message=request.POST['msg_body'])

                if rest != 0:
                     send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                              request.user.email,
                              email_list[4* send_number:email_list.__len__()], fail_silently=False,
                              html_message=request.POST['msg_body'])



            return HttpResponse(
                json.dumps([{'sended': 1}]),
                content_type="application/json"
            )
        except:
            return HttpResponse(
                json.dumps([{'sended': 0}]),
                content_type="application/json"
            )
    elif request.method == 'POST' and request.POST['msg_body'].__len__() > 4000:
        return HttpResponse(
            json.dumps([{'sended': 2}]),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps([{'sended': 0}]),
            content_type="application/json"
        )

def ajx_auto_request(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if request.method == 'POST':
        alphabet = string.ascii_letters + string.digits
        request_id = ''.join(secrets.choice(alphabet) for i in range(50))
        try:
            user = User.objects.get(email=request.POST['email'])
            return HttpResponse(
                json.dumps([{'requested': 2}]),
                content_type="application/json"
            )
        except User.DoesNotExist:
            try:
                requester = Requester.objects.get(email=request.POST['email'])
                return HttpResponse(
                    json.dumps([{'requested': 2}]),
                    content_type="application/json"
                )

            except Requester.DoesNotExist:
                requester = Requester(
                    program=program,
                    first_name=request.POST['name'],
                    last_name=request.POST['surename'],
                    email = request.POST['email'],
                    phone = request.POST['phone'],
                    dni = request.POST['dni'],
                    gender = request.POST['gender'],
                    theme=request.POST['theme'],
                    request_id=request_id,
                    birthdate=request.POST['birthdate'],
                    planned_end_year = request.POST['planned_end_year'],
                    line=request.POST['line'],
                )
                requester.save()

                request_send_email(request, 'request_received',requester, program)

                return HttpResponse(
                    json.dumps([{'requested': 1}]),
                    content_type="application/json"
                )
    else:
        return HttpResponse(
            json.dumps([{'requested': 0}]),
            content_type="application/json"
        )

def confirm_auto_request(request, program_slug, request_id):
    program = Program.objects.get(slug=program_slug)

    try:
        # Identificar edición activa para solicitante MSc y dip
        try: 
            edition = ProgramEdition.objects.filter(program=program).latest('init_date')
        except ProgramEdition.DoesNotExist:
            messages.error(request, 'No hay una edición activa disponible para este programa')
            return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
        try:
            requester = Requester.objects.get(program=program, request_id=request_id) 
        except Requester.DoesNotExist:
            messages.error(request, "La solicitud no existe o ya fue procesada")
            if program.type == 'phd':
                return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
            elif program.type == 'msc':
                return HttpResponseRedirect(reverse('programs:msc_index', args=[program_slug]))
            elif program.type == 'dip':
                return HttpResponseRedirect(reverse('programs:dip_index', args=[program_slug]))
        try:
            user = User.objects.get(email=requester.email)
            try:
                if program.type == 'phd':
                    Student.objects.get(user=user, program=program)
                elif program.type == 'msc':
                    MscStudent.objects.get(user=user, program=program)
                elif program.type == 'dip':
                    DipStudent.objects.get(user=user, program=program)
                requester.delete()
                messages.error(request, 'Ha habido un error, quizá usted ya haya confirmado su solicitud de ingreso')
                if program.type == 'phd':
                    return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
                elif program.type == 'msc':
                    return HttpResponseRedirect(reverse('programs:msc_index', args=[program_slug]))
                elif program.type == 'dip':
                    return HttpResponseRedirect(reverse('programs:dip_index', args=[program_slug]))
            except (Student.DoesNotExist, MscStudent.DoesNotExist):
                if program.type == 'phd':
                    student = Student(
                        user=user,
                        program=program,
                        gender=requester.gender,
                        dni=requester.dni,
                        birth_date=requester.birthdate,
                    )
                    student.save()
                    formation_plan = StudentFormationPlan(
                        phdstudent=student,
                        elaboration_date=now(),
                        last_update_date=now(),
                        planned_end_year=requester.planned_end_year,
                    )
                    formation_plan.save()
                    new_student = PhdStudent(
                        student=student,
                        status='solicitante',
                        category='',
                        center='',
                    )
                    new_student.save()
                    student_theme = PhdStudentTheme(
                        phd_student=new_student,
                        line=InvestigationLine.objects.get(pk=requester.line),
                        description=requester.theme,
                    )
                    student_theme.save()
                elif program.type == 'msc':
                    student = MscStudent(
                        user=user,
                        program=program,
                        edition=edition,
                        gender=requester.gender,
                        dni=requester.dni,
                        birth_date=requester.birthdate,
                        status='solicitante',
                    )
                    student.save()
                    student_theme = MscStudentTheme(
                        student=student,
                        line=InvestigationLine.objects.get(pk=requester.line),
                        description=requester.theme,
                    )
                    student_theme.save()

                elif program.type == 'dip':
                    student = DipStudent(
                        user=user,
                        program=program,
                        edition=edition,
                        gender=requester.gender,
                        dni=requester.dni,
                        birth_date=requester.birthdate,
                        status='solicitante',
                    )
                    student.save()
                    #student_theme = DipStudentTheme(
                        #student=student,
                        #line=InvestigationLine.objects.get(pk=requester.line),
                        #description=requester.theme,
                    #)
                    #student_theme.save()


                utils_send_email(request, 'wm', program.email, student, '', '', program, '*********')

                for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):
                    if program.type == 'phd':
                        new_student_requirement = StudentFileDocument(
                            student=student,
                            program_file_document=requirement,
                        )
                    elif program.type == 'msc':
                        new_student_requirement = StudentFileDocument(
                            msc_student=student,
                            program_file_document=requirement,
                        )
                    elif program.type == 'dip':
                        new_student_requirement = StudentFileDocument(
                            dip_student=student,
                            program_file_document=requirement,
                        )
                    new_student_requirement.save()

                requester.delete()
                messages.success(request, 'La solicitud se ha confirmado, revise el correo provisto por usted en busca de más orientaciones')
                if program.type == 'phd':
                    return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
                elif program.type == 'msc':
                    return HttpResponseRedirect(reverse('programs:msc_index', args=[program_slug]))
                elif program.type == 'dip':
                    return HttpResponseRedirect(reverse('programs:dip_index', args=[program_slug]))
        except User.DoesNotExist:
            passwd = program_slug + str(random.randint(1000000, 9999999))
            user = User.objects.create_user(
                requester.email,
                requester.email,
                passwd,
            )
            user.first_name = requester.first_name
            user.last_name = requester.last_name
            user.save()

            if program.type == 'phd':
                student = Student(
                    user=user,
                    program=program,
                    gender=requester.gender,
                    dni=requester.dni,
                    birth_date=requester.birthdate,
                )
                student.save()
                formation_plan = StudentFormationPlan(
                    phdstudent=student,
                    elaboration_date=now(),
                    last_update_date=now(),
                    planned_end_year=requester.planned_end_year,
                )
                formation_plan.save()
                new_student = PhdStudent(
                    student=student,
                    status='solicitante',
                    category='',
                    center='',
                )
                new_student.save()
                new_theme = PhdStudentTheme(
                    phd_student=new_student,
                    description=requester.theme,
                    line=InvestigationLine.objects.get(pk=requester.line),
                )
                new_theme.save()
            elif program.type == 'msc':
                student = MscStudent(
                    user=user,
                    program=program,
                    edition=edition,
                    gender=requester.gender,
                    dni=requester.dni,
                    birth_date=requester.birthdate,
                    status='solicitante',
                )
                student.save()
                new_theme = MscStudentTheme(
                    student=student,
                    description=requester.theme,
                    line=InvestigationLine.objects.get(pk=requester.line),
                )
                new_theme.save()

            elif program.type == 'dip':
                student = DipStudent(
                    user=user,
                    program=program,
                    edition=edition,
                    gender=requester.gender,
                    dni=requester.dni,
                    birth_date=requester.birthdate,
                    status='solicitante',
                )
                student.save()

            utils_send_email(request, 'wm', program.email, student, '', '', program, passwd)

            for requirement in ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True):
                if program.type == 'phd':
                    new_student_requirement = StudentFileDocument(
                        student=student,
                        program_file_document=requirement,
                    )
                elif program.type == 'msc':
                    new_student_requirement = StudentFileDocument(
                        msc_student=student,
                        program_file_document=requirement,
                    )
                
                elif program.type == 'dip':
                    new_student_requirement = StudentFileDocument(
                        msc_student=student,
                        program_file_document=requirement,
                    )
                new_student_requirement.save()

            requester.delete()
            messages.success(request, 'La solicitud se ha confirmado, revise el correo provisto por usted en busca de más orientaciones')
            if program.type == 'phd':
                return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
            elif program.type == 'msc':
                return HttpResponseRedirect(reverse('programs:msc_index', args=[program_slug]))
            elif program.type == 'dip':
                return HttpResponseRedirect(reverse('programs:dip_index', args=[program_slug]))
    except Requester.DoesNotExist:
        messages.error(request, 'Ha habido un error, quizá usted ya haya confirmado su solicitud de ingreso')
        if program.type == 'phd':
            return HttpResponseRedirect(reverse('programs:index', args=[program_slug]))
        elif program.type == 'msc':
            return HttpResponseRedirect(reverse('programs:msc_index', args=[program_slug]))
        elif program.type == 'dip':
            return HttpResponseRedirect(reverse('programs:dip_index', args=[program_slug]))
    
    # para msc
    


def error_offlogin(request, program, error_message):
    context = {
        'program': program,
        'error_message':error_message,
    }
    return render(request, 'programs/error_500_offlogin.html', context)

def success_offlogin(request, program, message):
    context = {
        'program': program,
        'message': message,
    }
    return render(request, 'programs/success.html', context)

@login_required
def ajx_everybody_massive_msg(request, program_slug ):
    program=Program.objects.get(slug=program_slug)
    if request.method == 'POST' and request.POST['msg_body'].__len__() <= 1000:
        try:
            email_list = []
            for member in ProgramMember.objects.filter(program=program):
                email_list.append(member.user.email)
                if program.type == 'phd':
                    for student in Student.objects.filter(program=program):
                        email_list.append(student.user.email)
                elif program.type == 'msc':
                    for student in MscStudent.objects.filter(program=program):
                        email_list.append(student.user.email)
                elif program.type == 'dip':
                    for student in DipStudent.objects.filter(program=program):
                        email_list.append(student.user.email)

            if email_list.__len__()<=10:
                send_mail(request.POST['msg_subject'], request.POST['msg_body'],request.user.get_full_name + request.user.email,
                          email_list, fail_silently=False, html_message=request.POST['msg_body'])
            else:
                count = email_list.__len__() // 10
                rest = email_list.__len__() % 10

                for i in range(count):
                    print(i)
                    send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                              request.user.get_full_name + request.user.email,
                              email_list[10 * i:10 * (i + 1)], fail_silently=False, html_message=request.POST['msg_body'])

                    if rest != 0:
                        send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                                  request.user.get_full_name + request.user.email,
                                  email_list[10 * count:10 * count + rest], fail_silently=False,
                                  html_message=request.POST['msg_body'])



            return HttpResponse(
                json.dumps([{'sended': 1}]),
                content_type="application/json"
            )
        except:
            return HttpResponse(
                json.dumps([{'sended': 0}]),
                content_type="application/json"
            )
    elif request.method == 'POST' and request.POST['msg_body'].__len__() > 1000:
        return HttpResponse(
            json.dumps([{'sended': 2}]),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps([{'sended': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_students_massive_msg(request, program_slug ):
    program=Program.objects.get(slug=program_slug)
    if request.method == 'POST' and request.POST['msg_body'].__len__() <= 4000:
        try:
            email_list = []
            for professor in ProgramMember.objects.filter(Q(role='Coordinador')|Q(role='Secretario'), program=program):
                email_list.append(professor.user.email)

            if program.type == 'phd':


                if request.POST['msg_scope'] == 'requesters':
                    for student in Student.objects.filter(program=program ,phdstudent__status='solicitante' ):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            phd_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope']=='aproved':
                    for student in Student.objects.filter(program=program ,phdstudent__status='doctorando' ):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            phd_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope']=='graduated':
                    for student in Student.objects.filter(program=program, phdstudent__status='graduado'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            phd_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope']=='all':
                    for student in Student.objects.filter(program=program):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            phd_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

            elif program.type == 'msc':
                if request.POST['msg_scope'] == 'requesters':
                    for student in MscStudent.objects.filter(program=program, status='solicitante'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            msc_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()
                elif request.POST['msg_scope'] == 'aproved':
                    for student in MscStudent.objects.filter(program=program, status='maestrante'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            msc_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope'] == 'graduated':
                    for student in MscStudent.objects.filter(program=program, status='graduado'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            msc_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope'] == 'all':
                    for student in MscStudent.objects.filter(program=program):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            msc_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()
            elif program.type == 'dip':
                if request.POST['msg_scope'] == 'requesters':
                    for student in DipStudent.objects.filter(program=program, status='solicitante'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            dip_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()
                elif request.POST['msg_scope'] == 'aproved':
                    for student in DipStudent.objects.filter(program=program, status='diplomante'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            dip_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope'] == 'graduated':
                    for student in DipStudent.objects.filter(program=program, status='graduado'):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            dip_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

                elif request.POST['msg_scope'] == 'all':
                    for student in DipStudent.objects.filter(program=program):
                        email_list.append(student.user.email)
                        new_message = Message(
                            sender=request.user,
                            dip_student_receiver=student,
                            subject=request.POST['msg_subject'],
                            body=request.POST['msg_body'],

                        )
                        new_message.save()

            if email_list.__len__()<=20:
                send_mail(request.POST['msg_subject'], request.POST['msg_body'],request.user.email,
                          email_list, fail_silently=False, html_message=request.POST['msg_body'])
            else:
                count = email_list.__len__() // 4
                rest = email_list.__len__() % 4

                for i in range(count):
                    print(i)
                    send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                              request.user.email,
                              email_list[count * i:count * (i + 1)], fail_silently=False, html_message=request.POST['msg_body'])

                    if rest != 0:
                        send_mail(request.POST['msg_subject'], request.POST['msg_body'],
                                  request.user.email,
                                  email_list[4* count:email_list.__len__()], fail_silently=False,
                                  html_message=request.POST['msg_body'])

            sended_message = MessageSended(
                sender=request.user,
                context='students',
                subject=request.POST['msg_subject'],
                body=request.POST['msg_body'],

            )
            sended_message.save()

            return HttpResponse(
                json.dumps([{'sended': 1}]),
                content_type="application/json"
            )
        except:
            print(Exception)
            return HttpResponse(
                json.dumps([{'sended': 0}]),
                content_type="application/json"
            )
    elif request.method == 'POST' and request.POST['msg_body'].__len__() > 4000:
        return HttpResponse(
            json.dumps([{'sended': 2}]),
            content_type="application/json"
        )
    else:
        return HttpResponse(
            json.dumps([{'sended': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_student_personal_msg(request, program_slug ):

    # ==== [CAMBIOS INICIO] Validaciones adicionales ====
    if not request.method == 'POST':
        return JsonResponse({'sended': 0, 'error': 'Método no permitido'}, status=405)
    
    if 'student_id' not in request.POST or 'msg_body' not in request.POST or 'msg_subject' not in request.POST:
        return JsonResponse({'sended': 0, 'error': 'Datos incompletos'}, status=400)
    # ==== [CAMBIOS FIN] ====

    


    program=Program.objects.get(slug=program_slug)
    if program.programmember_set.all().filter(role='Coordinador').count()>0:
        sender_email = program.programmember_set.filter(role='Coordinador')[0].user.email
    elif program.email is not None:
        sender_email = program.email
    else:
        sender_email = NO_REPLY_EMAIL
    print(sender_email)
    
    if request.method == 'POST' and request.POST['msg_body'].__len__() <= 4000:
        try:

            # ==== [CAMBIOS INICIO] Manejo dinámico de modelos ====
            student_model = {
                'phd': Student,
                'msc': MscStudent,
                'dip': DipStudent
            }.get(program.type)
            
            if not student_model:
                return JsonResponse({'sended': 0, 'error': 'Tipo de programa no válido'}, status=400)
                
            student = student_model.objects.get(pk=request.POST['student_id'])
            # ==== [CAMBIOS FIN] ====        

            

            if program.type == 'phd':
                new_message = Message(
                    sender=request.user,
                    phd_student_receiver=student,
                    #phd_student_receiver=Student.objects.get(pk=request.POST['student_id']),
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                new_message.save()
                sended_message = MessageSended(
                    sender= request.user,
                    context= 'personal',
                    phd_student_receiver=Student.objects.get(pk=request.POST['student_id']),
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                sended_message.save()
                send_mail(request.POST['msg_subject'], request.POST['msg_body'],request.user.email,

                          [student.user.email, sender_email],  # Más limpio**  
                          #[Student.objects.get(pk=request.POST['student_id']).user.email,sender_email],
                          fail_silently=False,html_message=request.POST['msg_body'])

            elif program.type == 'msc':
                new_message = Message(
                    sender=request.user,
                    msc_student_receiver=student,  # Usamos la variable ya obtenida**
                    #msc_student_receiver=MscStudent.objects.get(pk=request.POST['student_id']),
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                new_message.save()
                sended_message = MessageSended(
                    sender=request.user,
                    context='personal',
                    msc_student_receiver=MscStudent.objects.get(pk=request.POST['student_id']),
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                sended_message.save()
                send_mail(request.POST['msg_subject'], request.POST['msg_body'], request.user.email,
                          [student.user.email, sender_email],  # Más limpio** 
                          #[MscStudent.objects.get(pk=request.POST['student_id']).user.email, sender_email],
                          fail_silently=False, html_message=request.POST['msg_body'])
            elif program.type == 'dip':
                new_message = Message(
                    sender=request.user,
                    dip_student_receiver=student,  # Usamos la variable ya obtenida**
                    #dip_student_receiver=DipStudent.objects.get(pk=request.POST['student_id']),
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                new_message.save()
                sended_message = MessageSended(
                    sender=request.user,
                    context='personal',
                    dip_student_receiver=DipStudent.objects.get(pk=request.POST['student_id']),
                    subject=request.POST['msg_subject'],
                    body=request.POST['msg_body'],

                )
                sended_message.save()
                send_mail(request.POST['msg_subject'], request.POST['msg_body'], request.user.email,
                          [student.user.email, sender_email],  # Más limpio** 
                          #[DipStudent.objects.get(pk=request.POST['student_id']).user.email, sender_email],
                          fail_silently=False, html_message=request.POST['msg_body'])

            return HttpResponse(
                json.dumps([{'sended': 1}]),
                content_type="application/json"
            )
        except student_model.DoesNotExist:
            return JsonResponse({'sended': 0, 'error': 'Estudiante no encontrado'}, status=404)
        
        except Exception as e:
            print(f"Error inesperado: {str(e)}")  # Mejor logging**    
            return HttpResponse(
                json.dumps([{'sended': 0}]),
                content_type="application/json"
            )
    elif request.method == 'POST' and request.POST['msg_body'].__len__() > 4000:
        
        return JsonResponse({'sended': 2, 'error': 'Mensaje demasiado largo'})  # Más informativo**
        #return HttpResponse(
         #   json.dumps([{'sended': 2}]),
          #  content_type="application/json"
        #)
    else:
        return HttpResponse(
            json.dumps([{'sended': 0}]),
            content_type="application/json"
        )

@login_required
def ajx_members_by_age(request, program_slug):
    response_data=[]
    labels = ['<30 años','30-40','40-50','>50 años']
    data = []

    # locale.setlocale(locale.LC_ALL, 'es-ES')
    i=0
    data.append(ProgramMember.objects.filter(program=Program.objects.get(slug=program_slug ),birth_date__year__gt=now().year-30 ).__len__())
    data.append(ProgramMember.objects.filter(program=Program.objects.get(slug=program_slug ),birth_date__year__gt=now().year-40, birth_date__year__lt=now().year-30 ).__len__() )
    data.append(ProgramMember.objects.filter(program=Program.objects.get(slug=program_slug ),birth_date__year__gt=now().year-50, birth_date__year__lt=now().year-40 ).__len__() )
    data.append(ProgramMember.objects.filter(program=Program.objects.get(slug=program_slug ), birth_date__year__lt=now().year-50 ).__len__() )

    response_data.append(labels)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_students_by_state(request, program_slug):
    response_data=[]
    data = []
    program=Program.objects.get(slug=program_slug)

    if program.type == 'phd':
        labels = ['Graduados', 'Solicitantes', 'Doctorandos']
        data.append(PhdStudent.objects.filter(student__program=program, status='graduado').__len__())
        data.append(PhdStudent.objects.filter(student__program=program, status='solicitante').__len__())
        data.append(PhdStudent.objects.filter(student__program=program, status='doctorando').__len__())

        response_data.append(labels)
        response_data.append(data)
    elif program.type == 'msc':
        labels = ['Graduados', 'Solicitantes', 'Maestrantes']
        data.append(MscStudent.objects.filter(program=program, status='graduado').__len__())
        data.append(MscStudent.objects.filter(program=program, status='solicitante').__len__())
        data.append(MscStudent.objects.filter(program=program, status='maestrante').__len__())

        response_data.append(labels)
        response_data.append(data)
    elif program.type == 'dip':
        labels = ['Graduados', 'Solicitantes', 'Diplomantes']
        data.append(DipStudent.objects.filter(program=program, status='graduado').__len__())
        data.append(DipStudent.objects.filter(program=program, status='solicitante').__len__())
        data.append(DipStudent.objects.filter(program=program, status='diplomante').__len__())

        response_data.append(labels)
        response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )

@login_required
def ajx_students_by_line_donut(request, program_slug):
    response_data=[]
    data = []
    program=Program.objects.get(slug=program_slug)

    if program.type == 'phd':
        labels = []
        for line in InvestigationLine.objects.filter(program=program):
            labels.append(line.name)
            data.append(PhdStudent.objects.filter(student__program=program, phdstudenttheme__line=line).__len__())

        response_data.append(labels)
        response_data.append(data)
    elif program.type == 'msc':
        labels = ['Graduados', 'Solicitantes', 'Maestrantes']
        data.append(MscStudent.objects.filter(program=program, status='graduado').__len__())
        data.append(MscStudent.objects.filter(program=program, status='solicitante').__len__())
        data.append(MscStudent.objects.filter(program=program, status='maestrante').__len__())

        response_data.append(labels)
        response_data.append(data)
    elif program.type == 'dip':
        labels = ['Graduados', 'Solicitantes', 'Diplomantes']
        data.append(DipStudent.objects.filter(program=program, status='graduado').__len__())
        data.append(DipStudent.objects.filter(program=program, status='solicitante').__len__())
        data.append(DipStudent.objects.filter(program=program, status='diplomante').__len__())

        response_data.append(labels)
        response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )



@login_required
def ajx_members_by_grade(request, program_slug):
    response_data=[]
    grades=[]
    data=[]

    program = Program.objects.get(slug=program_slug)
    for member in ProgramMember.objects.filter(program=program):
        if not member.degree in grades:
            grades.append(member.degree)

    for degree in grades:
        data.append(ProgramMember.objects.filter(program=program, degree=degree).__len__())

    response_data.append(grades)
    response_data.append(data)

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_students_by_sex(request, program_slug):

    response_data=[]
    program=Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        response_data.append(Student.objects.filter(program=program, gender='f').__len__())
        response_data.append(Student.objects.filter(program=program, gender='m').__len__())
    elif program.type == 'msc':
        response_data.append(MscStudent.objects.filter(program=program, gender='f').__len__())
        response_data.append(MscStudent.objects.filter(program=program,gender='m').__len__())
    elif program.type == 'dip':
        response_data.append(DipStudent.objects.filter(program=program, gender='f').__len__())
        response_data.append(DipStudent.objects.filter(program=program, gender='m').__len__())

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


@login_required
def ajx_students_by_category(request, program_slug):

    response_data=[]
    program=Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        response_data.append(PhdStudent.objects.filter(student__program=program, category='interno').__len__())
        response_data.append(PhdStudent.objects.filter(student__program=program, category='externo').__len__())
    elif program.type == 'msc':
        response_data.append(MscStudent.objects.filter(program=program, category='interno').__len__())
        response_data.append(MscStudent.objects.filter(program=program, category='externo').__len__())
    elif program.type == 'dip':
        response_data.append(DipStudent.objects.filter(program=program, category='interno').__len__())
        response_data.append(DipStudent.objects.filter(program=program, category='externo').__len__())

    return HttpResponse(
        json.dumps(response_data),
        content_type="application/json"
    )


def program_member_picture(request, program_slug, member_id):
    fs = FileSystemStorage()
    # filename = Papers.objects.get(pk=paper_id).file_url +  str(Papers.objects.get(pk=paper_id).file)
    filename = ProgramMember.objects.get(pk=member_id).picture.url
    if fs.exists(filename):
        with fs.open(filename) as img:
            response = HttpResponse(img, content_type='image/jpeg')
            return response
    else:
        return HttpResponse('Error')

def program_new_picture(request, program_slug, new_id):
    fs = FileSystemStorage()
    # filename = Papers.objects.get(pk=paper_id).file_url +  str(Papers.objects.get(pk=paper_id).file)
    filename = New.objects.get(pk=new_id).img.url
    if fs.exists(filename):
        with fs.open(filename) as img:
            response = HttpResponse(img, content_type='image/jpeg')
            return response
    else:
        return HttpResponse('Error')

def program_student_picture(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    fs = FileSystemStorage()
    # filename = Papers.objects.get(pk=paper_id).file_url +  str(Papers.objects.get(pk=paper_id).file)
    if program.type == 'phd':
        filename = Student.objects.get(pk=student_id).picture.url
    elif program.type == 'msc':
        filename = MscStudent.objects.get(pk=student_id).picture.url
    elif program.type == 'dip':
        filename = DipStudent.objects.get(pk=student_id).picture.url

    if fs.exists(filename):
        with fs.open(filename) as img:
            response = HttpResponse(img, content_type='image/jpeg')
            return response
    else:
        return HttpResponse('Error')

def program_background(request, program_slug, background_id):
    fs = FileSystemStorage()
    # filename = Papers.objects.get(pk=paper_id).file_url +  str(Papers.objects.get(pk=paper_id).file)
    filename = ProgramBackgrounds.objects.get(pk=background_id).background.path
    print(filename)
    if fs.exists(filename):
        return FileResponse(open(filename, 'rb'))

    else:
        return HttpResponse('Error')

@login_required
def program_statistics(request, program_slug):
    program=Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        context={
            'program':program,
            'member': ProgramMember.objects.get(user=request.user, program=program),
        }
        return render(request, 'programs/statistics.html', context)

@login_required
def create_program_edition(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user,program):
        if program.type != 'phd':
            if request.method == 'POST':
                new_edition= ProgramEdition(
                    program=program,
                    init_date=request.POST['init_date'],
                    end_date=request.POST['end_date'],
                    observations=request.POST['observations'],
                    order=request.POST['order'],

                )
                try:
                    new_edition.order = ProgramEdition.objects.filter(program=program).__len__()+1
                except:
                    pass
                new_edition.save()
                return HttpResponseRedirect(reverse('programs:editions_list', args=[program_slug]))
            else:
                context={
                    'program': program,
                    'member':ProgramMember.objects.get(user=request.user, program=program),
                }
                return render(request, 'programs/create_edition.html', context)
        else:
            return  error_500(request,program,'Los programas doctorales de nuevo tipo no tienen ediciones')
    else:
        return error_500(request, program, 'Usted no tiene privilegios para agregar ediciones')


@login_required
def editions_list(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    context = {
        'program': program,
        'editions': ProgramEdition.objects.filter(program=program)
    }
    try:
        context['member']= ProgramMember.objects.get(user=request.user, program = program)
    except ProgramMember.DoesNotExist:
        pass

    if program.type == 'msc':
        try:
            context['student']=MscStudent.objects.get(user=request.user)
        except MscStudent.DoesNotExist:
            pass
    elif program.type == 'dip':
        try:
            context['student']=DipStudent.objects.get(user=request.user)
        except DipStudent.DoesNotExist:
            pass


    return render(request, 'programs/editions_list.html', context)

@login_required
def edit_program_edition(request, program_slug, edition_id):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if program.type == 'msc' or program.type == 'dip':
            if request.method == 'POST':
                ProgramEdition.objects.filter(pk=edition_id).update(
                    init_date = request.POST['init_date'],
                    end_date = request.POST['end_date'],
                    observations= request.POST['observations'],
                    order= request.POST['order'],

                )
                return HttpResponseRedirect(reverse('programs:editions_list', args=[program_slug]))
            else:
                context={
                    'program':program,
                    'edition': ProgramEdition.objects.get(pk=edition_id)
                }
                return render(request, 'programs/edit_edition.html', context)
        else:
            return error_500(request,program, 'Este tipo de programas no tiene ediciones.')
    else:
        return error_500(request, program, 'Usted no tiene privilegios para modificar ediciones en este programa.')



@login_required
def create_program_doc(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            try:

                new_doc = ProgramDocument(
                    program=program,
                    doc=request.FILES['doc'],
                    type=request.POST['type'],
                    year=request.POST['year'],
                    month=request.POST['month'],
                    description=request.POST['description']
                )
                try:
                    if request.POST['is_public']== 'on':
                        new_doc.is_public = True
                except:
                    pass

                new_doc.save()
                return HttpResponseRedirect(reverse('programs:program_docs_by_year', args=[program_slug, new_doc.year]))

            except:

                return error_500(request,program, 'Ha ocurrido un error al crear la nueva acta')


        else:
            meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 7: "Julio",
                     8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}
            context = {
                'program':program,
                'months': ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio','Agosto', 'Septiembre',
                           'Octubre', 'Noviembre', 'Diciembre'],
                'current_month': meses[now().month],
                'years':range(now().year-10,now().year+1),
                'current_year':now().year,
                'member': ProgramMember.objects.get(user=request.user, program = program)
            }
            return render(request, 'programs/create_program_doc.html',context)
    else:
        return error_500(request,'Usted no tiene privilegios para agregar actas.')

@login_required
def program_documents(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_member(request.user, program):
        years=[]
        for document in ProgramDocument.objects.filter(program=program):
            if not document.year in years:
                years.append(document.year)

        context={
            'program':program,
            'years': sorted(years),
            'documents': ProgramDocument.objects.filter(program=program),
            'member':ProgramMember.objects.get(user=request.user, program=program),
        }
        return render(request, 'programs/program_documents_list.html',context)

    elif user_is_program_student(request.user, program):
        years = []
        context = {
            'program': program,
            'years': sorted(years),
            'documents': ProgramDocument.objects.filter(program=program, is_public=True),
            'student': Student.objects.get(program=program, user=request.user)
        }
        return render(request, 'programs/program_documents_list.html',context)
    else:
        return error_500(request,program,'Usted no puede ver los documentos de este programa')

@login_required
def autoedit_member_profile(request, program_slug, member_id):
    program = Program.objects.get(slug=program_slug)
    member = ProgramMember.objects.get(pk=member_id)
    if request.user == member.user:
        if request.method == 'POST':
            if member.user.email == request.POST['prof_email']:
                request.user.first_name = request.POST['prof_name']
                request.user.last_name = request.POST['prof_lastname']
                member.birth_date = request.POST['prof_bdate']
                member.degree = request.POST['prof_degree']
                member.institution = request.POST['prof_inst']
                member.sex = request.POST['prof_gender']
                member.phone = request.POST['prof_phone']
                try:
                    member.picture = request.FILES['prof_picture']
                except:
                    pass
                request.user.save()
                member.save()

                return HttpResponseRedirect(reverse('programs:view_program_member_profile', args=[program_slug, member_id]))
            else:
                if not ProgramMember.objects.filter(user__email=request.POST['prof_email']):
                    request.user.first_name = request.POST['prof_name']
                    request.user.last_name = request.POST['prof_lastname']
                    request.user.email = request.POST['prof_email']
                    request.user.username = request.POST['prof_email']
                    member.birth_date = request.POST['prof_bdate']
                    member.degree = request.POST['prof_degree']
                    member.institution = request.POST['prof_inst']
                    member.sex = request.POST['prof_gender']
                    member.phone = request.POST['prof_phone']
                    try:
                        member.picture = request.FILES['prof_picture']
                    except:
                        pass
                    request.user.save()
                    member.save()

                    return HttpResponseRedirect(
                        reverse('programs:view_program_member_profile', args=[program_slug, member_id]))

                else:
                    return error_500(request, program, 'El correo introducido esta en uso por otro usuario')
        else:
            return error_500(request, program, 'Usted no tiene acceso a esta funcionalidad por GET')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta funcionalidad')


@login_required
def autoedit_student_profile(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    student =''
    if program.type == 'phd':
        student = Student.objects.get(pk=student_id)
    elif program.type == 'msc':
        student = MscStudent.objects.get(pk=student_id)
    elif program.type == 'dip':
        student = DipStudent.objects.get(pk=student_id)

    if request.user == student.user:
        if request.method == 'POST':
            if student.user.email == request.POST['student_email']:
                request.user.first_name = request.POST['student_name']
                request.user.last_name = request.POST['student_lastname']
                student.birth_date = request.POST['student_bdate']
                student.gender = request.POST['student_gender']
                student.phone = request.POST['student_phone']
                if program.type == 'phd':
                    theme = student.phdstudent.phdstudenttheme
                    theme.description = request.POST['student_theme']
                    theme.save()

                try:
                    student.picture = request.FILES['student_picture']
                except:
                    pass
                request.user.save()
                student.save()

                return HttpResponseRedirect(reverse('programs:view_student_profile', args=[program_slug, student_id]))
            else:
                if not User.objects.filter(email=request.POST['student_email']):
                    request.user.first_name = request.POST['student_name']
                    request.user.last_name = request.POST['student_lastname']
                    request.user.email = request.POST['student_email']
                    request.user.username = request.POST['student_email']
                    student.birth_date = request.POST['student_bdate']
                    student.sex = request.POST['student_gender']
                    student.phone = request.POST['student_phone']
                    try:
                        student.picture = request.FILES['student_picture']
                    except:
                        pass
                    request.user.save()
                    student.save()

                    return HttpResponseRedirect(
                        reverse('programs:view_student_profile', args=[program_slug, student_id]))

                else:
                    return error_500(request, program, 'El correo introducido esta en uso por otro usuario')
        else:
            return error_500(request, program, 'Usted no tiene acceso a esta funcionalidad por GET')

    else:
        return error_500(request, program, 'Usted no tiene acceso a esta funcionalidad')


@login_required
def program_cgc_documents(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    years = []
    if user_is_program_member(request.user, program):

        for document in CGCDocument.objects.filter(is_public=True):
            if not document.year in years:
                years.append(document.year)


        context={
            'program':program,
            'years': sorted(years),
            'documents': CGCDocument.objects.filter(is_public=True),
            'member': ProgramMember.objects.filter(user=request.user).first(),
        }
        return render(request, 'programs/program_cgc_documents_list.html',context)
    elif user_is_program_student(request.user, program):

        for document in CGCDocument.objects.filter(is_public=True):
            if not document.year in years:
                years.append(document.year)


        context={
            'program':program,
            'years': sorted(years),
            'documents': CGCDocument.objects.filter(is_public=True),
            'student': Student.objects.get(user=request.user, program=program),
        }
        return render(request, 'programs/program_cgc_documents_list.html',context)
    else:
        return error_500(request,program, 'Usted no puede ver documentos de la CGC')


@login_required
def program_docs_by_year(request, program_slug, year):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_member(request.user, program):
        years=[]
        for doc in ProgramDocument.objects.filter(program=program):
            if not doc.year in years:
                years.append(doc.year)

        context={
            'year': year,
            'program':program,
            'years': sorted(years),
            'documents': ProgramDocument.objects.filter(program=program, year=year),
        }
        return render(request, 'programs/program_documents_list.html', context)
    else:
        return error_500(request,'Usted no puede ver las actas de este programa')


@login_required
def edit_program_doc(request,program_slug, doc_id):
    program = Program.objects.get(slug=program_slug)
    doc = ProgramDocument.objects.get(pk=doc_id)

    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            try:
                old_year=doc.year
                old_month=doc.month
                old_type = doc.type

                index = ProgramDocument.objects.filter(program=program, month=request.POST['month'],
                                                       year=request.POST['year']).__len__()

                if str(old_year) != request.POST['year'] or old_month != request.POST['month'] or old_type != request.POST['type'] :
                    initial_path = ProgramDocument.objects.get(pk=doc_id).doc.path



                    doc_ext = initial_path.split('.')[initial_path.split('.').__len__() - 1]




                    doc_name = '{0}-{1}-{2}-{3}-{4}.{5}'.format(request.POST['type'].capitalize(), program_slug, request.POST['month'] , request.POST['year'] ,str(index+1), doc_ext)
                    doc.doc.name = 'program_{0}/documents/{1}/{2}/{3}'.format(program_slug, request.POST['year'], request.POST['month'],doc_name)
                    new_path= MEDIA_ROOT+ '/program_{0}/documents/{1}/{2}/{3}'.format(program_slug, request.POST['year'], request.POST['month'], doc_name)



                    os.renames(initial_path, new_path)
                    doc.year = request.POST['year']
                    doc.month = request.POST['month']
                    doc.type = request.POST['type']
                    doc.description = request.POST['description']

                    try:
                        if request.POST['is_public'] == 'on':
                            doc.is_public = True
                    except:
                        doc.is_public = False

                    doc.save()
                else:
                    doc.description = request.POST['description']
                    try:
                        if request.POST['is_public'] == 'on':
                            doc.is_public = True
                    except:
                        doc.is_public = False

                    doc.save()

                try:
                    doc_file = request.FILES['doc']
                    fs = FileSystemStorage()
                    doc_ext = doc_file.name.split('.')[doc_file.name.split('.').__len__() - 1]

                    doc_name = '{0}-{1}-{2}-{3}-{4}.{5}'.format(request.POST['type'].capitalize(), program_slug, request.POST['month'],request.POST['year'], str(index + 1), doc_ext)
                    new_doc_name = 'program_{0}/documents/{1}/{2}/{3}'.format(program_slug, request.POST['year'], request.POST['month'],doc_name)

                    doc.doc.delete()

                    filename = fs.save(new_doc_name, doc_file)

                    doc.doc = filename
                    doc.save()

                except Exception:
                    print('Exception: No se puede actualizar el archivo')

                return HttpResponseRedirect(reverse('programs:program_docs_by_year', args=[program_slug, ProgramDocument.objects.get(pk=doc_id).year]))

            except:
                return error_500(request,program, 'Ha ocurrido un error al editar el documento')


        else:
            meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 7: "Julio",
                     8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}
            context = {
                'program':program,
                'member':ProgramMember.objects.get(user=request.user, program=program),
                'months': ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio','Agosto', 'Septiembre',
                           'Octubre', 'Noviembre', 'Diciembre'],
                'current_month': meses[now().month],
                'years':range(now().year-10,now().year+1),
                'current_year':now().year,
                'doc': ProgramDocument.objects.get(pk=doc_id),
            }
            return render(request, 'programs/edit_program_doc.html', context)
    else:
        return error_500(request,'Usted no tiene privilegios para agregar actas.')

@login_required
def program_doc_view(request,program_slug, doc_id):
    program = Program.objects.get(slug=program_slug)
    document = ProgramDocument.objects.get(pk=doc_id)

    if user_is_program_member(request.user, program) or (user_is_program_student(request.user, program) and document.is_public):


        fs = FileSystemStorage()

        filename = document.doc.url

        if fs.exists(filename):
            doc_ext = filename.split('.')[filename.split('.').__len__() - 1]

            if doc_ext == 'doc' or doc_ext == 'docx' or doc_ext == 'odt':

                with fs.open(filename) as doc:
                    response = HttpResponse(doc, content_type='application/doc')
                    response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                        filename.split('/').__len__() - 1] + '"'

                    return response

            elif doc_ext == 'pdf':
                with fs.open(filename) as doc:
                    response = HttpResponse(doc, content_type='application/pdf')
                    response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                        filename.split('/').__len__() - 1] + '"'

                    return response

        else:

            return error_500(request, program, 'No existe el archivo solicitado')
    else:
        return error_500(request,program, 'Solo los miembros del claustro pueden ver todos los documentos del mismo. Los estudiantes solo pueden ver los documentos públicos')

def public_program_doc_view(request, program_slug,doc_id):
    program = Program.objects.get(slug=program_slug)

    document = ProgramDocument.objects.get(pk=doc_id)

    fs = FileSystemStorage()

    filename = document.doc.url

    if fs.exists(filename):
        doc_ext = filename.split('.')[filename.split('.').__len__() - 1]

        if doc_ext == 'doc' or doc_ext == 'docx' or doc_ext == 'odt':

            with fs.open(filename) as doc:
                response = HttpResponse(doc, content_type='application/doc')
                response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                    filename.split('/').__len__() - 1] + '"'

                return response

        elif doc_ext == 'pdf':
            with fs.open(filename) as doc:
                response = HttpResponse(doc, content_type='application/pdf')
                response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                    filename.split('/').__len__() - 1] + '"'

                return response

    else:

        return error_500(request, program, 'No existe el archivo solicitado')

@login_required
def student_filedoc_view(request, program_slug,student_id, doc_id):
    program = Program.objects.get(slug=program_slug)


    if program.type == 'phd':
        student = Student.objects.get(pk=student_id)
        document = StudentFileDocument.objects.get(student=student, program_file_document=ProgramFileDoc.objects.get(pk=doc_id))

    elif program.type == 'msc':
        student = MscStudent.objects.get(pk=student_id)
        document = StudentFileDocument.objects.get(msc_student=student,
                                                   program_file_document=ProgramFileDoc.objects.get(pk=doc_id))
    elif program.type == 'dip':
        student = DipStudent.objects.get(pk=student_id)
        document = StudentFileDocument.objects.get(dip_student=student,
                                                   program_file_document=ProgramFileDoc.objects.get(pk=doc_id))


    if request.user == student.user or user_is_program_cs(request.user, program):
        fs = FileSystemStorage()

        filename = document.file.url

        if fs.exists(filename):
            doc_ext = (filename.split('.')[filename.split('.').__len__() - 1]).lower()

            if doc_ext == 'doc' or doc_ext == 'docx' or doc_ext == 'odt':

                with fs.open(filename) as doc:
                    response = HttpResponse(doc, content_type='application/doc')
                    response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                        filename.split('/').__len__() - 1] + '"'

                    return response

            elif doc_ext == 'pdf':
                with fs.open(filename) as doc:
                    response = HttpResponse(doc, content_type='application/pdf')
                    response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                        filename.split('/').__len__() - 1] + '"'

                    return response
            else:
                with fs.open(filename) as doc:
                    response = HttpResponse(doc, content_type='application')
                    response['Content-Disposition'] = "inline; filename=" + '"' + filename.split('/')[
                        filename.split('/').__len__() - 1] + '"'

                    return response


        else:

            return error_500(request, program, 'No existe el archivo solicitado')
    else:
        return error_500(request, program, 'No tiene privilegios para ver este archivo.')


def program_cgc_document_view(request,program_slug, document_id):
    program = Program.objects.get(slug=program_slug)

    brieffing = CGCDocument.objects.get(pk=document_id)

    fs = FileSystemStorage()

    filename =brieffing.doc.url

    if fs.exists(filename):
        brief_ext =filename.split('.')[filename.split('.').__len__()-1]


        if brief_ext =='doc' or brief_ext=='docx' or brief_ext == 'odt':

            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/doc')
                response['Content-Disposition'] =  "inline; filename=" + '"'+filename.split('/')[filename.split('/').__len__()-1]+'"'

                return response

        elif brief_ext == 'pdf' :
            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/pdf')
                response['Content-Disposition'] = "inline; filename=" + '"'+filename.split('/')[filename.split('/').__len__()-1] + '"'

                return response

    else:


        return error_500(request,program, 'No existe el archivo solicitado')

def program_cgc_brief_view(request,program_slug, brief_id):
    program = Program.objects.get(slug=program_slug)

    brieffing = CGCBrief.objects.get(pk=brief_id)

    fs = FileSystemStorage()

    filename =brieffing.brief.url

    if fs.exists(filename):
        brief_ext =filename.split('.')[filename.split('.').__len__()-1]


        if brief_ext =='doc' or brief_ext=='docx' or brief_ext == 'odt':

            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/doc')
                response['Content-Disposition'] =  "inline; filename=" + '"'+filename.split('/')[filename.split('/').__len__()-1]+'"'

                return response

        elif brief_ext == 'pdf' :
            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/pdf')
                response['Content-Disposition'] = "inline; filename=" + '"'+filename.split('/')[filename.split('/').__len__()-1] + '"'

                return response

    else:


        return error_500(request,program, 'No existe el archivo solicitado')

def program_cngc_brief_view(request,program_slug, brief_id):
    program = Program.objects.get(slug=program_slug)

    brieffing = CNGCBrief.objects.get(pk=brief_id)

    fs = FileSystemStorage()

    filename =brieffing.brief.url

    if fs.exists(filename):
        brief_ext =filename.split('.')[filename.split('.').__len__()-1]


        if brief_ext =='doc' or brief_ext=='docx' or brief_ext == 'odt':

            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/doc')
                response['Content-Disposition'] =  "inline; filename=" + '"'+filename.split('/')[filename.split('/').__len__()-1]+'"'

                return response

        elif brief_ext == 'pdf' :
            with fs.open(filename) as brief:
                response = HttpResponse(brief, content_type='application/pdf')
                response['Content-Disposition'] = "inline; filename=" + '"'+filename.split('/')[filename.split('/').__len__()-1] + '"'

                return response

    else:


        return error_500(request,program, 'No existe el archivo solicitado')



@login_required
def ajx_delete_program_document(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program ):
        if request.method=='POST':
            doc_id=request.POST['doc_id']
            try:

                ProgramDocument.objects.get(pk=doc_id).doc.delete()
                ProgramDocument.objects.get(pk=doc_id).delete()


                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 0}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 2}]),
            content_type="application/json"
        )

@login_required()
def program_download_docs(request,program_slug):
    program = Program.objects.get(slug=program_slug)

    brief_list = ProgramDocument.objects.filter(program=program)
    zipname = "Documentos-"+slugify(program.short_name)

    if brief_list.count() > 0:
        # Files (local path) to put in the .zip
        filenames = []
        i = int(0)
        # Folder name in ZIP archive which contains the above files
        # E.g [thearchive.zip]/somefiles/file2.txt

        zip_filename = "%s.zip" % zipname

        # # Open StringIO to grab in-memory ZIP contents
        # s = StringIO()
        #
        # # The zip compressor
        zf = zipfile.ZipFile(MEDIA_ROOT + '/' + zip_filename, "w")

        for brief in brief_list:


            try:
                fpath = MEDIA_ROOT +'/'+ brief.doc.name
                fdir, fname = os.path.split(fpath)
                zip_subdir = str(brief.year)
                zip_path = os.path.join(zip_subdir, brief.doc.name.split('/')[brief.doc.name.split('/').__len__()-1])


                # zip_path = os.path.join(zip_subdir, fname)
                # Add file, at correct path
                zf.write(fpath, zip_path)
            except:
                print('Excepcion:algo paso al agregar el archivo', brief.brief.name)


        zf.close()

        # Must close zip for all contents to be written
        fs = FileSystemStorage()

        with fs.open(zip_filename) as zip:
            # response = HttpResponse(pdf, content_type='application/pdf')
            # response['Content-Disposition'] = "inline; filename=" + '"' + filename + '"'
            # Grab ZIP file from in-memory, make response with correct MIME-type
            resp = HttpResponse(zip, content_type="application/x-zip-compressed")
            # ..and correct content-disposition
            resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename

        return resp

    else:
        return error_500(request, program, 'No hay actas para descargar')


@login_required()
def program_by_year_doc_download(request,program_slug, year):
    program = Program.objects.get(slug=program_slug)

    brief_list = ProgramDocument.objects.filter(program=program, year=year)
    zipname = "Documentos-de-"+slugify(program.short_name)+'-'+str(year)

    if brief_list.count() > 0:
        # Files (local path) to put in the .zip
        filenames = []
        i = int(0)
        # Folder name in ZIP archive which contains the above files
        # E.g [thearchive.zip]/somefiles/file2.txt

        zip_filename = "%s.zip" % zipname

        # # Open StringIO to grab in-memory ZIP contents
        # s = StringIO()
        #
        # # The zip compressor
        zf = zipfile.ZipFile(MEDIA_ROOT + '/' + zip_filename, "w")

        for brief in brief_list:
            try:
                fpath = MEDIA_ROOT +'/'+ brief.doc.name
                fdir, fname = os.path.split(fpath)
                zip_subdir = str(brief.month)
                zip_path = os.path.join(zip_subdir, brief.doc.name.split('/')[brief.doc.name.split('/').__len__()-1])


                # zip_path = os.path.join(zip_subdir, fname)
                # Add file, at correct path
                zf.write(fpath, zip_path)
            except:
                print('Excepcion:algo paso al agregar el archivo', brief.doc.name)


        zf.close()

        # Must close zip for all contents to be written
        fs = FileSystemStorage()

        with fs.open(zip_filename) as zip:
            # response = HttpResponse(pdf, content_type='application/pdf')
            # response['Content-Disposition'] = "inline; filename=" + '"' + filename + '"'
            # Grab ZIP file from in-memory, make response with correct MIME-type
            resp = HttpResponse(zip, content_type="application/x-zip-compressed")
            # ..and correct content-disposition
            resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename

        return resp

    else:
        return error_500(request, program, 'No hay actas para descargar')

@login_required
def docx_program_report(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        document = Document()

        document.add_heading('Resumen de datos del Programa', level=1)

        document.add_heading(program.full_name, level=2)
        if program.type == 'phd':
            document.add_heading('Estudiantes por estado', level=3)
            if PhdStudent.objects.filter(student__program=program):
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Solicitantes'
                hdr_cells[1].text = 'Doctorandos'
                hdr_cells[2].text = 'Graduados'
                hdr_cells[3].text = 'Total'

                row_cells = table.add_row().cells
                row_cells[0].text = str(PhdStudent.objects.filter(student__program=program, status='Solicitante').__len__())
                row_cells[1].text = str(PhdStudent.objects.filter(student__program=program, status='Doctorando').__len__())
                row_cells[2].text = str(PhdStudent.objects.filter(student__program=program, status='Graduado').__len__())
                row_cells[3].text = str(PhdStudent.objects.filter(student__program=program).__len__())
            else:
                document.add_heading('No hay estudiantes registrados en el programa', level=5)

            document.add_heading('Solicitantes de ingreso', level=3)
            if PhdStudent.objects.filter(student__program=program, status='Solicitante'):
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de solicitud'
                hdr_cells[2].text = 'Requisitos de ingreso'
                hdr_cells[3].text = 'Tema'

                for student in PhdStudent.objects.filter(student__program=program, status='Solicitante'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.student.user.get_full_name())
                    row_cells[1].text = str(student.student.request_date)
                    if init_requirements_accomplished(student.student, program):
                        row_cells[2].text = 'Cumplidos'
                    else:
                        row_cells[2].text = 'Incumplidos'
                    row_cells[3].text = str(student.phdstudenttheme)


            else:
                document.add_heading('No hay solicitantes registrados en el programa', level=5)

            document.add_heading('Doctorandos', level=3)
            if PhdStudent.objects.filter(student__program=program, status='Doctorando'):
                table = document.add_table(rows=1, cols=5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de ingreso'
                hdr_cells[2].text = 'Año de defensa'
                hdr_cells[3].text = 'Requisitos de egreso'
                hdr_cells[4].text = 'Tema'

                for student in PhdStudent.objects.filter(student__program=program, status='Doctorando').order_by('student__studentformationplan__planned_end_year'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.student.user.get_full_name())
                    row_cells[1].text = str(student.student.init_date)
                    try:
                        row_cells[2].text = str(student.student.studentformationplan.planned_end_year)
                    except:
                        row_cells[2].text = "No declarada"

                    if finish_requirements_accomplished(student.student, program):
                        row_cells[3].text = 'Cumplidos'
                    else:
                        row_cells[3].text = 'Incumplidos'
                    row_cells[4].text = str(student.phdstudenttheme)
            else:
                document.add_heading('No hay doctorandos registrados en el programa', level=5)

            document.add_heading('Graduados', level=3)
            if PhdStudent.objects.filter(student__program=program, status='Graduado'):
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de ingreso'
                hdr_cells[2].text = 'Fecha de egreso'
                hdr_cells[3].text = 'Tema'

                for student in PhdStudent.objects.filter(student__program=program, status='Graduado').order_by('student__graduate_date'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.student.user.get_full_name())
                    row_cells[1].text = str(student.student.init_date)
                    row_cells[2].text = str(student.student.graduate_date)
                    row_cells[3].text = str(student.phdstudenttheme)

            else:
                document.add_heading('No hay graduados registrados en el programa', level=5)

        elif program.type == 'msc':
            if MscStudent.objects.filter(program=program):
                document.add_heading('Estudiantes por estado', level=3)
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Solicitantes'
                hdr_cells[1].text = 'Maestrantes'
                hdr_cells[2].text = 'Graduados'
                hdr_cells[3].text = 'Total'

                row_cells = table.add_row().cells
                row_cells[0].text = str(
                    MscStudent.objects.filter(program=program, status='Solicitante').__len__())
                row_cells[1].text = str(
                    MscStudent.objects.filter(program=program, status='Maestrante').__len__())
                row_cells[2].text = str(
                    MscStudent.objects.filter(program=program, status='Graduado').__len__())
                row_cells[3].text = str(
                    MscStudent.objects.filter(program=program).__len__())
            else:
                document.add_heading('No hay estudiantes registrados en el programa', level=5)

            document.add_heading('Solicitantes de ingreso', level=3)
            if MscStudent.objects.filter(program=program, status='Solicitante'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y Apellidos'
                hdr_cells[1].text = 'Fecha de solicitud'
                hdr_cells[2].text = 'Requisitos de ingreso'

                for student in MscStudent.objects.filter(program=program, status='Solicitante'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.user.get_full_name())
                    row_cells[1].text = str(student.request_date)
                    if init_requirements_accomplished(student, program):
                        row_cells[2].text = 'Cumplidos'
                    else:
                        row_cells[2].text = 'Incumplidos'
            else:
                document.add_heading('No hay solicitantes registrados en el programa', level=5)

            document.add_heading('Graduados', level=3)
            if MscStudent.objects.filter(program=program, status='Graduado'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de ingreso'
                hdr_cells[2].text = 'Fecha de egreso'

                for student in MscStudent.objects.filter(program=program, status='Graduado'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.user.get_full_name())
                    row_cells[1].text = str(student.init_date)
                    row_cells[2].text = str(student.graduate_date)

            else:
                document.add_heading('No hay graduados registrados en el programa', level=5)

            document.add_heading('Maestrantes', level=3)
            if MscStudent.objects.filter(program=program, status='Maestrante'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de ingreso'
                hdr_cells[2].text = 'Requisitos de egreso'

                for student in MscStudent.objects.filter(program=program, status='Maestrante'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.user.get_full_name())
                    row_cells[1].text = str(student.init_date)
                    if finish_requirements_accomplished(student, program):
                        row_cells[2].text = 'Cumplidos'
                    else:
                        row_cells[2].text = 'Incumplidos'
            else:
                document.add_heading('No hay maestrantes registrados en el programa', level=5)
        elif program.type == 'dip':
            if DipStudent.objects.filter(program=program):
                document.add_heading('Estudiantes por estado', level=3)
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Solicitantes'
                hdr_cells[1].text = 'Diplomantes'
                hdr_cells[2].text = 'Graduados'
                hdr_cells[3].text = 'Total'

                row_cells = table.add_row().cells
                row_cells[0].text = str(
                    DipStudent.objects.filter(program=program, status='Solicitante').__len__())
                row_cells[1].text = str(
                    DipStudent.objects.filter(program=program, status='Diplomante').__len__())
                row_cells[2].text = str(
                    DipStudent.objects.filter(program=program, status='Graduado').__len__())
                row_cells[3].text = str(
                    DipStudent.objects.filter(program=program).__len__())
            else:
                document.add_heading('No hay estudiantes registrados en el programa', level=5)

            document.add_heading('Solicitantes de ingreso', level=3)
            if DipStudent.objects.filter(program=program, status='Solicitante'):
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y Apellidos'
                hdr_cells[1].text = 'Fecha de solicitud'


                for student in DipStudent.objects.filter(program=program, status='Solicitante'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.user.get_full_name())
                    row_cells[1].text = str(student.request_date)

            else:
                document.add_heading('No hay solicitantes registrados en el programa', level=5)

            document.add_heading('Graduados', level=3)
            if DipStudent.objects.filter(program=program, status='Graduado'):
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de ingreso'
                hdr_cells[2].text = 'Fecha de egreso'
                hdr_cells[3].text = 'Edición'

                for student in DipStudent.objects.filter(program=program, status='Graduado'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.user.get_full_name())
                    row_cells[1].text = str(student.init_date)
                    row_cells[2].text = str(student.graduate_date)
                    row_cells[3].text = str(student.edition.order)

            else:
                document.add_heading('No hay graduados registrados en el programa', level=5)

            document.add_heading('Diplomantes', level=3)
            if DipStudent.objects.filter(program=program, status='Diplomante'):
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Fecha de ingreso'
                hdr_cells[2].text = 'Edición'

                for student in DipStudent.objects.filter(program=program, status='Diplomante'):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(student.user.get_full_name())
                    row_cells[1].text = str(student.init_date)
                    row_cells[2].text = str(student.edition.order)

            else:
                document.add_heading('No hay diplomantes registrados en el programa', level=5)

        document.add_heading('Claustro', level=3)
        document.add_heading('Generales', level=4)
        if ProgramMember.objects.filter(program=program):
            if program.type == 'phd':
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Miembros'
                hdr_cells[1].text = 'Tutores de doctorandos'
                hdr_cells[2].text = 'Tutores de solicitantes'
                hdr_cells[3].text = 'Tutores (histórico)'

                doc_tuthor = 0
                req_tuthor = 0
                his_tuthor = 0
                members = ProgramMember.objects.filter(program=program).__len__()
                doc_tuthors = []
                req_tuthors = []
                his_tuthors = []
                for member in ProgramMember.objects.filter(program=program):
                    if member.tuthor_set.all():
                        if not member in his_tuthors:
                            his_tuthors.append(member)
                        for tuthor in member.tuthor_set.all():
                            student = tuthor.phd_student
                            if student.status == 'doctorando' and not member in doc_tuthors:
                                doc_tuthors.append(member)

                            elif student.status == 'solicitante' and not member in req_tuthors:
                                req_tuthors.append(member)

                row_cells = table.add_row().cells
                row_cells[0].text = str(members)
                row_cells[1].text = str(doc_tuthors.__len__())
                row_cells[2].text = str(req_tuthors.__len__())
                row_cells[3].text = str(his_tuthors.__len__())


                document.add_heading('Individuales', level=4)

                table = document.add_table(rows=1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Rol'
                hdr_cells[2].text = 'Aspirantes activos'
                hdr_cells[3].text = 'Estado'
                hdr_cells[4].text = 'Año de defensa'
                hdr_cells[5].text = 'Pais'

                for member in ProgramMember.objects.filter(program=program):

                    if member.tuthor_set.filter(Q(phd_student__status = 'solicitante')|Q(phd_student__status = 'doctorando')|Q(phd_student__status = 'graduado')).count()>0:
                        table = document.add_table(rows=1, cols=6)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = member.user.get_full_name()
                        hdr_cells[1].text = str(member.role)
                        first_aspirant = member.tuthor_set.filter(Q(phd_student__status = 'solicitante')|Q(phd_student__status = 'doctorando')|Q(phd_student__status = 'graduado'))[0].phd_student.student

                        hdr_cells[2].text = first_aspirant.user.get_full_name()
                        hdr_cells[3].text = first_aspirant.phdstudent.status[:3].upper()

                        if first_aspirant.phdstudent.status == 'graduado':
                            hdr_cells[4].text = str(first_aspirant.graduate_date.year)
                        else:
                            try:
                                hdr_cells[4].text = str(first_aspirant.studentformationplan.planned_end_year)
                            except StudentFormationPlan.DoesNotExist:
                                hdr_cells[4].text = '------'
                        hdr_cells[5].text = first_aspirant.country


                        for tuthor in member.tuthor_set.filter(Q(phd_student__status = 'solicitante')|Q(phd_student__status = 'doctorando')|Q(phd_student__status = 'graduado'))[1:]:
                            student = tuthor.phd_student.student
                            end_date = ''
                            if student.phdstudent.status == 'graduado':
                                end_date = str(student.graduate_date.year)
                            else:
                                try:
                                    end_date = str(student.studentformationplan.planned_end_year)
                                except StudentFormationPlan.DoesNotExist:
                                    end_date = '------'



                            row_cells = table.add_row().cells
                            row_cells[0].text = ''
                            row_cells[1].text = ''
                            row_cells[2].text = student.user.get_full_name()
                            row_cells[3].text = student.phdstudent.status[:3].upper()
                            row_cells[4].text = end_date
                            row_cells[5].text = student.country



            else:
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Rol'

                for member in ProgramMember.objects.filter(program=program):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(member.user.get_full_name())
                    row_cells[1].text = str(member.role)


        else:
            document.add_heading('No hay profesores registrados en el claustro del programa', level=5)

        docname = 'Reporte_Programa_'+program.slug.upper() + str(now().year) + '_' + str(now().month) + '.docx'
        # docpath = MEDIA_ROOT + '/cgc/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        program_path = MEDIA_ROOT + '/program_{0}'.format(program_slug)
        docpath = MEDIA_ROOT + '/program_{0}/{1}'.format(program_slug, docname)
        try:
            document.save(docpath)
        except:
            try:
                if os.path.isdir(program_path):
                    document.save(docpath)
                else:
                    os.mkdir(program_path)
                    document.save(docpath)

            except:
                return error_500(request, program, "Ha ocurrido un error al intentar guardar el documento solicitado")

        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="'+docname+ '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')

    else:
        return error_500(request,program, 'Solo el Coordinador y el Secretario pueden acceder a la vista de reportes')


@login_required
def docx_thesis_comments(request, program_slug, thesis_id):
    program = Program.objects.get(slug=program_slug)
    thesis = PhdStudentThesis.objects.get(pk=thesis_id)
    if user_is_program_cs(request.user, program):
        document = Document()
        document.add_heading(program.full_name, level=1)
        document.add_heading('Comentarios realizados a '+thesis.phd_student.phdstudenttheme.description, level=2)


        if program.type == 'phd':

            if PhdThesisComment.objects.filter(thesis =thesis):
                table = document.add_table(rows=1, cols=1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '(Fecha, Nombre , Correo): Texto del comentario'

                for comment in PhdThesisComment.objects.filter(thesis =thesis):

                    row_cells = table.add_row().cells
                    row_cells[0].text = '({0}, {1}, {2}): {3}'.format(str(comment.date), str(comment.commenter_name), str(comment.commenter_email),str(comment.text))
            else:
                document.add_heading('No se realizaron comentarios a esta tesis.', level=5)

        elif program.type == 'msc':
            pass
        elif program.type == 'dip':
            pass

        docname = 'Comentarios_{0}_{1}.docx'.format(thesis.phd_student.student.user.first_name, thesis.phd_student.student.user.last_name)
        # docpath = MEDIA_ROOT + '/cgc/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        program_path = MEDIA_ROOT + '/program_{0}'.format(program_slug)
        docpath = MEDIA_ROOT + '/program_{0}/{1}'.format(program_slug, docname)
        try:
            document.save(docpath)
        except:
            try:
                if os.path.isdir(program_path):
                    document.save(docpath)
                else:
                    os.mkdir(program_path)
                    document.save(docpath)

            except:
                return error_500(request, program, "Ha ocurrido un error al intentar guardar el documento solicitado")

        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="'+docname+ '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')

    else:
        return error_500(request,program, 'Solo el Coordinador y el Secretario pueden acceder a la vista de reportes')

@login_required
def print_edition_courses_registers(request, program_slug, edition_id):
    program = Program.objects.get(slug=program_slug)
    edition = ProgramEdition.objects.get(pk=edition_id)
    if user_is_program_cs(request.user, program):
        if program.type == 'msc' or program.type == 'dip':

            document = Document()

            for course in Course.objects.filter(program=program, edition=edition):
                document.add_heading(program.full_name.upper(), level=1)
                document.add_heading('Acta  de ' + course.name, level=2)

                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre y apellidos'
                hdr_cells[1].text = 'Evaluación'
                hdr_cells[2].text = 'Firma'

                if program.type == 'msc':
                    if MscStudent.objects.filter(Q(status='Maestrante')|Q(status='Graduado'),edition=edition):

                        for student in MscStudent.objects.filter(Q(status='Maestrante')|Q(status='Graduado'),edition=edition).order_by('user__last_name'):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(student.user.last_name)+','+student.user.first_name
                            try:
                                row_cells[1].text = str(CourseEvaluation.objects.get(course=course, mscstudent=student))
                            except CourseEvaluation.DoesNotExist:
                                row_cells[1].text = '   '

                            row_cells[2].text = '   '
                    else:
                        document.add_heading('No hay estudiantes en esta edición', level=5)

                elif program.type == 'dip':
                    if DipStudent.objects.filter(Q(status='Diplomante')|Q(status='Graduado'),edition=edition):

                        for student in DipStudent.objects.filter(Q(status='Diplomante')|Q(status='Graduado'), edition=edition).order_by('user__last_name'):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(student.user.last_name)+','+student.user.first_name
                            try:
                                row_cells[1].text = str(CourseEvaluation.objects.get(course=course, dipstudent=student))
                            except CourseEvaluation.DoesNotExist:
                                row_cells[1].text = '   '
                            row_cells[2].text = '   '
                    else:
                        document.add_heading('No hay estudiantes en esta edición', level=5)

                for professor in course.courseprofessor_set.all():
                    document.add_paragraph('')
                    document.add_paragraph('__________________________')
                    document.add_paragraph(professor.professor.user.get_full_name())

                document.add_page_break()



        docname = 'Actas_' + '_' + program.slug.upper() +'_'+str(edition)+'a_ed'+ '.docx'
        # docpath = MEDIA_ROOT + '/cgc/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        docpath = MEDIA_ROOT + '/program_{0}/{1}'.format(program_slug, docname)
        try:
            document.save(docpath)
        except:
            return error_500(request, program,
                             "Ha ocurrido un error al intentar guardar el documento solicitado")
        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="' + docname + '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')
    else:
        return error_500(request, program, 'En este tipo de programas no se puede exportar las actas.')

@login_required
def print_course_register(request, program_slug, course_id):
    program = Program.objects.get(slug=program_slug)
    course = Course.objects.get(pk=course_id)

    if user_is_program_cs(request.user, program):
        if program.type == 'msc' or program.type == 'dip':
            edition = course.edition

            document = Document()
            document.add_heading(program.full_name.upper(), level=1)
            document.add_heading('Acta  de ' + course.name, level=2)

            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nombre y apellidos'
            hdr_cells[1].text = 'Evaluación'
            hdr_cells[2].text = 'Firma'

            if program.type == 'msc':
                if MscStudent.objects.filter(Q(status='Maestrante') | Q(status='Graduado'), edition=edition):

                    for student in MscStudent.objects.filter(Q(status='Maestrante') | Q(status='Graduado'),
                                                             edition=edition).order_by('user__last_name'):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.last_name) + ',' + student.user.first_name
                        try:
                            row_cells[1].text = str(CourseEvaluation.objects.get(course=course, mscstudent=student))
                        except CourseEvaluation.DoesNotExist:
                            row_cells[1].text = '   '

                        row_cells[2].text = '   '
                else:
                    document.add_heading('No hay estudiantes en esta edición', level=5)

            elif program.type == 'dip':
                if DipStudent.objects.filter(Q(status='Diplomante') | Q(status='Graduado'), edition=edition):

                    for student in DipStudent.objects.filter(Q(status='Diplomante') | Q(status='Graduado'),
                                                             edition=edition).order_by('user__last_name'):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(student.user.last_name) + ',' + student.user.first_name
                        try:
                            row_cells[1].text = str(CourseEvaluation.objects.get(course=course, dipstudent=student))
                        except CourseEvaluation.DoesNotExist:
                            row_cells[1].text = '   '
                        row_cells[2].text = '   '
                else:
                    document.add_heading('No hay estudiantes en esta edición', level=5)

            for professor in course.courseprofessor_set.all():
                document.add_paragraph('')
                document.add_paragraph('__________________________')
                document.add_paragraph(professor.professor.user.get_full_name())

            document.add_page_break()

        docname = 'Acta_' +slugify(course.name).upper()+ '_' + program.slug.upper() + '_' + str(edition) + 'a_ed' + '.docx'
        # docpath = MEDIA_ROOT + '/cgc/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        docpath = MEDIA_ROOT + '/program_{0}/{1}'.format(program_slug, docname)
        try:
            document.save(docpath)
        except:
            return error_500(request, program,
                             "Ha ocurrido un error al intentar guardar el documento solicitado")
        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="' + docname + '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')
    else:
        return error_500(request, program, 'En este tipo de programas no se puede exportar las actas.')

@login_required
def print_student_evals(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)

    if program.type == 'phd':
        student=PhdStudent.objects.get(pk=student_id)
        student_name=student.student.user.get_full_name()
    elif program.type == 'msc':
        student=MscStudent.objects.get(pk=student_id)
        student_name=student.user.get_full_name()

    elif program.type == 'dip':
        student=DipStudent.objects.get(pk=student_id)
        student_name=student.user.get_full_name()


    if user_is_program_cs(request.user, program):
        document = Document()

        document.add_heading(program.full_name.upper(), level=1)
        document.add_heading('Evaluaciones de '+student_name, level=2)

        if program.type == 'phd':
            if CourseEvaluation.objects.filter(phdstudent=student, course__program=program):
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre del componente'
                hdr_cells[1].text = 'Evaluación'

                for evaluation in CourseEvaluation.objects.filter(phdstudent=student, course__program=program):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(evaluation.course.name)
                    row_cells[1].text = str(evaluation.value)
            else:
                document.add_heading('No tiene evaluaciones', level=5)

        elif program.type == 'msc':
            if CourseEvaluation.objects.filter(mscstudent=student, course__program=program, course__edition=student.edition):
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre del curso'
                hdr_cells[1].text = 'Evaluación'

                for evaluation in CourseEvaluation.objects.filter(mscstudent=student, course__program=program, course__edition=student.edition):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(evaluation.course.name)
                    row_cells[1].text = str(evaluation.value)
            else:
                document.add_heading('No tiene evaluaciones', level=5)

        elif program.type == 'dip':
            if CourseEvaluation.objects.filter(dipstudent=student, course__program=program, course__edition=student.edition):
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Nombre del curso'
                hdr_cells[1].text = 'Evaluación'

                for evaluation in CourseEvaluation.objects.filter(dipstudent=student, course__program=program, course__edition=student.edition):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(evaluation.course.name)
                    row_cells[1].text = str(evaluation.value)
            else:
                document.add_heading('No tiene evaluaciones', level=5)

        docname = 'Evaluaciones_'+slugify(student_name).upper()+'_'+program.slug.upper() + '.docx'
        # docpath = MEDIA_ROOT + '/cgc/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        docpath = MEDIA_ROOT + '/program_{0}/{1}'.format(program_slug, docname)
        try:
            document.save(docpath)
        except:
            return error_500(request, program, "Ha ocurrido un error al intentar guardar el documento solicitado")
        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="'+docname+ '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')

    else:
        return error_500(request,program, 'Solo el Coordinador y el Secretario pueden acceder a la vista de reportes')

@login_required
def print_student_plan(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)

    if program.type == 'phd':
        student=Student.objects.get(pk=student_id)
        student_name=student.user.get_full_name()

        if user_is_program_cs(request.user, program) or request.user == student.user:
            document = Document()

            document.add_heading(program.full_name.upper(), level=1)
            document.add_heading('Plan de formación de '+student_name, level=2)

            try:
                formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                table = document.add_table(rows=1, cols=3)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Fecha de elaboración'
                hdr_cells[1].text = 'Última actualización'
                hdr_cells[2].text = 'Año de defensa previsto'

                row_cells = table.add_row().cells
                row_cells[0].text = str(formation_plan.elaboration_date)
                row_cells[1].text = str(formation_plan.last_update_date)
                row_cells[2].text = str(formation_plan.planned_end_year)

                document.add_heading('Tareas del plan de formación.' , level=3)

                if formation_plan.formationplanactivities_set.all():
                    table = document.add_table(rows=1, cols=4)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'No.'
                    hdr_cells[1].text = 'Tarea'
                    hdr_cells[2].text = 'Fecha de inicio'
                    hdr_cells[3].text = 'Fecha de fin'
                    index = 1
                    for activity in formation_plan.formationplanactivities_set.all():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(index)
                        row_cells[1].text = str(activity.description)
                        row_cells[2].text = str(activity.init_date)
                        row_cells[3].text = str(activity.end_date)
                        index += 1



                else:
                    document.add_heading('No tiene tareas', level=5)

                document.add_heading('' , level=3)

                if student.phdstudent.tuthor_set.all().count() > 0:
                    for tuthor in student.phdstudent.tuthor_set.all():

                        table = document.add_table(rows=1, cols=2)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Nombre y apellidos del tutor'
                        hdr_cells[1].text = 'Firma del tutor'

                        row_cells = table.add_row().cells
                        row_cells[0].text = 'Dr.C. ' + str(tuthor.professor.user.get_full_name())
                        row_cells[1].text = '                          '

                        document.add_heading('', level=3)



                else:
                    document.add_heading('No tiene tutores designados', level=5)




            except StudentFormationPlan.DoesNotExist:
                return error_500(request, program,'El estudiante no tiene plan de formación.')

        docname = 'Plan_formacion_'+slugify(student_name).upper()+'_'+program.slug.upper() + '.docx'
        # docpath = MEDIA_ROOT + '/cgc/reports/{0}/{1}/{2}'.format(now().year,now().month,docname)
        docpath = MEDIA_ROOT + '/program_{0}/{1}'.format(program_slug, docname)
        try:
            document.save(docpath)
        except:
            return error_500(request, program, "Ha ocurrido un error al intentar guardar el documento solicitado")
        fs = FileSystemStorage()

        filename = docpath

        if fs.exists(filename):

            with fs.open(filename) as docx:
                response = HttpResponse(docx, content_type='application/docx')
                response['Content-Disposition'] = 'attachment; filename="'+docname+ '"'
                return response
        else:
            return error_500(request, 'No se ha encontrado el archivo del reporte correspondiente')

    else:
        return error_500(request,program, 'Solo el Coordinador y el Secretario pueden acceder a la vista de reportes')


# View para crear componente de programa doctoral
@login_required
def create_program_course(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            component = Course(
                program=program,
                name=request.POST['course_name'],
                description=request.POST['course_description'],

            )

            component.save()
            return HttpResponseRedirect(reverse('programs:create_program_course', args=[program_slug]))
        else:
            context={
                'program':program,
                'member': ProgramMember.objects.get(user=request.user, program=program),
            }
            return render(request, 'programs/create_program_course.html', context)
    else:
        return error_500(request,program, 'Usted no tiene privilegios para crear componentes en este programa doctoral')


@login_required
def create_edition_course(request, program_slug, edition_id):
    program = Program.objects.get(slug=program_slug)
    edition = ProgramEdition.objects.get(pk=edition_id)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            component = Course(
                program=program,
                edition=edition,
                name=request.POST['course_name'],
                description=request.POST['course_description'],
                init_date=request.POST['init_date'],
                end_date=request.POST['end_date'],

            )

            component.save()
            for i in range(1,int(request.POST['total_prof'])+1):
                course_professor = CourseProfessor(
                    course=component,
                    professor=ProgramMember.objects.get(program=program, user__email=request.POST['prof_email_'+str(i)])
                )
                course_professor.save()

            return HttpResponseRedirect(reverse('programs:create_edition_course', args=[program_slug, edition_id]))
        else:
            context={
                'program':program,
                'edition':edition,
                'member': ProgramMember.objects.get(user=request.user, program=program),
            }
            if edition.order > 1:
                context['editions'] = ProgramEdition.objects.filter(program=program,order__lt=edition.order)

            return render(request, 'programs/create_edition_course.html', context)
    else:
        return error_500(request,program, 'Usted no tiene privilegios para crear componentes en este programa doctoral')


@login_required
def program_courses(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    context={
        'program':program,
        'components': Course.objects.filter(program=program),
    }
    try:
        context['member'] = ProgramMember.objects.get(user=request.user, program=program)
    except ProgramMember.DoesNotExist:
        try:
            context['student'] = Student.objects.get(user=request.user, program=program)
        except Student.DoesNotExist:
            logout(request)
            raise Http404('No hay profesor o estudiante de este programa con ese usuario')

    return render(request, 'programs/courses_list.html', context)


@login_required
def edition_courses(request, program_slug, edition_id):
    program = Program.objects.get(slug=program_slug)
    edition = ProgramEdition.objects.get(pk=edition_id)
    context={
        'program':program,
        'edition':edition,
        'components': Course.objects.filter(program=program, edition=edition),
    }
    try:
        context['member'] = ProgramMember.objects.get(user=request.user, program=program)
    except ProgramMember.DoesNotExist:
        try:
            context['student'] = Student.objects.get(user=request.user, program=program)
        except Student.DoesNotExist:
            logout(request)
            raise Http404('No hay profesor o estudiante de este programa con ese usuario')

    return render(request, 'programs/courses_list.html', context)

@login_required
def edit_program_course(request, program_slug, course_id):
    program = Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            if program.type== 'phd':
                Course.objects.filter(pk=course_id).update(
                    name=request.POST['course_name'],
                    description=request.POST['course_description'],

                )

                return HttpResponseRedirect(reverse('programs:program_courses', args=[program_slug]))
            else:
                Course.objects.filter(pk=course_id).update(
                    name=request.POST['course_name'],
                    description=request.POST['course_description'],
                    init_date = request.POST['init_date'],
                    end_date = request.POST['end_date'],

                )
                if int(request.POST['total_new_prof']) > 0:
                    for i in range(1,int(request.POST['total_new_prof'])+1):
                        new_professor = CourseProfessor(
                            course=Course.objects.get(pk=course_id),
                            professor=ProgramMember.objects.get(program=program, user__email=request.POST['new_prof_email_'+str(i)])
                        )
                        new_professor.save()


                return HttpResponseRedirect(reverse('programs:edition_courses', args=[program_slug,Course.objects.get(pk=course_id).edition.id ]))

        else:
            if program.type == 'phd':
                context={
                    'program':program,
                    'member': ProgramMember.objects.get(user=request.user, program=program),
                    'component': Course.objects.get(pk=course_id),
                }
            else:
                context = {
                    'program': program,
                    'member': ProgramMember.objects.get(user=request.user, program=program),
                    'component': Course.objects.get(pk=course_id),
                    'edition': Course.objects.get(pk=course_id).edition,
                }

            return render(request, 'programs/edit_program_course.html', context)
    else:
        return error_500(request, program, 'Usted no tiene privilegios para crear componentes en este programa doctoral')


@login_required
def evaluate_student(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            evaluation = CourseEvaluation(
                course=Course.objects.get(pk=request.POST['course']),
                value=request.POST['eval'],
            )
            if program.type == 'phd':
                try:
                    evaluation.phdstudent = PhdStudent.objects.get(pk=student_id)
                except PhdStudent.DoesNotExist:
                    return error_500(request, program, 'No existe el estudiante de doctorado a evaluar.')
            elif program.type == 'msc':
                try:
                    evaluation.mscstudent = MscStudent.objects.get(pk=student_id)
                except MscStudent.DoesNotExist:
                    return error_500(request, program, 'No existe el estudiante de maestría a evaluar.')
            elif program.type == 'dip':
                try:
                    evaluation.dipstudent = DipStudent.objects.get(pk=student_id)
                except DipStudent.DoesNotExist:
                    return error_500(request, program, 'No existe el estudiante de diplomado a evaluar.')
            evaluation.save()

            return  HttpResponseRedirect(reverse('programs:evaluate_student', args=[program_slug, student_id]))
        else:
            if program.type == 'phd':
                student=PhdStudent.objects.get(pk=student_id)
            elif program.type == 'msc':
                student = MscStudent.objects.get(pk=student_id)
            elif program.type == 'dip':
                student = DipStudent.objects.get(pk=student_id)

            pending_courses=[]

            if program.type == 'phd':
                for course in Course.objects.filter(program=program):
                    try:
                        evaluation=CourseEvaluation.objects.get(course=course, phdstudent=student)
                    except CourseEvaluation.DoesNotExist:
                        pending_courses.append(course)
            elif program.type == 'msc':
                for course in Course.objects.filter(program=program, edition=student.edition):
                    try:
                        evaluation=CourseEvaluation.objects.get(course=course, mscstudent=student)
                    except CourseEvaluation.DoesNotExist:
                        pending_courses.append(course)
            elif program.type == 'dip':
                for course in Course.objects.filter(program=program, edition=student.edition):
                    try:
                        evaluation=CourseEvaluation.objects.get(course=course, dipstudent=student)
                    except CourseEvaluation.DoesNotExist:
                        pending_courses.append(course)


            context={
                'program':program,
                'pending_courses':pending_courses,
                'student':student,
            }
            return render(request, 'programs/evaluate_student_course.html', context)
    else:
        return error_500(request, program, 'Usted no tiene privilegios para asignar evaluaciones a estudiantes en este programa')


@login_required
def student_evals(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    if program.type == 'phd':
        if user_is_program_member(request.user, program) or PhdStudent.objects.get(pk=student_id).student.user == request.user:
            context={
                'program': program,
                'student': PhdStudent.objects.get(pk=student_id),
                'evals': PhdStudent.objects.get(pk=student_id).courseevaluation_set.all(),
            }
            if user_is_program_member(request.user, program):
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            return render(request, 'programs/student_evals.html', context)
        else:
            return error_500(request, program, 'Usted no puede ver las evaluaciones de estudiantes')
    elif program.type == 'msc':
        if user_is_program_member(request.user, program) or MscStudent.objects.get(pk=student_id).user == request.user:
            context={
                'program': program,
                'student': MscStudent.objects.get(pk=student_id),
                'evals': MscStudent.objects.get(pk=student_id).courseevaluation_set.all(),
            }
            if user_is_program_member(request.user, program):
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            return render(request, 'programs/student_evals.html', context)
        else:
            return error_500(request, program, 'Usted no puede ver las evaluaciones de estudiantes')
    elif program.type == 'dip':
        if user_is_program_member(request.user, program) or DipStudent.objects.get(pk=student_id).user == request.user:
            context={
                'program': program,
                'student': DipStudent.objects.get(pk=student_id),
                'evals': DipStudent.objects.get(pk=student_id).courseevaluation_set.all(),
            }
            if user_is_program_member(request.user, program):
                context['member']=ProgramMember.objects.get(user=request.user, program=program)
            return render(request, 'programs/student_evals.html', context)
        else:
            return error_500(request, program, 'Usted no puede ver las evaluaciones de estudiantes')
    else:
        pass


@login_required
def create_formation_plan(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)
    if program.type == 'phd':
        if request.user == student.user:
            if request.method == 'POST':

                new_plan = StudentFormationPlan(
                    phdstudent=student,
                    elaboration_date=now().date(),
                    last_update_date=now().date(),
                    planned_end_year=request.POST['planned_end_year'],
                )
                new_plan.save()
                return HttpResponseRedirect(reverse('programs:edit_formation_plan', args=[program_slug, student_id]))

            else:
                try:
                    formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
                    return HttpResponseRedirect(
                        reverse('programs:edit_formation_plan', args=[program_slug, student_id]))
                except StudentFormationPlan.DoesNotExist:
                    context = {
                        'program': program,
                        'student': student,
                        'years': range(now().year, now().year + 6)
                    }
                    return render(request, 'programs/create_formation_plan.html', context)
        else:
            return error_500(request, program,
                             'Usted no tiene privilegios para crear el plan de formación de este estudiante.')
    else:
        return error_500(request, program,
                         'El programa no es un doctorado, los estudiantes no tienen plan de formación.')

@login_required
def edit_formation_plan(request, program_slug, student_id):
    program = Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)

    try:
        formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
        if request.user == student.user:
            if request.method == 'POST':
                formation_plan.elaboration_date = request.POST['elaboration_date']
                formation_plan.last_update_date = now().date()
                formation_plan.planned_end_year = request.POST['planned_end_year']

                formation_plan.save()

                return HttpResponseRedirect(reverse('programs:home', args=[program_slug]))
            else:
                context = {
                    'program': program,
                    'student': student,
                    'years': range(now().year, now().year + 6),
                    'formation_plan': formation_plan,

                }
                return render(request, 'programs/edit_formation_plan.html', context)
        else:
            return error_500(request, program,
                             'Usted no tiene privilegios para editar el plan de formación de este estudiante.')
    except StudentFormationPlan.DoesNotExist:
        return HttpResponseRedirect(reverse('programs:create_formation_plan', args=[program_slug, student_id]))

@login_required
def view_formation_plan(request, program_slug,student_id):
    program = Program.objects.get(slug=program_slug)
    student = Student.objects.get(pk=student_id)
    try:
        formation_plan = StudentFormationPlan.objects.get(phdstudent=student)
        context = {
            'program': program,
            'years': range(now().year, now().year + 6),
            'formation_plan': formation_plan,

        }
        if user_is_program_member(request.user, program):
            context['member']=ProgramMember.objects.get(user=request.user, program=program)
        elif user_is_program_student(request.user, program):
            context['student']=student
        else:
            return error_500(request, program, 'Usted no tiene acceso a este plan de formación')

        return render(request, 'programs/view_formation_plan.html', context)
    except StudentFormationPlan.DoesNotExist:
        return HttpResponseRedirect(reverse('programs:create_formation_plan', args=[program_slug, student_id]))

@login_required
def create_new(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            new=New(
                program=program,
                title=request.POST['new_title'],
                body=request.POST['new_body'],
            )
            try:
                new.img = request.FILES['img_new']
            except:
                pass
            new.save()
            return HttpResponseRedirect(reverse('programs:news_list', args=[program_slug]))
        else:
            context = {
                'program':program,
                'member': ProgramMember.objects.get(user=request.user, program=program)
            }
            return render(request,'programs/create_new.html', context)
    else:
        return error_500(request,program,'Usted no tiene provilegios para agregar noticias en este programa')

@login_required
def news_list(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    context={
        'program':program,
        'news': New.objects.filter(program=program),
    }
    if user_is_program_member(request.user, program):
        context['member']=ProgramMember.objects.get(user=request.user, program=program)
    elif user_is_program_student(request.user, program):
        if program.type == 'phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type == 'msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type == 'dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)

    return render(request, 'programs/news_list.html', context)

@login_required
def read_new(request, program_slug, new_id):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_member( request.user, program) or user_is_program_student(request.user, program):
        context={
            'program':program,
            'new': New.objects.get(pk=new_id)
        }
        if user_is_program_member(request.user, program):
            context['member']=ProgramMember.objects.get(user=request.user, program=program)
        elif user_is_program_student(request.user, program):
            if program.type == 'phd':
                context['student']=Student.objects.get(user=request.user, program=program)
            elif program.type == 'msc':
                context['student'] = MscStudent.objects.get(user=request.user, program=program)
            elif program.type == 'dip':
                context['student'] = DipStudent.objects.get(user=request.user, program=program)

        return render(request, 'programs/read_new.html',context)


@login_required
def edit_new(request, program_slug, new_id):
    program = Program.objects.get(slug=program_slug)
    new = New.objects.get(pk=new_id)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            new.title = request.POST['new_title']
            new.body = request.POST['new_body']
            try:
                new.img = request.FILES['img_new']
            except:
                pass
            new.save()
            return HttpResponseRedirect(reverse('programs:read_new', args=[program_slug, new_id]))
        else:
            context={
                'program':program,
                'new':new,
                'member':ProgramMember.objects.get(user=request.user, program=program)
            }
            return render(request, 'programs/edit_new.html', context)
    else:
        return error_500(request, program, 'Usted no tiene privilegios para editar noticias en este programa')

@login_required
def edit_program(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            program.full_name = request.POST['program_fullname']
            program.short_name = request.POST['program_shortname']
            program.email = request.POST['program_email']
            program.phone = request.POST['program_phone']
            program.address = request.POST['program_address']
            program.center = request.POST['program_center']
            program.branch = request.POST['program_branch']
            program.code = request.POST['program_code']
            try:
                if request.POST['is_self_request'] == 'True':
                    program.self_request = True
                else:
                    program.self_request = False
            except:
                program.self_request = False

            program.save()

            return HttpResponseRedirect(reverse('programs:edit_program', args=[program_slug]))
        else:
            context = {
                'program': program,
                'program_init_requirenments': ProgramFileDoc.objects.filter(program=program, is_init_requirenment=True),
                'program_finish_requirenments': ProgramFileDoc.objects.filter(program=program, is_finish_requirenment=True),
                'member': ProgramMember.objects.get(user=request.user, program=program)
            }
            return render(request,'programs/edit_program.html', context)
    else:
        return  error_500(request, program, 'Usted no tiene privilegios para editar los ajustes de este programa')

@login_required
def export_csv_students(request, program_slug, scope):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if program.type == 'phd':
            if scope == "all":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Todos.csv"'
                writer = csv.writer(response)
                writer.writerow(["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                                 "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                                 "Forma de Posgrado: DrC, MSc, EPG, DIP",
                                 "Nombre del Programa",
                                 "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                                 "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                                 "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i=1
                for student in Student.objects.filter(program=program):
                    try:
                        options = [str(i), student.dni, student.user.first_name,
                                   student.user.last_name.split(" ")[0],
                                   student.user.last_name.split(" ")[1], str(student.gender).upper(),
                                   "MES", str(student.country)[:2],
                                   "DrC",
                                   program,
                                   program.code,
                                   program.branch, "N",
                                   "N"]
                        if student.phdstudent.status == 'graduado':
                            options.append('S')
                        else:
                            options.append('N')
                        writer.writerow(options)
                    except:
                        options = [str(i), student.dni, student.user.first_name,
                                   student.user.last_name.split(" ")[0],
                                   '--------', str(student.gender).upper(),
                                   "MES", str(student.country)[:2],
                                   "DrC",
                                   program,
                                   program.code,
                                   program.branch, "N",
                                   "N"]
                        if student.phdstudent.status == 'graduado':
                            options.append('S')
                        else:
                            options.append('N')
                        writer.writerow(options)
                    i+=1
                return response
            elif scope == "requesters":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Todos.csv"'
                writer = csv.writer(response)
                writer.writerow(['', 'Listado de Solicitantes, Doctorandos, Graduados y Rechazados en ' + program.short_name])
                writer.writerow(["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                                 "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                                 "Forma de Posgrado: DrC, MSc, EPG, DIP",
                                 "Nombre del Programa",
                                 "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                                 "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                                 "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in PhdStudent.objects.filter(student__program=program, status="solicitante").order_by('student__request_date'):
                    try:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name, student.student.user.last_name.split(" ")[0],
                             student.student.user.last_name.split(" ")[1], str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    except:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             "--------", str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    i += 1
                return response

            elif scope == "aproved":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Doctorandos.csv"'
                writer = csv.writer(response)
                writer.writerow(["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                                 "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                                 "Forma de Posgrado: DrC, MSc, EPG, DIP",
                                 "Nombre del Programa",
                                 "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                                 "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                                 "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in PhdStudent.objects.filter(student__program=program, status="doctorando").order_by('student__init_date'):
                    try:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             student.student.user.last_name.split(" ")[1], str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    except:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             "--------", str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    i += 1
                return response

            elif scope == "graduated":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Graduados.csv"'
                writer = csv.writer(response)
                writer.writerow(["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                                 "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                                 "Forma de Posgrado: DrC, MSc, EPG, DIP",
                                 "Nombre del Programa",
                                 "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                                 "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                                 "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in PhdStudent.objects.filter(student__program=program, status="graduado").order_by('student__graduate_date'):
                    try:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             student.student.user.last_name.split(" ")[1], str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "S"])
                    except:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             "--------", str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "S"])
                    i += 1
                return response

            elif scope == "denied":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Denegados.csv"'
                writer = csv.writer(response)
                writer.writerow(["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                                 "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                                 "Forma de Posgrado: DrC, MSc, EPG, DIP",
                                 "Nombre del Programa",
                                 "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                                 "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                                 "Ejecutado en otro País: S o N", "Graduado: S o N"])

                i = 1
                for student in PhdStudent.objects.filter(student__program=program, status="denegado").order_by('student__request_date'):
                    try:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             student.student.user.last_name.split(" ")[1], str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    except:
                        writer.writerow(
                            [str(i), student.student.dni, student.student.user.first_name,
                             student.student.user.last_name.split(" ")[0],
                             "--------", str(student.student.gender).upper(),
                             "MES", str(student.student.country)[:2],
                             "DrC",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    i += 1

                return response

            else:
                return error_500(request, program,
                                 "No se reconoce el contexto "+scope)
        elif program.type == 'msc':
            if scope == "all":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Todos.csv"'
                writer = csv.writer(response)
                writer.writerow(
                    ["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                     "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                     "Forma de Posgrado: DrC, MSc, EPG, DIP",
                     "Nombre del Programa",
                     "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                     "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                     "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in MscStudent.objects.filter(program=program).order_by('request_date'):
                    try:
                        options = [str(i), student.dni, student.user.first_name,
                                   student.user.last_name.split(" ")[0],
                                   student.user.last_name.split(" ")[1], str(student.gender).upper(),
                                   "MES", str(student.country)[:2],
                                   "MSc",
                                   program,
                                   program.code,
                                   program.branch, "N",
                                   "N"]
                        if student.status == 'graduado':
                            options.append('S')
                        else:
                            options.append('N')
                        writer.writerow(options)
                    except:
                        options = [str(i), student.dni, student.user.first_name,
                                   student.user.last_name.split(" ")[0],
                                   '------', str(student.gender).upper(),
                                   "MES", str(student.country)[:2],
                                   "MSc",
                                   program,
                                   program.code,
                                   program.branch, "N",
                                   "N"]
                        if student.status == 'graduado':
                            options.append('S')
                        else:
                            options.append('N')
                        writer.writerow(options)
                    i += 1
                return response
            elif scope == "requesters":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Todos.csv"'
                writer = csv.writer(response)
                writer.writerow(
                    ['', 'Listado de Solicitantes, Doctorandos, Graduados y Rechazados en ' + program.short_name])
                writer.writerow(
                    ["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                     "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                     "Forma de Posgrado: DrC, MSc, EPG, DIP",
                     "Nombre del Programa",
                     "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                     "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                     "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in MscStudent.objects.filter(program=program, status = 'solicitante').order_by('request_date'):
                    try:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             student.user.last_name.split(" ")[1], str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    except:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             "-------", str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    i += 1
                return response

            elif scope == "aproved":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Doctorandos.csv"'
                writer = csv.writer(response)
                writer.writerow(
                    ["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                     "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                     "Forma de Posgrado: DrC, MSc, EPG, DIP",
                     "Nombre del Programa",
                     "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                     "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                     "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in MscStudent.objects.filter(program=program, status = 'maestrante').order_by('init_date'):
                    try:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             student.user.last_name.split(" ")[1], str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    except:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             "-------", str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    i += 1
                return response

            elif scope == "graduated":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Graduados.csv"'
                writer = csv.writer(response)
                writer.writerow(
                    ["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                     "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                     "Forma de Posgrado: DrC, MSc, EPG, DIP",
                     "Nombre del Programa",
                     "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                     "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                     "Ejecutado en otro País: S o N", "Graduado: S o N"])
                i = 1
                for student in MscStudent.objects.filter(program=program, status = 'graduado').order_by('graduate_date'):
                    try:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             student.user.last_name.split(" ")[1], str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "S"])
                    except:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             "-------", str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "S"])
                    i += 1
                return response

            elif scope == "denied":
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename="Denegados.csv"'
                writer = csv.writer(response)
                writer.writerow(
                    ["Consecutivo", "Carné de Identidad", "Nombres", "Apellido 1", "Apellido 2", "Sexo: F o M",
                     "Sigla del Organismo", "Sigla del País: según Codificador de Países",
                     "Forma de Posgrado: DrC, MSc, EPG, DIP",
                     "Nombre del Programa",
                     "Código del Programa: según Codificador de Maestrías y Especialidades de Posgrado",
                     "Rama del Programa: CT, CNE, CBM, CA, CE, CSH, CP, CCF, ARTE", "Ejecutado en CUM: S o N",
                     "Ejecutado en otro País: S o N", "Graduado: S o N"])

                i = 1
                for student in MscStudent.objects.filter(program=program, status = 'denegado').order_by('request_date'):
                    try:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             student.user.last_name.split(" ")[1], str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    except:
                        writer.writerow(
                            [str(i), student.dni, student.user.first_name,
                             student.user.last_name.split(" ")[0],
                             "-------", str(student.gender).upper(),
                             "MES", str(student.country)[:2],
                             "MSc",
                             program,
                             program.code,
                             program.branch, "N",
                             "N", "N"])
                    i += 1

                return response

            else:
                return error_500(request, program,
                                 "No se reconoce el contexto " + scope)



    else:
        return error_500(request, program, "Usted no tiene privilegios para exportar listado de estudiantes del programa")


@login_required
def download_evidences(request, program_slug):

    try:

        program = Program.objects.get(slug=program_slug)

        fpath = "{0}/users/{1}/downloads".format(MEDIA_ROOT, program.slug)

        if not os.path.exists(fpath):
            os.makedirs(fpath)


        if program.type == 'phd':
            zpath = "{0}/users/{1}/downloads/Evidences.zip".format(MEDIA_ROOT, program_slug)

            zf = zipfile.ZipFile(zpath, "w")

            for student in program.student_set.all():
                for requirement in ProgramFileDoc.objects.filter(program=program):
                    try:
                        s_f_d = StudentFileDocument.objects.get(student=student, program_file_document=requirement)
                        if s_f_d.file:
                            print(s_f_d.file.path)
                            if os.path.exists(s_f_d.file.path):
                                fpath = s_f_d.file.path
                                fdir,fname = os.path.split(fpath)
                                zip_subdir = "{0}/{1}".format(student.user.get_full_name(), slugify(requirement))
                                zip_path = os.path.join(zip_subdir, fname)
                                zf.write(fpath, zip_path)

                        else:
                            print('No hay archivo')
                    except StudentFileDocument.DoesNotExist:
                        print("Creando documento de estudiante")
                        s_f_d = StudentFileDocument(
                            student = student,
                            program_file_document = requirement
                        )
                        s_f_d.save()

            zf.close()

        fs = FileSystemStorage()

        with fs.open(zpath) as zip:
            # response = HttpResponse(pdf, content_type='application/pdf')
            # response['Content-Disposition'] = "inline; filename=" + '"' + filename + '"'
            # Grab ZIP file from in-memory, make response with correct MIME-type
            resp = HttpResponse(zip, content_type="application/x-zip-compressed")
            # ..and correct content-disposition
            resp['Content-Disposition'] = 'attachment; filename=%s' % '{0}_Evidencias_Estudiantes.zip'.format(program_slug)

        return resp
    except Program.DoesNotExist:
        messages.error(request, 'El programa no existe')
        return HttpResponse('El programa no existe')
    

    # Gestion de preguntas frecuentes

@login_required
def create_faq(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            faq=FAQ(
                program=program,
                title=request.POST['faq.title'],
                content=request.POST['faq_content'],
            )
            try:
                faq.img = request.FILES['img_faq']
            except:
                pass
            faq.save()
            return HttpResponseRedirect(reverse('programs:faq_list', args=[program_slug]))
        else:
            context = {
                'program':program,
                'member': ProgramMember.objects.get(user=request.user, program=program)
            }
            return render(request,'programs/create_faq.html', context)
    else:
        return error_500(request,program,'Usted no tiene provilegios para agregar noticias en este programa')

@login_required
def faq_list(request, program_slug):
    program = Program.objects.get(slug=program_slug)
    context={
        'program':program,
        'faqs': FAQ.objects.filter(program=program),
    }
    if user_is_program_member(request.user, program):
        context['member']=ProgramMember.objects.get(user=request.user, program=program)
    elif user_is_program_student(request.user, program):
        if program.type == 'phd':
            context['student'] = Student.objects.get(user=request.user, program=program)
        elif program.type == 'msc':
            context['student'] = MscStudent.objects.get(user=request.user, program=program)
        elif program.type == 'dip':
            context['student'] = DipStudent.objects.get(user=request.user, program=program)

    return render(request, 'programs/faq_list.html', context)

@login_required
def read_faq(request, program_slug, faq_id):
    program = Program.objects.get(slug=program_slug)
    if user_is_program_member( request.user, program) or user_is_program_student(request.user, program):
        context={
            'program':program,
            'faq': FAQ.objects.get(pk=faq_id)
        }
        if user_is_program_member(request.user, program):
            context['member']=ProgramMember.objects.get(user=request.user, program=program)
        elif user_is_program_student(request.user, program):
            if program.type == 'phd':
                context['student']=Student.objects.get(user=request.user, program=program)
            elif program.type == 'msc':
                context['student'] = MscStudent.objects.get(user=request.user, program=program)
            elif program.type == 'dip':
                context['student'] = DipStudent.objects.get(user=request.user, program=program)

        return render(request, 'programs/read_faq.html',context)


@login_required
def edit_faq(request, program_slug, faq_id):
    program = Program.objects.get(slug=program_slug)
    faq = FAQ.objects.get(pk=faq_id)
    if user_is_program_cs(request.user, program):
        if request.method == 'POST':
            faq.title = request.POST['faq.title']
            faq.content = request.POST['faq.content']
            try:
                faq.img = request.FILES['img_faq']
            except:
                pass
            faq.save()
            return HttpResponseRedirect(reverse('programs:read_faq', args=[program_slug, faq_id]))
        else:
            context={
                'program':program,
                'faq':faq,
                'member':ProgramMember.objects.get(user=request.user, program=program)
            }
            return render(request, 'programs/edit_faq.html', context)
    else:
        return error_500(request, program, 'Usted no tiene privilegios para editar noticias en este programa')
    
@login_required
def ajx_delete_faq(request, program_slug):
    program=Program.objects.get(slug=program_slug)

    if user_is_program_cs(request.user,program ):
        if request.method=='POST':
            faq_id=request.POST['faq_id']
            try:
                FAQ.objects.get(pk=faq_id).delete()
                return HttpResponse(
                    json.dumps([{'deleted': 1}]),
                    content_type="application/json"
                )
            except:
                return HttpResponse(
                    json.dumps([{'deleted': 0}]),
                    content_type="application/json"
                )
        else:
            return HttpResponse(
                json.dumps([{'deleted': 2}]),
                content_type="application/json"
            )
    else:
        return HttpResponse(
            json.dumps([{'deleted': 3}]),
            content_type="application/json"
        )
    
def msc_index(request, program_slug=None):
    """
    Vista para la página principal de maestrías
    program_slug: Parámetro que identifica el programa (ej: 'educacion-superior')
    """
    context = {
        'program_slug': program_slug,
        # Agrega aquí cualquier otro contexto necesario
    }
    return render(request, 'msc_index.html', context)


def dip_index(request, program_slug=None):
    """
    Vista para la página principal de maestrías
    program_slug: Parámetro que identifica el programa (ej: 'educacion-superior')
    """
    context = {
        'program_slug': program_slug,
        # Agrega aquí cualquier otro contexto necesario
    }
    return render(request, 'dip_index.html', context)




